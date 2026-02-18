"""
AI求职助手 - Streamlit 完整版
整合所有功能：简历分析 + 岗位推荐 + 自动投递 + 投递记录 + 数据统计
"""
import streamlit as st
import sys
import os
import asyncio
import io
import requests
import pandas as pd
import time
from datetime import datetime
import uuid

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(__file__))

st.set_page_config(
    page_title="AI求职助手",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 全局样式 - Gemini 极简风格
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Google+Sans:wght@400;500;700&family=Noto+Sans+SC:wght@400;500;700&display=swap');

:root {
    --gemini-blue: #1a73e8;
    --gemini-blue-hover: #1557b0;
    --text: #1f1f1f;
    --text-light: #5f6368;
    --border: #e8eaed;
    --bg: #ffffff;
    --bg-hover: #f8f9fa;
}

* {
    font-family: 'Google Sans', 'Noto Sans SC', sans-serif;
}

#MainMenu, footer, header {visibility: hidden}

.main .block-container {
    max-width: 1000px;
    padding: 1rem 1.5rem 3rem;
}

/* 顶部 Logo */
.logo {
    font-size: 1.375rem;
    font-weight: 500;
    color: var(--text);
    padding: 1rem 0;
    margin-bottom: 2rem;
}

/* Hero */
.hero {
    margin-bottom: 3rem;
}

.hero h1 {
    font-size: 2.75rem;
    font-weight: 400;
    color: var(--text);
    line-height: 1.3;
    margin-bottom: 0.75rem;
}

.hero p {
    font-size: 1rem;
    color: var(--text-light);
    line-height: 1.5;
}

/* 按钮 */
.stButton > button {
    background: var(--gemini-blue);
    color: white;
    border: none;
    border-radius: 24px;
    padding: 0.625rem 1.5rem;
    font-size: 0.875rem;
    font-weight: 500;
    transition: background 0.2s;
}

.stButton > button:hover {
    background: var(--gemini-blue-hover);
}

/* 输入框 */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 0.75rem;
    font-size: 0.875rem;
    transition: border 0.2s;
}

.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: var(--gemini-blue);
    outline: none;
}

.stTextArea > div > div > textarea {
    min-height: 180px;
}

/* 标签页 */
.stTabs [data-baseweb="tab-list"] {
    gap: 0;
    border-bottom: 1px solid var(--border);
}

.stTabs [data-baseweb="tab"] {
    padding: 0.75rem 1rem;
    font-size: 0.875rem;
    font-weight: 500;
    color: var(--text-light);
    border: none;
}

.stTabs [aria-selected="true"] {
    color: var(--gemini-blue);
    border-bottom: 2px solid var(--gemini-blue);
}

/* 岗位卡片 */
.job-card {
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 1rem;
    margin-bottom: 1rem;
    transition: box-shadow 0.2s;
}

.job-card:hover {
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}

/* 指标卡片 */
.metric-card {
    background: var(--bg-hover);
    border-radius: 8px;
    padding: 1rem;
    text-align: center;
}

.metric-value {
    font-size: 2rem;
    font-weight: 700;
    color: var(--gemini-blue);
}

.metric-label {
    font-size: 0.875rem;
    color: var(--text-light);
    margin-top: 0.5rem;
}
</style>
""", unsafe_allow_html=True)

# 配置 API Key
os.environ['OPENAI_API_KEY'] = 'sk-SnQQxqPPxqxqxqxqxqxqxqxqxqxqxqxqxqxqxqxqxqxqxqxq'
os.environ['OPENAI_BASE_URL'] = 'https://oneapi.gemiaude.com/v1'

# 后端 API 地址
BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8000")

# 文件解析函数
def parse_uploaded_file(uploaded_file):
    """解析上传的文件 - 支持 PDF/Word/图片（OCR）"""
    try:
        file_content = uploaded_file.read()
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        resume_text = ""

        if file_ext == '.txt':
            try:
                resume_text = file_content.decode('utf-8')
            except:
                try:
                    resume_text = file_content.decode('gbk', errors='ignore')
                except:
                    resume_text = file_content.decode('latin-1', errors='ignore')

        elif file_ext == '.pdf':
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                if len(pdf_reader.pages) == 0:
                    st.error("PDF 文件为空")
                    return None
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        resume_text += text + "\n"
                if not resume_text.strip():
                    st.warning("PDF 可能是扫描件，尝试使用图片上传方式")
                    return None
            except Exception as e:
                st.error(f"PDF 解析失败: {str(e)}")
                return None

        elif file_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_content))
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        resume_text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                resume_text += cell.text + " "
                        resume_text += "\n"
                if not resume_text.strip():
                    st.error("Word 文档为空或无法提取文字")
                    return None
            except Exception as e:
                st.error(f"Word 文档解析失败: {str(e)}")
                return None

        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
            try:
                from PIL import Image
                import pytesseract
                image = Image.open(io.BytesIO(file_content))
                st.image(image, caption="上传的图片", use_container_width=True)
                with st.spinner("识别图片中的文字..."):
                    resume_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                if not resume_text.strip():
                    st.error("图片识别失败，未能提取到文字")
                    return None
            except ImportError:
                st.error("图片 OCR 功能未安装")
                st.info("请安装: pip install pytesseract")
                return None
            except Exception as e:
                st.error(f"图片识别失败: {str(e)}")
                return None

        else:
            st.error(f"不支持的文件格式: {file_ext}")
            return None

        if resume_text and len(resume_text.strip()) < 50:
            st.warning("提取的文字内容较少，可能影响分析质量")

        return resume_text.strip() if resume_text else None

    except Exception as e:
        st.error(f"文件解析失败: {str(e)}")
        return None

# 异步函数包装器
def run_async(coro):
    """运行异步函数的同步包装器"""
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        result = loop.run_until_complete(coro)
        loop.close()
        return result
    except Exception as e:
        st.error(f"执行出错: {str(e)}")
        return None

# 初始化 session state
if 'user_id' not in st.session_state:
    st.session_state.user_id = str(uuid.uuid4())
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'current_jobs' not in st.session_state:
    st.session_state.current_jobs = []
if 'applications' not in st.session_state:
    st.session_state.applications = []

# Logo
st.markdown('<div class="logo">AI 求职助手</div>', unsafe_allow_html=True)

# Hero
st.markdown('''
<div class="hero">
    <h1>让 AI 帮你找到理想工作</h1>
    <p>分析简历，推荐岗位，自动投递，追踪进度</p>
</div>
''', unsafe_allow_html=True)

# 标签页
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "简历分析",
    "岗位推荐",
    "自动投递",
    "投递记录",
    "数据统计"
])

# ==================== Tab1: 简历分析 ====================
with tab1:
    st.markdown("### 简历分析")

    method = st.radio("", ["文本输入", "上传文件"], horizontal=True, label_visibility="collapsed")

    if method == "文本输入":
        resume_text = st.text_area(
            "",
            height=180,
            placeholder="粘贴你的简历内容...",
            label_visibility="collapsed"
        )

        if resume_text and st.button("分析", type="primary", key="analyze_text"):
            if len(resume_text.strip()) < 50:
                st.warning("简历内容较少，建议至少 50 字以上")
            else:
                with st.spinner("分析中..."):
                    try:
                        from app.core.multi_ai_debate import JobApplicationPipeline
                        pipeline = JobApplicationPipeline()
                        results = run_async(pipeline.process_resume(resume_text))

                        if results:
                            st.session_state.analysis_results = results
                            st.success("完成")

                            result_tabs = st.tabs(["职业分析", "岗位推荐", "简历优化", "面试准备", "模拟面试", "技能分析"])

                            with result_tabs[0]:
                                st.markdown(results.get('career_analysis', '暂无数据'))
                            with result_tabs[1]:
                                st.markdown(results.get('job_recommendations', '暂无数据'))
                            with result_tabs[2]:
                                st.markdown(results.get('resume_optimization', '暂无数据'))
                            with result_tabs[3]:
                                st.markdown(results.get('interview_preparation', '暂无数据'))
                            with result_tabs[4]:
                                st.markdown(results.get('mock_interview', '暂无数据'))
                            with result_tabs[5]:
                                st.markdown(results.get('skill_gap_analysis', '暂无数据'))

                    except Exception as e:
                        st.error(f"分析失败: {str(e)}")

    else:  # 上传文件
        uploaded_file = st.file_uploader(
            "",
            type=["pdf", "doc", "docx", "png", "jpg", "jpeg", "txt"],
            label_visibility="collapsed"
        )

        if uploaded_file:
            if st.button("分析", type="primary", key="analyze_file"):
                with st.spinner("解析中..."):
                    resume_text = parse_uploaded_file(uploaded_file)

                if resume_text:
                    with st.spinner("分析中..."):
                        try:
                            from app.core.multi_ai_debate import JobApplicationPipeline
                            pipeline = JobApplicationPipeline()
                            results = run_async(pipeline.process_resume(resume_text))

                            if results:
                                st.session_state.analysis_results = results
                                st.success("完成")

                                with st.expander("职业分析", expanded=True):
                                    st.write(results.get('career_analysis', '暂无数据'))
                                with st.expander("岗位推荐"):
                                    st.write(results.get('job_recommendations', '暂无数据'))
                                with st.expander("简历优化"):
                                    st.write(results.get('resume_optimization', '暂无数据'))
                                with st.expander("面试准备"):
                                    st.write(results.get('interview_preparation', '暂无数据'))
                                with st.expander("模拟面试"):
                                    st.write(results.get('mock_interview', '暂无数据'))
                                with st.expander("技能分析"):
                                    st.write(results.get('skill_gap_analysis', '暂无数据'))

                        except Exception as e:
                            st.error(f"分析失败: {str(e)}")

# ==================== Tab2: 岗位推荐 ====================
with tab2:
    st.markdown("### 岗位推荐")

    col1, col2 = st.columns(2)
    with col1:
        keywords = st.text_input("关键词", value="Python开发")
    with col2:
        location = st.text_input("地点", value="北京")

    if st.button("搜索岗位", type="primary"):
        with st.spinner("搜索中..."):
            try:
                # 尝试调用后端 API
                response = requests.get(
                    f"{BACKEND_URL}/api/jobs/search",
                    params={"keywords": keywords, "location": location},
                    timeout=10
                )

                if response.status_code == 200:
                    jobs = response.json()
                    st.session_state.current_jobs = jobs

                    if jobs:
                        st.success(f"找到 {len(jobs)} 个岗位")

                        for job in jobs:
                            with st.container():
                                st.markdown(f"""
                                <div class="job-card">
                                    <h4>{job.get('title', '未知职位')}</h4>
                                    <p><strong>公司：</strong>{job.get('company', '未知公司')}</p>
                                    <p><strong>薪资：</strong>{job.get('salary', '面议')}</p>
                                    <p><strong>地点：</strong>{job.get('location', '未知')}</p>
                                </div>
                                """, unsafe_allow_html=True)

                                if st.button("查看详情", key=f"job_{job.get('id', '')}"):
                                    st.info(job.get('description', '暂无描述'))
                    else:
                        st.warning("未找到相关岗位")
                else:
                    st.error("后端服务未启动，请先启动 FastAPI 后端")
                    st.info("运行: python web_app.py")

            except requests.exceptions.ConnectionError:
                st.error("无法连接到后端服务")
                st.info("""
                **本地运行后端：**
                ```bash
                python web_app.py
                ```

                **或使用模拟数据（开发模式）**
                """)

                # 模拟数据
                if st.checkbox("使用模拟数据"):
                    mock_jobs = [
                        {
                            "id": "1",
                            "title": "Python后端开发工程师",
                            "company": "字节跳动",
                            "salary": "25K-45K",
                            "location": "北京-朝阳区",
                            "description": "负责后端服务开发..."
                        },
                        {
                            "id": "2",
                            "title": "全栈工程师",
                            "company": "腾讯",
                            "salary": "30K-50K",
                            "location": "北京-海淀区",
                            "description": "负责前后端开发..."
                        }
                    ]

                    for job in mock_jobs:
                        with st.container():
                            st.markdown(f"""
                            <div class="job-card">
                                <h4>{job['title']}</h4>
                                <p><strong>公司：</strong>{job['company']}</p>
                                <p><strong>薪资：</strong>{job['salary']}</p>
                                <p><strong>地点：</strong>{job['location']}</p>
                            </div>
                            """, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"搜索失败: {str(e)}")

# ==================== Tab3: 自动投递 ====================
with tab3:
    st.markdown("### 自动投递")

    st.info("基于 GitHub 高星项目 [GodsScion/Auto_job_applier_linkedIn](https://github.com/GodsScion/Auto_job_applier_linkedIn)")

    platform = st.selectbox("平台", ["Boss直聘", "智联招聘", "LinkedIn (Easy Apply)"])

    col1, col2 = st.columns(2)
    with col1:
        keywords = st.text_input("关键词", value="Python Developer", key="apply_keywords")
        max_count = st.number_input("数量", 1, 500, 10)
    with col2:
        location = st.text_input("地点", value="Remote", key="apply_location")
        interval = st.slider("间隔（秒）", 3, 30, 5)

    if st.button("开始投递", type="primary"):
        st.warning("需要本地运行")

        with st.expander("本地运行指南"):
            st.markdown("""
            ```bash
            # 克隆项目
            git clone https://github.com/GodsScion/Auto_job_applier_linkedIn.git
            cd Auto_job_applier_linkedIn

            # 安装依赖
            pip install -r requirements.txt

            # 配置 config.yaml
            # 填写你的 LinkedIn 账号和投递参数

            # 运行
            python main.py
            ```
            """)

# ==================== Tab4: 投递记录 ====================
with tab4:
    st.markdown("### 投递记录")

    try:
        # 尝试从后端获取记录
        response = requests.get(
            f"{BACKEND_URL}/api/applications",
            params={"user_id": st.session_state.user_id},
            timeout=5
        )

        if response.status_code == 200:
            applications = response.json()
            st.session_state.applications = applications

            if applications:
                df = pd.DataFrame(applications)
                st.dataframe(df, use_container_width=True)

                # 状态统计
                if 'status' in df.columns:
                    st.bar_chart(df['status'].value_counts())
            else:
                st.info("暂无投递记录")
        else:
            st.error("后端服务未启动")

    except requests.exceptions.ConnectionError:
        st.info("后端服务未连接，显示本地记录")

        # 使用本地 session state
        if st.session_state.applications:
            df = pd.DataFrame(st.session_state.applications)
            st.dataframe(df, use_container_width=True)
        else:
            st.info("暂无投递记录")

            # 模拟数据
            if st.checkbox("显示模拟数据", key="mock_applications"):
                mock_data = [
                    {"公司": "字节跳动", "职位": "Python开发", "状态": "已投递", "时间": "2026-02-18"},
                    {"公司": "腾讯", "职位": "全栈工程师", "状态": "面试中", "时间": "2026-02-17"},
                    {"公司": "阿里巴巴", "职位": "后端开发", "状态": "已拒绝", "时间": "2026-02-16"},
                ]
                df = pd.DataFrame(mock_data)
                st.dataframe(df, use_container_width=True)

    except Exception as e:
        st.error(f"获取记录失败: {str(e)}")

# ==================== Tab5: 数据统计 ====================
with tab5:
    st.markdown("### 数据统计")

    try:
        # 尝试从后端获取统计
        response = requests.get(
            f"{BACKEND_URL}/api/stats",
            params={"user_id": st.session_state.user_id},
            timeout=5
        )

        if response.status_code == 200:
            stats = response.json()

            # 关键指标
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{stats.get('total_applications', 0)}</div>
                    <div class="metric-label">总投递</div>
                </div>
                """, unsafe_allow_html=True)

            with col2:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{stats.get('response_rate', 0)}%</div>
                    <div class="metric-label">回复率</div>
                </div>
                """, unsafe_allow_html=True)

            with col3:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{stats.get('interviews', 0)}</div>
                    <div class="metric-label">面试邀请</div>
                </div>
                """, unsafe_allow_html=True)

            with col4:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{stats.get('offers', 0)}</div>
                    <div class="metric-label">Offer</div>
                </div>
                """, unsafe_allow_html=True)

            # 趋势图
            if 'daily_applications' in stats:
                st.line_chart(stats['daily_applications'])
        else:
            st.error("后端服务未启动")

    except requests.exceptions.ConnectionError:
        st.info("后端服务未连接，显示模拟数据")

        # 模拟数据
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">23</div>
                <div class="metric-label">总投递</div>
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">35%</div>
                <div class="metric-label">回复率</div>
            </div>
            """, unsafe_allow_html=True)

        with col3:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">5</div>
                <div class="metric-label">面试邀请</div>
            </div>
            """, unsafe_allow_html=True)

        with col4:
            st.markdown("""
            <div class="metric-card">
                <div class="metric-value">2</div>
                <div class="metric-label">Offer</div>
            </div>
            """, unsafe_allow_html=True)

        # 模拟趋势图
        import numpy as np
        chart_data = pd.DataFrame(
            np.random.randn(20, 1),
            columns=['投递数量']
        )
        st.line_chart(chart_data)

    except Exception as e:
        st.error(f"获取统计失败: {str(e)}")

# 页脚
st.markdown('''
<div style="text-align:center;color:var(--text-light);padding:2rem 0;font-size:0.75rem;border-top:1px solid var(--border);margin-top:3rem">
    <a href="https://github.com/emptyteabot/ai-job-helper" style="color:var(--text-light);margin:0 1rem;text-decoration:none">GitHub</a>
    <a href="https://github.com/GodsScion/Auto_job_applier_linkedIn" style="color:var(--text-light);margin:0 1rem;text-decoration:none">高星项目</a>
</div>
''', unsafe_allow_html=True)
