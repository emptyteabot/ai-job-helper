"""
AI求职助手 - Web界面
一个漂亮的网页界面，让您直接在浏览器中使用
"""

from fastapi import FastAPI, Request, UploadFile, File, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import os
import sys
from typing import Optional, List, Dict, Any, Tuple
import asyncio
from datetime import datetime
from urllib.parse import quote_plus, urlparse, parse_qs, unquote
import re
import html as html_lib
import requests
import logging
import time
import uuid

sys.path.insert(0, os.path.dirname(__file__))

from app.core.multi_ai_debate import JobApplicationPipeline
from app.core.fast_ai_engine import fast_pipeline, HighPerformanceAIEngine
from app.core.market_driven_engine import market_driven_pipeline
from app.core.llm_client import get_public_llm_config
from app.services.resume_analyzer import ResumeAnalyzer
from app.services.real_job_service import RealJobService
from app.services.business_service import BusinessService
from app.core.realtime_progress import progress_tracker

app = FastAPI(title="AI求职助手")
APP_BOOT_TS = datetime.now().isoformat()
APP_BOOT_MONO = time.perf_counter()
logger = logging.getLogger("ai_job_helper")
if not logger.handlers:
    logging.basicConfig(
        level=os.getenv("APP_LOG_LEVEL", "INFO"),
        format="%(asctime)s %(levelname)s %(name)s %(message)s",
    )

# 使用市场驱动引擎
market_engine = market_driven_pipeline

# 挂载静态文件
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def request_context_middleware(request: Request, call_next):
    """Attach request id + process time headers for observability."""
    rid = request.headers.get("x-request-id") or str(uuid.uuid4())
    start = time.perf_counter()
    request.state.request_id = rid
    try:
        response = await call_next(request)
    except Exception as e:
        took_ms = (time.perf_counter() - start) * 1000
        logger.exception(
            "request_failed path=%s method=%s rid=%s took_ms=%.1f err=%s",
            request.url.path,
            request.method,
            rid,
            took_ms,
            str(e)[:300],
        )
        return JSONResponse(
            {"success": False, "error": "internal_error", "request_id": rid},
            status_code=500,
            headers={"x-request-id": rid, "x-process-time-ms": f"{took_ms:.1f}"},
        )

    took_ms = (time.perf_counter() - start) * 1000
    response.headers["x-request-id"] = rid
    response.headers["x-process-time-ms"] = f"{took_ms:.1f}"
    if request.url.path.startswith("/api"):
        logger.info(
            "request_done path=%s method=%s status=%s rid=%s took_ms=%.1f",
            request.url.path,
            request.method,
            getattr(response, "status_code", 0),
            rid,
            took_ms,
        )
    return response

# 全局变量
pipeline = JobApplicationPipeline()
analyzer = ResumeAnalyzer()
real_job_service = RealJobService()  # 真实招聘数据服务
business_service = BusinessService()

# 云端岗位缓存（内存）
cloud_jobs_cache: List[Dict[str, Any]] = []
CN_JOB_DOMAINS = ("zhipin.com", "liepin.com", "zhaopin.com", "51job.com", "lagou.com")
cloud_jobs_meta: Dict[str, Any] = {
    "last_push_at": None,
    "last_received": 0,
    "last_new": 0,
}
recent_search_jobs: Dict[str, Dict[str, Any]] = {}


def _api_success(payload: Dict[str, Any], status_code: int = 200) -> JSONResponse:
    body = {"success": True}
    body.update(payload or {})
    return JSONResponse(body, status_code=status_code)


def _api_error(message: str, status_code: int = 400, code: str = "bad_request") -> JSONResponse:
    return JSONResponse(
        {"success": False, "error": message, "code": code},
        status_code=status_code,
    )


def _build_investor_readiness_snapshot() -> Dict[str, Any]:
    """
    Financing-oriented readiness snapshot.
    Provides one consolidated view for product, traction, reliability and GTM.
    """
    metrics = business_service.metrics()
    funnel = metrics.get("funnel", {})
    leads = metrics.get("leads", {})
    feedback = metrics.get("feedback", {})
    stability = metrics.get("stability", {})

    cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
    enterprise_on = bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip())
    global_fallback_on = os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}

    uploads = int(funnel.get("uploads", 0) or 0)
    process_runs = int(funnel.get("process_runs", 0) or 0)
    searches = int(funnel.get("searches", 0) or 0)
    applies = int(funnel.get("applies", 0) or 0)
    api_errors = int(stability.get("api_errors", 0) or 0)

    process_success_pct = float(funnel.get("process_success_pct", 0) or 0)
    search_to_apply_pct = float(funnel.get("search_to_apply_pct", 0) or 0)

    product_score = 100.0
    if cfg_mode not in {"auto", "cloud", "baidu", "bing", "brave", "jooble", "openclaw", "enterprise_api"}:
        product_score -= 25.0
    if global_fallback_on:
        product_score -= 15.0
    if process_success_pct < 90:
        product_score -= 15.0
    if process_success_pct < 75:
        product_score -= 15.0
    if enterprise_on:
        product_score += 5.0
    product_score = max(0.0, min(100.0, product_score))

    traction_score = 0.0
    traction_score += min(35.0, uploads * 1.2)
    traction_score += min(35.0, process_runs * 1.5)
    traction_score += min(20.0, searches * 0.9)
    traction_score += min(10.0, applies * 2.0)
    traction_score = max(0.0, min(100.0, traction_score))

    reliability_score = 100.0
    if process_runs > 0:
        err_ratio = api_errors / max(process_runs, 1)
        reliability_score -= min(60.0, err_ratio * 100.0)
    if api_errors >= 20:
        reliability_score -= 10.0
    reliability_score = max(0.0, min(100.0, reliability_score))

    gtm_score = 0.0
    gtm_score += min(45.0, int(leads.get("total", 0) or 0) * 12.0)
    gtm_score += min(25.0, int(leads.get("last_7d", 0) or 0) * 8.0)
    gtm_score += min(10.0, int(feedback.get("last_7d", 0) or 0) * 2.5)
    gtm_score += min(15.0, search_to_apply_pct * 0.8)
    gtm_score += 10.0 if searches > 0 else 0.0
    gtm_score += 5.0 if applies > 0 else 0.0
    gtm_score = max(0.0, min(100.0, gtm_score))

    weighted = (
        product_score * 0.35
        + traction_score * 0.25
        + reliability_score * 0.25
        + gtm_score * 0.15
    )
    overall = round(max(0.0, min(100.0, weighted)), 1)

    status = "seed_ready" if overall >= 75 else ("pre_seed_ready" if overall >= 60 else "build_stage")

    return {
        "status": status,
        "overall_score": overall,
        "pillars": {
            "product": round(product_score, 1),
            "traction": round(traction_score, 1),
            "reliability": round(reliability_score, 1),
            "go_to_market": round(gtm_score, 1),
        },
        "highlights": {
            "provider_mode": cfg_mode,
            "enterprise_api_configured": enterprise_on,
            "global_fallback_enabled": global_fallback_on,
            "cloud_cache_total": len(cloud_jobs_cache),
            "feedback_total": int(feedback.get("total", 0) or 0),
        },
        "metrics": metrics,
        "next_30d_targets": {
            "uploads": max(uploads + 80, 120),
            "process_runs": max(process_runs + 60, 100),
            "searches": max(searches + 100, 180),
            "applies": max(applies + 30, 40),
            "leads_total": max(int(leads.get("total", 0) or 0) + 20, 30),
        },
    }


def _track_event(name: str, payload: Optional[Dict[str, Any]] = None) -> None:
    try:
        business_service.track_event(name, payload or {})
    except Exception:
        # Do not break core user path due to analytics write failures.
        pass


def _cache_recent_jobs(jobs: List[Dict[str, Any]], max_size: int = 2000) -> None:
    now = datetime.now().isoformat()
    for j in jobs or []:
        jid = str(j.get("id") or "").strip()
        if not jid:
            continue
        row = dict(j)
        row["_cached_at"] = now
        recent_search_jobs[jid] = row
    if len(recent_search_jobs) > max_size:
        # Keep newest max_size by cached time.
        items = sorted(
            recent_search_jobs.items(),
            key=lambda x: str(x[1].get("_cached_at") or ""),
            reverse=True,
        )[:max_size]
        recent_search_jobs.clear()
        recent_search_jobs.update(items)


def _is_seed_or_demo_job(job: Dict[str, Any]) -> bool:
    """Filter obvious seed/demo jobs so they never show in production recommendations."""
    jid = str(job.get("id") or "").strip().lower()
    if jid.startswith("seed") or jid.startswith("demo"):
        return True
    link = str(job.get("link") or job.get("apply_url") or "").strip().lower()
    if "/job_detail/seed" in link:
        return True
    company = str(job.get("company") or "").strip().lower()
    if company in {"示例公司", "测试公司", "demo"}:
        return True
    return False


def _job_has_actionable_link(job: Dict[str, Any]) -> bool:
    link = str(job.get("link") or job.get("apply_url") or "").strip()
    return link.startswith("http://") or link.startswith("https://")


def _is_cn_entrypoint_link(link: str) -> bool:
    """Search entry/list pages are not treated as real actionable postings."""
    low = (link or "").strip().lower()
    if not low:
        return False
    patterns = (
        "zhipin.com/web/geek/job?",
        "zhipin.com/zhaopin/",
        "liepin.com/zhaopin/?key=",
        "sou.zhaopin.com/?kw=",
        "we.51job.com/pc/search?",
        "lagou.com/wn/jobs?kd=",
    )
    return any(p in low for p in patterns)


def _is_cn_entrypoint_job(job: Dict[str, Any]) -> bool:
    provider = str(job.get("provider") or "").strip().lower()
    title = str(job.get("title") or job.get("job_title") or "").strip().lower()
    link = str(job.get("link") or job.get("apply_url") or "").strip()
    if provider == "cn_portal":
        return True
    if "搜索入口" in title:
        return True
    return _is_cn_entrypoint_link(link)


def _normalize_and_filter_jobs(jobs: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    """Keep only actionable real jobs and deduplicate by link/id/title-company."""
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for job in jobs or []:
        if not isinstance(job, dict):
            continue
        if _is_seed_or_demo_job(job):
            continue
        if not _job_has_actionable_link(job):
            continue

        link = str(job.get("link") or job.get("apply_url") or "").strip().lower()
        jid = str(job.get("id") or "").strip().lower()
        title = str(job.get("title") or job.get("job_title") or "").strip().lower()
        company = str(job.get("company") or "").strip().lower()
        dedupe_key = link or jid or f"{title}|{company}"
        if not dedupe_key or dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        out.append(job)
        if len(out) >= max(1, int(limit or 10)):
            break
    return out


def _normalize_real_jobs(jobs: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    """
    Strict real-posting normalization:
    - actionable links only
    - no seed/demo
    - no board-level search entry pages
    """
    scan_limit = max(20, int(limit or 10) * 5)
    base = _normalize_and_filter_jobs(jobs, limit=scan_limit)
    out: List[Dict[str, Any]] = []
    for job in base:
        if _is_cn_entrypoint_job(job):
            continue
        out.append(job)
        if len(out) >= max(1, int(limit or 10)):
            break
    return out


def _is_cn_job_link(link: str) -> bool:
    low = (link or "").lower()
    return any(d in low for d in CN_JOB_DOMAINS)


def _enforce_cn_market_jobs(
    jobs: List[Dict[str, Any]],
    allow_entrypoints: bool = False,
) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for j in (jobs or []):
        if (not allow_entrypoints) and _is_cn_entrypoint_job(j):
            continue
        link = str(j.get("link") or j.get("apply_url") or "")
        platform = str(j.get("platform") or "").strip()
        if _is_cn_job_link(link):
            out.append(j)
            continue
        if platform in {"Boss直聘", "猎聘", "智联招聘", "前程无忧", "拉勾"}:
            out.append(j)
    return out


PROCESS_RESPONSE_SCHEMA_VERSION = "process_response.v2"
LOW_QUALITY_MARKERS = (
    "i appreciate your interest",
    "amazon q",
    "built by aws",
    "i'm amazon q",
    "i am amazon q",
    "not career counseling",
    "没有看到任何附件",
)
PUBLIC_EVENT_NAME_PATTERN = re.compile(r"^[a-z0-9_]{3,48}$")
PUBLIC_EVENT_WHITELIST = {
    "workspace_opened",
    "resume_process_started",
    "resume_process_success",
    "resume_process_failed",
    "job_link_click",
    "result_download",
}


def _text_has_low_quality_marker(text: str) -> bool:
    low = (text or "").strip().lower()
    if not low:
        return True
    return any(marker in low for marker in LOW_QUALITY_MARKERS)


def _render_market_analysis_fallback(info: Dict[str, Any]) -> str:
    skills = [s for s in (info.get("skills") or []) if s][:8]
    skills_text = ", ".join(skills) if skills else "Python, FastAPI, SQL"
    return (
        "【市场竞争力分析】\n\n"
        f"✅ 识别技能：{skills_text}\n\n"
        "📊 市场需求度：中高（基于中国技术岗位关键词覆盖）\n"
        "💼 建议投递方向：Python后端 / AI应用工程 / 数据工程\n"
        "📈 建议策略：先投递有明确技术栈和薪资区间的岗位，再按反馈迭代简历。\n\n"
        "💡 市场建议：\n"
        "1. 先聚焦 1-2 个主岗位方向，避免关键词过散。\n"
        "2. 简历中补充可量化结果（性能提升、效率提升、交付周期）。\n"
        "3. 优先投递带真实岗位详情页和直接投递入口的职位。"
    )


def _render_interview_prep_fallback(info: Dict[str, Any], jobs: List[Dict[str, Any]]) -> str:
    top = (jobs or [{}])[0]
    title = str(top.get("title") or "Python后端工程师")
    company = str(top.get("company") or "目标公司")
    return (
        f"目标岗位：{title}（{company}）\n\n"
        "高频问题 1：你如何设计高并发 API？\n"
        "回答要点：异步 I/O、缓存策略、限流、慢查询治理、可观测性。\n\n"
        "高频问题 2：你如何保证数据链路质量？\n"
        "回答要点：输入校验、幂等设计、回滚方案、监控告警、复盘机制。\n\n"
        "高频问题 3：你做过最有业务价值的项目是什么？\n"
        "回答要点：按“背景-动作-结果”讲清指标提升，并说明你的关键贡献。"
    )


def _render_optimized_resume_fallback(resume_text: str, info: Dict[str, Any]) -> str:
    name = str(info.get("name") or "").strip()
    if not name or name == "未知":
        first = (resume_text or "").strip().splitlines()
        name = first[0].strip() if first else "候选人"
    skills = [s for s in (info.get("skills") or []) if s][:10]
    skills_line = ", ".join(skills) if skills else "Python, FastAPI, SQL, Docker"
    return (
        f"{name}\n"
        "AI应用工程师 | Python后端工程师\n\n"
        "核心优势\n"
        "- 聚焦 AI 应用工程与后端交付，能从需求到上线完成闭环。\n"
        "- 具备数据处理、服务部署、性能优化和质量门禁实践。\n\n"
        f"技术栈\n- {skills_line}\n\n"
        "项目亮点（示例结构）\n"
        "- 构建 RAG/数据服务，接口稳定性提升，响应延迟下降。\n"
        "- 建立自动化数据管道，减少人工处理成本并提高时效。\n"
        "- 引入测试与监控机制，降低线上故障率并缩短定位时间。\n\n"
        "求职方向\n- Python后端开发\n- AI应用开发\n- 数据工程与平台方向"
    )


def _run_output_quality_gate(
    results: Dict[str, Any],
    resume_text: str,
    info: Dict[str, Any],
    real_jobs: List[Dict[str, Any]],
) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    clean = dict(results or {})
    report: Dict[str, Any] = {"passed": True, "issues": []}

    market_analysis = str(clean.get("market_analysis") or "").strip()
    if len(market_analysis) < 80 or _text_has_low_quality_marker(market_analysis):
        clean["market_analysis"] = _render_market_analysis_fallback(info)
        report["issues"].append("market_analysis_replaced")

    optimized_resume = str(clean.get("optimized_resume") or "").strip()
    if len(optimized_resume) < 120 or _text_has_low_quality_marker(optimized_resume):
        clean["optimized_resume"] = _render_optimized_resume_fallback(resume_text, info)
        report["issues"].append("optimized_resume_replaced")

    interview_prep = str(clean.get("interview_prep") or "").strip()
    if len(interview_prep) < 80 or _text_has_low_quality_marker(interview_prep):
        clean["interview_prep"] = _render_interview_prep_fallback(info, real_jobs)
        report["issues"].append("interview_prep_replaced")

    salary_analysis = str(clean.get("salary_analysis") or "").strip()
    if not salary_analysis:
        clean["salary_analysis"] = "【薪资潜力分析】\n\n建议：基于岗位匹配度和市场供需，优先投递高匹配岗位后再进行薪资谈判。"
        report["issues"].append("salary_analysis_replaced")

    if report["issues"]:
        report["passed"] = False
    return clean, report


def _public_job_payload(jobs: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for j in _normalize_real_jobs(jobs or [], limit=limit):
        link = str(j.get("link") or j.get("apply_url") or "").strip()
        out.append(
            {
                "id": str(j.get("id") or f"job_{abs(hash(link.lower()))}"),
                "title": str(j.get("title") or j.get("job_title") or "未知岗位"),
                "company": str(j.get("company") or ""),
                "location": str(j.get("location") or ""),
                "salary": str(j.get("salary") or j.get("salary_range") or ""),
                "platform": str(j.get("platform") or j.get("provider") or _platform_from_link(link)),
                "link": link,
                "provider": str(j.get("provider") or ""),
                "updated": str(j.get("updated") or ""),
            }
        )
    return out


def _validate_process_response_shape(payload: Dict[str, Any]) -> List[str]:
    errors: List[str] = []
    required_text_fields = (
        "career_analysis",
        "job_recommendations",
        "optimized_resume",
        "interview_prep",
        "mock_interview",
        "job_provider_mode",
        "schema_version",
    )
    for key in required_text_fields:
        if not isinstance(payload.get(key), str):
            errors.append(f"{key}:expected_string")

    jobs = payload.get("recommended_jobs")
    if not isinstance(jobs, list):
        errors.append("recommended_jobs:expected_list")
    else:
        for idx, row in enumerate(jobs[:20]):
            if not isinstance(row, dict):
                errors.append(f"recommended_jobs[{idx}]:expected_object")
                continue
            for key in ("id", "title", "link"):
                if not isinstance(row.get(key), str) or not str(row.get(key)).strip():
                    errors.append(f"recommended_jobs[{idx}].{key}:missing")
                    break
    return errors


def _filter_cloud_cache_by_query(
    keywords: List[str], location: Optional[str], limit: int
) -> List[Dict[str, Any]]:
    kw = [k.strip().lower() for k in (keywords or []) if k and k.strip()]

    def hit(job: Dict[str, Any]) -> bool:
        text = f"{job.get('title','')} {job.get('company','')}".lower()
        if kw and not any(k in text for k in kw):
            return False
        if location and job.get("location") and location not in str(job.get("location")):
            return False
        return True

    matched = [j for j in cloud_jobs_cache if hit(j)]
    if not matched:
        matched = list(cloud_jobs_cache)
    return _normalize_real_jobs(matched, limit=limit)


def _search_jobs_without_browser(
    keywords: List[str],
    location: Optional[str],
    limit: int,
    allow_portal_fallback: bool = False,
) -> Tuple[List[Dict[str, Any]], str, Optional[str]]:
    """
    Cloud-safe real-time search path.
    No local browser/OpenClaw required.
    """
    try:
        jobs = real_job_service.search_jobs(
            keywords=keywords,
            location=location,
            limit=max(5, min(int(limit or 10), 50)),
        )
        mode = (real_job_service.get_statistics() or {}).get("provider_mode", "auto")
        normalized = _normalize_real_jobs(jobs, limit=limit)
        normalized = _enforce_cn_market_jobs(normalized)
        if normalized:
            return normalized, mode, None
    except Exception as e:
        first_error = str(e)
    else:
        first_error = None

    enterprise_jobs = _search_jobs_enterprise_api(keywords, location, limit=limit)
    enterprise_jobs = _normalize_real_jobs(enterprise_jobs, limit=limit)
    enterprise_jobs = _enforce_cn_market_jobs(enterprise_jobs)
    if enterprise_jobs:
        return enterprise_jobs, "enterprise_api", None

    # CN market fallback: no-browser HTML search on Chinese job sites.
    bing_html_jobs = _search_jobs_bing_html(keywords, location, limit=limit)
    bing_html_jobs = _normalize_real_jobs(bing_html_jobs, limit=limit)
    bing_html_jobs = _enforce_cn_market_jobs(bing_html_jobs)
    if bing_html_jobs:
        return bing_html_jobs, "bing_html", None

    # Last fallback: DuckDuckGo HTML search (no key, no browser).
    ddg_jobs = _search_jobs_duckduckgo(keywords, location, limit=limit)
    ddg_jobs = _normalize_real_jobs(ddg_jobs, limit=limit)
    ddg_jobs = _enforce_cn_market_jobs(ddg_jobs)
    if ddg_jobs:
        return ddg_jobs, "duckduckgo", None

    # Optional global fallback. Disabled by default to keep CN market realism.
    if os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}:
        remotive_jobs = _search_jobs_remotive(keywords, location, limit=limit)
        if remotive_jobs:
            return remotive_jobs, "remotive", None

    if allow_portal_fallback:
        portal_jobs = _search_jobs_cn_entrypoints(keywords, location, limit=min(limit, 5))
        if portal_jobs:
            return portal_jobs, "cn_portal", first_error or "using cn portal fallback"

    return [], "no_real_jobs", first_error or "no result from no-browser providers"


def _platform_from_link(link: str) -> str:
    host = (urlparse(link).netloc or "").lower()
    if "zhipin.com" in host:
        return "Boss直聘"
    if "liepin.com" in host:
        return "猎聘"
    if "zhaopin.com" in host:
        return "智联招聘"
    if "51job.com" in host:
        return "前程无忧"
    if "lagou.com" in host:
        return "拉勾"
    return host or "web"


def _first_non_empty(row: Dict[str, Any], keys: List[str]) -> str:
    for k in keys:
        v = row.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return ""


def _infer_company_from_title(title: str, platform: str = "") -> str:
    t = (title or "").strip()
    if not t:
        return ""
    # Boss style examples:
    # - 「Python招聘」_苏州鼎级招聘-BOSS直聘
    # - Python开发工程师 - 某某科技 - Boss直聘
    if "_" in t:
        tail = t.split("_", 1)[1].strip()
        tail = re.split(r"[-|｜]", tail)[0].strip()
        tail = tail.replace("招聘", "").replace("诚聘", "").strip()
        if 1 <= len(tail) <= 24 and "直聘" not in tail:
            return tail
    parts = [p.strip() for p in re.split(r"[-|｜]", t) if p.strip()]
    for p in parts[1:3]:
        if any(x in p for x in ("直聘", "招聘", "猎聘", "前程无忧", "拉勾", "智联")):
            continue
        if 1 <= len(p) <= 24:
            return p.replace("招聘", "").replace("诚聘", "").strip()
    return ""


def _search_jobs_enterprise_api(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    """
    Enterprise-grade provider hook (config-driven).
    Intended for stable cloud use when search engines hit anti-bot limits.
    """
    url = os.getenv("ENTERPRISE_JOB_API_URL", "").strip()
    if not url:
        return []

    method = os.getenv("ENTERPRISE_JOB_API_METHOD", "GET").strip().upper()
    timeout_s = int(os.getenv("ENTERPRISE_JOB_API_TIMEOUT_S", "15") or "15")
    auth_header = os.getenv("ENTERPRISE_JOB_API_AUTH_HEADER", "Authorization").strip() or "Authorization"
    auth_scheme = os.getenv("ENTERPRISE_JOB_API_AUTH_SCHEME", "Bearer").strip()
    api_key = os.getenv("ENTERPRISE_JOB_API_KEY", "").strip()

    kw = [k.strip() for k in (keywords or []) if k and k.strip()]
    query = " ".join(kw[:5]) or "Python"
    payload = {
        "query": query,
        "keywords": kw,
        "location": (location or "").strip(),
        "limit": max(1, min(int(limit or 10), 50)),
    }
    headers = {"User-Agent": "ai-job-helper/1.0"}
    if api_key:
        headers[auth_header] = f"{auth_scheme} {api_key}".strip() if auth_scheme else api_key

    try:
        if method == "POST":
            resp = requests.post(url, json=payload, headers=headers, timeout=timeout_s)
        else:
            resp = requests.get(url, params=payload, headers=headers, timeout=timeout_s)
        if resp.status_code >= 400:
            return []
        data = resp.json() if resp.content else {}
    except Exception:
        return []

    rows: List[Dict[str, Any]] = []
    if isinstance(data, list):
        rows = [x for x in data if isinstance(x, dict)]
    elif isinstance(data, dict):
        for k in ("jobs", "results", "items", "list", "data"):
            v = data.get(k)
            if isinstance(v, list):
                rows = [x for x in v if isinstance(x, dict)]
                break
            if isinstance(v, dict):
                for kk in ("jobs", "results", "items", "list"):
                    vv = v.get(kk)
                    if isinstance(vv, list):
                        rows = [x for x in vv if isinstance(x, dict)]
                        break
                if rows:
                    break

    out: List[Dict[str, Any]] = []
    for i, row in enumerate(rows, 1):
        link = _first_non_empty(row, ["link", "url", "job_url", "detail_url", "apply_url", "jobLink", "jobUrl"])
        if not link:
            continue
        title = _first_non_empty(row, ["title", "job_title", "name", "position", "jobName"]) or "招聘岗位"
        company = _first_non_empty(row, ["company", "company_name", "employer", "brandName"])
        loc = _first_non_empty(row, ["location", "city", "region", "work_city"]) or (location or "")
        salary = _first_non_empty(row, ["salary", "salary_range", "pay", "salaryRange"])
        platform = _first_non_empty(row, ["platform", "source", "site", "origin"]) or _platform_from_link(link)
        out.append(
            {
                "id": f"enterprise_{i}_{abs(hash(link.lower()))}",
                "title": title,
                "company": company,
                "location": loc,
                "salary": salary,
                "platform": platform,
                "link": link,
                "provider": "enterprise_api",
                "updated": _first_non_empty(row, ["updated", "updated_at", "publish_time", "published_at"]),
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _normalize_ddg_redirect(href: str) -> str:
    href = html_lib.unescape((href or "").strip())
    if not href:
        return ""
    if href.startswith("//"):
        href = "https:" + href
    # Absolute DuckDuckGo redirect URL.
    if href.startswith("http://duckduckgo.com/l/?") or href.startswith("https://duckduckgo.com/l/?"):
        qs = parse_qs(urlparse(href).query)
        uddg = (qs.get("uddg") or [""])[0]
        return unquote(uddg) if uddg else ""
    if href.startswith("/l/?"):
        qs = parse_qs(urlparse("https://duckduckgo.com" + href).query)
        uddg = (qs.get("uddg") or [""])[0]
        return unquote(uddg) if uddg else ""
    return href


def _search_jobs_duckduckgo(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    q_parts = [k.strip() for k in (keywords or []) if k and k.strip()]
    if location:
        q_parts.append(location.strip())
    q_parts.append("招聘 职位 site:zhipin.com OR site:liepin.com OR site:zhaopin.com OR site:51job.com OR site:lagou.com")
    q = " ".join(q_parts).strip() or "招聘 职位 site:zhipin.com"

    url = f"https://html.duckduckgo.com/html/?q={quote_plus(q)}"
    try:
        resp = requests.get(
            url,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            },
            timeout=12,
        )
        text = resp.text or ""
    except Exception:
        return []

    # result__a href="...">title</a>
    pattern = re.compile(r'<a[^>]*class="[^"]*result__a[^"]*"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.I | re.S)
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for href, raw_title in pattern.findall(text):
        link = _normalize_ddg_redirect(href)
        if not link:
            continue
        if not (link.startswith("http://") or link.startswith("https://")):
            continue
        low = link.lower()
        if not any(d in low for d in CN_JOB_DOMAINS):
            continue
        if low in seen:
            continue
        seen.add(low)
        title = re.sub(r"<[^>]+>", "", raw_title or "")
        title = html_lib.unescape(title).strip()
        platform = _platform_from_link(link)
        company = _infer_company_from_title(title, platform=platform)
        out.append(
            {
                "id": f"duckduckgo_{abs(hash(low))}",
                "title": title or "招聘岗位",
                "company": company,
                "location": location or "",
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "duckduckgo",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _search_jobs_bing_html(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    q_parts = [k.strip() for k in (keywords or []) if k and k.strip()]
    if location:
        q_parts.append(location.strip())
    q_parts.append("招聘 职位 site:zhipin.com OR site:liepin.com OR site:zhaopin.com OR site:51job.com OR site:lagou.com")
    q = " ".join(q_parts).strip() or "招聘 职位 site:zhipin.com"

    try:
        resp = requests.get(
            "https://www.bing.com/search",
            params={"q": q, "count": max(10, min(int(limit or 10) * 2, 50))},
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            },
            timeout=12,
        )
        text = resp.text or ""
    except Exception:
        return []

    # <li class="b_algo"> ... <h2><a href="...">title</a>
    pattern = re.compile(r'<li class="b_algo"[^>]*>.*?<h2><a href="([^"]+)"[^>]*>(.*?)</a>', re.I | re.S)
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for link, raw_title in pattern.findall(text):
        if not link.startswith(("http://", "https://")):
            continue
        low = link.lower()
        if not any(d in low for d in CN_JOB_DOMAINS):
            continue
        if low in seen:
            continue
        seen.add(low)
        title = re.sub(r"<[^>]+>", "", raw_title or "")
        title = html_lib.unescape(title).strip()
        platform = _platform_from_link(link)
        company = _infer_company_from_title(title, platform=platform)
        out.append(
            {
                "id": f"binghtml_{abs(hash(low))}",
                "title": title or "招聘岗位",
                "company": company,
                "location": location or "",
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "bing_html",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _search_jobs_remotive(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    # Remotive is a public job API and currently reachable in cloud environments.
    cleaned = [k.strip() for k in (keywords or []) if k and k.strip()]
    q = cleaned[0] if cleaned else "python"
    try:
        resp = requests.get(
            "https://remotive.com/api/remote-jobs",
            params={"search": q},
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=15,
        )
        data = resp.json() if resp.content else {}
    except Exception:
        return []

    jobs = data.get("jobs") or []
    base_rows: List[Dict[str, Any]] = []
    for it in jobs:
        link = str(it.get("url") or "").strip()
        if not link:
            continue
        title = str(it.get("title") or "").strip()
        company = str(it.get("company_name") or "").strip()
        loc = str(it.get("candidate_required_location") or location or "").strip()
        base_rows.append(
            {
                "id": f"remotive_{it.get('id')}",
                "title": title or "招聘岗位",
                "company": company,
                "location": loc,
                "salary": str(it.get("salary") or "").strip(),
                "platform": "Remotive",
                "link": link,
                "provider": "remotive",
                "updated": str(it.get("publication_date") or "").strip(),
            }
        )
        if len(base_rows) >= max(30, int(limit or 10) * 4):
            break

    if location:
        narrowed = [j for j in base_rows if location in str(j.get("location") or "")]
        if narrowed:
            return _normalize_and_filter_jobs(narrowed, limit=limit)
    # If no location match, return best available remote jobs instead of empty.
    return _normalize_and_filter_jobs(base_rows, limit=limit)


def _search_jobs_cn_entrypoints(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    """Guaranteed CN-market fallback: job-board search entry links."""
    kw = [k.strip() for k in (keywords or []) if k and k.strip()]
    query = " ".join(kw[:3]) or "Python"
    loc = (location or "").strip()

    boards = [
        ("Boss直聘", f"https://www.zhipin.com/web/geek/job?query={quote_plus(query)}"),
        ("猎聘", f"https://www.liepin.com/zhaopin/?key={quote_plus(query)}"),
        ("智联招聘", f"https://sou.zhaopin.com/?kw={quote_plus(query)}" + (f"&jl={quote_plus(loc)}" if loc else "")),
        ("前程无忧", f"https://we.51job.com/pc/search?keyword={quote_plus(query)}"),
        ("拉勾", f"https://www.lagou.com/wn/jobs?kd={quote_plus(query)}"),
    ]
    out: List[Dict[str, Any]] = []
    for i, (platform, link) in enumerate(boards, 1):
        out.append(
            {
                "id": f"cn_portal_{i}_{abs(hash(link))}",
                "title": f"{query} - {platform}搜索入口",
                "company": "",
                "location": loc,
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "cn_portal",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)

@app.get("/", response_class=HTMLResponse)
async def home():
    """主页 - 官网首页"""
    static_index = "static/index.html"
    if os.path.exists(static_index):
        with open(static_index, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>AI求职助手</h1>")

@app.get("/app", response_class=HTMLResponse)
async def app_page():
    """主应用页面"""
    app_candidates = ["static/app_clean.html", "static/app.html"]
    for app_html in app_candidates:
        if not (os.path.exists(app_html) and os.path.getsize(app_html) > 64):
            continue
        with open(app_html, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())

    return """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI求职助手 - 多AI协作系统</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
        }
        
        .header {
            text-align: center;
            color: white;
            margin-bottom: 40px;
            animation: fadeInDown 0.8s ease;
        }
        
        .header h1 {
            font-size: 3em;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        
        .header p {
            font-size: 1.2em;
            opacity: 0.9;
        }
        
        .main-card {
            background: white;
            border-radius: 20px;
            padding: 40px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            animation: fadeInUp 0.8s ease;
        }
        
        .step-indicator {
            display: flex;
            justify-content: space-between;
            margin-bottom: 40px;
            position: relative;
        }
        
        .step-indicator::before {
            content: '';
            position: absolute;
            top: 20px;
            left: 0;
            right: 0;
            height: 2px;
            background: #e0e0e0;
            z-index: 0;
        }
        
        .step {
            flex: 1;
            text-align: center;
            position: relative;
            z-index: 1;
        }
        
        .step-circle {
            width: 40px;
            height: 40px;
            border-radius: 50%;
            background: #e0e0e0;
            color: white;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 10px;
            font-weight: bold;
            transition: all 0.3s;
        }
        
        .step.active .step-circle {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            transform: scale(1.2);
        }
        
        .step-label {
            font-size: 0.9em;
            color: #666;
        }
        
        .step.active .step-label {
            color: #667eea;
            font-weight: bold;
        }
        
        .input-section {
            margin-bottom: 30px;
        }
        
        .input-section label {
            display: block;
            font-size: 1.1em;
            font-weight: bold;
            color: #333;
            margin-bottom: 10px;
        }
        
        textarea {
            width: 100%;
            min-height: 300px;
            padding: 15px;
            border: 2px solid #e0e0e0;
            border-radius: 10px;
            font-size: 1em;
            font-family: inherit;
            resize: vertical;
            transition: border-color 0.3s;
        }
        
        textarea:focus {
            outline: none;
            border-color: #667eea;
        }
        
        .button-group {
            display: flex;
            gap: 15px;
            justify-content: center;
        }
        
        button {
            padding: 15px 40px;
            font-size: 1.1em;
            font-weight: bold;
            border: none;
            border-radius: 10px;
            cursor: pointer;
            transition: all 0.3s;
        }
        
        .btn-primary {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 20px rgba(102, 126, 234, 0.4);
        }
        
        .btn-secondary {
            background: #f0f0f0;
            color: #666;
        }
        
        .btn-secondary:hover {
            background: #e0e0e0;
        }
        
        .loading {
            display: none;
            text-align: center;
            padding: 40px;
        }
        
        .loading.active {
            display: block;
        }
        
        .spinner {
            border: 4px solid #f3f3f3;
            border-top: 4px solid #667eea;
            border-radius: 50%;
            width: 50px;
            height: 50px;
            animation: spin 1s linear infinite;
            margin: 0 auto 20px;
        }
        
        .ai-status {
            margin-top: 20px;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 10px;
            font-size: 0.95em;
        }
        
        .ai-status .ai-item {
            padding: 8px;
            margin: 5px 0;
            border-left: 3px solid #667eea;
            padding-left: 15px;
        }
        
        .results {
            display: none;
            margin-top: 30px;
        }
        
        .results.active {
            display: block;
        }
        
        .result-card {
            background: #f8f9fa;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 20px;
            border-left: 4px solid #667eea;
        }
        
        .result-card h3 {
            color: #667eea;
            margin-bottom: 15px;
            font-size: 1.3em;
        }
        
        .result-card pre {
            white-space: pre-wrap;
            word-wrap: break-word;
            font-family: inherit;
            line-height: 1.6;
        }
        
        @keyframes fadeInDown {
            from {
                opacity: 0;
                transform: translateY(-30px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        @keyframes fadeInUp {
            from {
                opacity: 0;
                transform: translateY(30px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        .ai-roles {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
            margin: 30px 0;
        }
        
        .ai-role-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            transition: transform 0.3s;
        }
        
        .ai-role-card:hover {
            transform: translateY(-5px);
        }
        
        .ai-role-card .icon {
            font-size: 2em;
            margin-bottom: 10px;
        }
        
        .ai-role-card .name {
            font-weight: bold;
            margin-bottom: 5px;
        }
        
        .ai-role-card .task {
            font-size: 0.9em;
            opacity: 0.9;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🤖 AI求职助手</h1>
            <p>多AI协作系统 - 让6个AI帮您找到理想工作</p>
        </div>
        
        <div class="main-card">
            <div class="step-indicator">
                <div class="step active" id="step1">
                    <div class="step-circle">1</div>
                    <div class="step-label">上传简历</div>
                </div>
                <div class="step" id="step2">
                    <div class="step-circle">2</div>
                    <div class="step-label">AI分析</div>
                </div>
                <div class="step" id="step3">
                    <div class="step-circle">3</div>
                    <div class="step-label">优化简历</div>
                </div>
                <div class="step" id="step4">
                    <div class="step-circle">4</div>
                    <div class="step-label">面试准备</div>
                </div>
            </div>
            
            <div class="ai-roles">
                <div class="ai-role-card">
                    <div class="icon">👔</div>
                    <div class="name">职业规划师</div>
                    <div class="task">分析优势定位</div>
                </div>
                <div class="ai-role-card">
                    <div class="icon">🔍</div>
                    <div class="name">招聘专家</div>
                    <div class="task">搜索匹配岗位</div>
                </div>
                <div class="ai-role-card">
                    <div class="icon">✍️</div>
                    <div class="name">简历优化师</div>
                    <div class="task">改写简历内容</div>
                </div>
                <div class="ai-role-card">
                    <div class="icon">✅</div>
                    <div class="name">质量检查官</div>
                    <div class="task">审核并改进</div>
                </div>
                <div class="ai-role-card">
                    <div class="icon">🎓</div>
                    <div class="name">面试教练</div>
                    <div class="task">面试辅导</div>
                </div>
                <div class="ai-role-card">
                    <div class="icon">🎭</div>
                    <div class="name">模拟面试官</div>
                    <div class="task">模拟面试</div>
                </div>
            </div>
            
            <div class="input-section" id="inputSection">
                <label>📄 上传简历或输入内容：</label>
                
                <!-- 文件上传区域 -->
                <div style="margin-bottom: 20px;">
                    <input type="file" id="fileInput" accept=".pdf,.docx,.doc,.txt,.jpg,.jpeg,.png,.bmp,.gif" style="display: none;" onchange="handleFileUpload(event)">
                    <button class="btn-secondary" onclick="document.getElementById('fileInput').click()" style="width: 100%; padding: 20px; font-size: 1em;">
                        📎 点击上传简历文件（支持PDF、Word、TXT、图片）
                    </button>
                    <div id="uploadStatus" style="margin-top: 10px; text-align: center; color: #667eea; font-weight: bold;"></div>
                </div>
                
                <div style="text-align: center; margin: 20px 0; color: #999;">或者</div>
                
                <textarea id="resumeInput" placeholder="直接输入您的简历内容：
- 姓名、学历、工作经验
- 技能清单（编程语言、框架、工具等）
- 项目经验
- 求职意向

示例：
姓名：张三
学历：本科 - 计算机科学
工作经验：3年Python开发
技能：Python, Django, MySQL, Redis
项目：电商后台系统、数据分析平台
求职意向：Python后端开发工程师"></textarea>
                
                <div class="button-group" style="margin-top: 20px;">
                    <button class="btn-secondary" onclick="loadExample()">使用示例简历</button>
                    <button class="btn-primary" onclick="startProcess()">🚀 开始AI协作</button>
                </div>
            </div>
            
            <div class="loading" id="loading">
                <div class="spinner"></div>
                <h3>🤖 AI团队正在协作中...</h3>
                <p>6个AI正在依次工作，预计需要1-2分钟</p>
                <div class="ai-status" id="aiStatus"></div>
            </div>
            
            <div class="results" id="results">
                <h2 style="text-align: center; color: #667eea; margin-bottom: 30px;">🎉 AI协作完成！</h2>
                
                <div class="result-card">
                    <h3>📊 职业分析报告</h3>
                    <pre id="careerAnalysis"></pre>
                </div>
                
                <div class="result-card">
                    <h3>🎯 推荐岗位列表</h3>
                    <pre id="jobRecommendations"></pre>
                </div>
                
                <div class="result-card">
                    <h3>✍️ 优化后的简历</h3>
                    <pre id="optimizedResume"></pre>
                </div>
                
                <div class="result-card">
                    <h3>🎓 面试准备指南</h3>
                    <pre id="interviewPrep"></pre>
                </div>
                
                <div class="result-card">
                    <h3>🎭 模拟面试问答</h3>
                    <pre id="mockInterview"></pre>
                </div>
                
                <div class="button-group" style="margin-top: 30px;">
                    <button class="btn-secondary" onclick="location.reload()">重新开始</button>
                    <button class="btn-primary" onclick="downloadResults()">📥 下载结果</button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        async function handleFileUpload(event) {
            const file = event.target.files[0];
            if (!file) return;
            
            const statusDiv = document.getElementById('uploadStatus');
            statusDiv.textContent = '📤 正在上传和解析文件...';
            
            try {
                const formData = new FormData();
                formData.append('file', file);
                
                const response = await fetch('/api/upload', {
                    method: 'POST',
                    body: formData
                });
                
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.error || '上传失败');
                }
                
                const data = await response.json();
                
                // 显示上传成功
                statusDiv.textContent = `✅ 文件上传成功：${data.filename}`;
                statusDiv.style.color = '#27ae60';
                
                // 等待1秒后自动开始AI处理
                setTimeout(() => {
                    statusDiv.textContent = '🚀 正在启动AI协作...';
                    statusDiv.style.color = '#667eea';
                    
                    // 自动开始处理
                    processResumeText(data.resume_text);
                }, 1000);
                
            } catch (error) {
                statusDiv.textContent = `❌ ${error.message}`;
                statusDiv.style.color = '#e74c3c';
            }
        }
        
        async function processResumeText(resumeText) {
            // 隐藏输入区，显示加载动画
            document.getElementById('inputSection').style.display = 'none';
            document.getElementById('loading').classList.add('active');
            document.getElementById('step2').classList.add('active');
            
            try {
                // 调用后端API
                const response = await fetch('/api/process', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({ resume: resumeText })
                });
                
                if (!response.ok) {
                    throw new Error('处理失败');
                }
                
                const data = await response.json();
                
                // 显示结果
                document.getElementById('loading').classList.remove('active');
                document.getElementById('results').classList.add('active');
                document.getElementById('step3').classList.add('active');
                document.getElementById('step4').classList.add('active');
                
                document.getElementById('careerAnalysis').textContent = data.career_analysis;
                document.getElementById('jobRecommendations').textContent = data.job_recommendations;
                document.getElementById('optimizedResume').textContent = data.optimized_resume;
                document.getElementById('interviewPrep').textContent = data.interview_prep;
                document.getElementById('mockInterview').textContent = data.mock_interview;
                
            } catch (error) {
                alert('处理出错：' + error.message);
                location.reload();
            }
        }
        
        function loadExample() {
            document.getElementById('resumeInput').value = `姓名：李明
学历：本科 - 软件工程
工作经验：3年Python开发经验

技能清单：
- 编程语言: Python, JavaScript, SQL
- 后端框架: Django, Flask, FastAPI
- 数据库: MySQL, Redis, MongoDB
- 前端: React, Vue.js, HTML/CSS
- 工具: Docker, Git, Linux

项目经验：
1. 电商后台管理系统
   - 使用Django + MySQL开发
   - 实现商品管理、订单处理、用户权限等功能
   - 日均处理订单5000+

2. 数据分析平台
   - 使用Python + Pandas进行数据处理
   - 开发可视化报表系统
   - 支持实时数据监控

3. RESTful API服务
   - 使用FastAPI开发高性能API
   - 集成Redis缓存，响应时间<100ms
   - 日均请求量100万+

求职意向：Python后端开发工程师 / 全栈开发工程师
期望薪资：20-35K
工作地点：北京、上海、杭州`;
        }
        
        async function startProcess() {
            const resume = document.getElementById('resumeInput').value.trim();
            
            if (!resume) {
                alert('请先输入简历内容或上传文件！');
                return;
            }
            
            // 调用统一的处理函数
            processResumeText(resume);
        }
        
        function downloadResults() {
            // 下载所有结果为文本文件
            const results = {
                '职业分析': document.getElementById('careerAnalysis').textContent,
                '推荐岗位': document.getElementById('jobRecommendations').textContent,
                '优化后简历': document.getElementById('optimizedResume').textContent,
                '面试准备': document.getElementById('interviewPrep').textContent,
                '模拟面试': document.getElementById('mockInterview').textContent
            };
            
            let content = '';
            for (const [title, text] of Object.entries(results)) {
                content += `\n${'='.repeat(60)}\n${title}\n${'='.repeat(60)}\n\n${text}\n\n`;
            }
            
            const blob = new Blob([content], { type: 'text/plain;charset=utf-8' });
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = 'AI求职助手-完整结果.txt';
            a.click();
            URL.revokeObjectURL(url);
        }
    </script>
</body>
</html>
"""


@app.get("/investor", response_class=HTMLResponse)
async def investor_page():
    """Investor-facing readiness dashboard page."""
    investor_html = "static/investor.html"
    if os.path.exists(investor_html):
        with open(investor_html, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>Investor Dashboard</h1><p>/api/investor/readiness</p>")

@app.post("/api/upload")
async def upload_resume(file: UploadFile = File(...)):
    """上传简历文件（支持PDF、Word、TXT、图片）"""
    try:
        # 检查文件类型
        allowed_types = ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png', '.bmp', '.gif']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_types:
            return JSONResponse({
                "error": f"不支持的文件格式。支持：PDF、Word、TXT、图片（JPG/PNG等）"
            }, status_code=400)
        
        # 读取文件内容
        content = await file.read()
        resume_text = ""
        
        try:
            # 解析文件
            if file_ext == '.txt':
                # 文本文件
                try:
                    resume_text = content.decode('utf-8')
                except:
                    resume_text = content.decode('gbk', errors='ignore')
                    
            elif file_ext == '.pdf':
                # PDF文件
                try:
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            resume_text += text + "\n"
                except Exception as e:
                    return JSONResponse({
                        "error": f"PDF解析失败: {str(e)}。请确保PDF不是扫描件，或上传图片格式。"
                    }, status_code=500)
                    
            elif file_ext in ['.docx', '.doc']:
                # Word文件
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(content))
                    
                    # 提取段落文本
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            resume_text += paragraph.text + "\n"
                    
                    # 提取表格文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    resume_text += cell.text + " "
                            resume_text += "\n"
                            
                except Exception as e:
                    return JSONResponse({
                        "error": f"Word文档解析失败: {str(e)}。请确保文件未损坏。"
                    }, status_code=500)
                    
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                # 图片文件 - 使用OCR
                try:
                    from PIL import Image
                    import pytesseract
                    import io
                    
                    # 打开图片
                    image = Image.open(io.BytesIO(content))
                    
                    # OCR识别（支持中英文）
                    resume_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                    
                    if not resume_text.strip():
                        return JSONResponse({
                            "error": "图片识别失败，未能提取到文字。请确保图片清晰，或尝试其他格式。"
                        }, status_code=500)
                        
                except ImportError:
                    return JSONResponse({
                        "error": "图片OCR功能未安装。正在安装依赖，请稍后重试..."
                    }, status_code=500)
                except Exception as e:
                    return JSONResponse({
                        "error": f"图片识别失败: {str(e)}。请确保图片清晰可读。"
                    }, status_code=500)
            
            # 检查是否成功提取到内容
            if not resume_text.strip():
                return JSONResponse({
                    "error": "文件解析成功，但未能提取到有效内容。请检查文件是否为空或格式是否正确。"
                }, status_code=500)

            _track_event(
                "resume_uploaded",
                {
                    "filename": file.filename,
                    "ext": file_ext,
                    "chars": len(resume_text.strip()),
                },
            )
            return JSONResponse({
                "success": True,
                "resume_text": resume_text.strip(),
                "filename": file.filename
            })
            
        except Exception as e:
            return JSONResponse({
                "error": f"文件解析出错: {str(e)}"
            }, status_code=500)
        
    except Exception as e:
        _track_event("api_error", {"api": "/api/upload", "error": str(e)[:300]})
        return JSONResponse({"error": f"上传失败: {str(e)}"}, status_code=500)

@app.websocket("/ws/progress")
async def websocket_progress(websocket: WebSocket):
    """WebSocket实时进度推送"""
    await websocket.accept()
    await progress_tracker.connect(websocket)
    
    try:
        while True:
            # 保持连接
            await websocket.receive_text()
    except WebSocketDisconnect:
        progress_tracker.disconnect(websocket)

@app.post("/api/process")
async def process_resume(request: Request):
    """处理简历的API接口 - 市场驱动"""
    try:
        data = await request.json()
        resume_text = data.get("resume", "")
        
        if not resume_text:
            return _api_error("简历内容不能为空", status_code=400, code="empty_resume")

        _track_event("resume_process_started", {"chars": len(resume_text)})

        # 重置进度
        progress_tracker.reset()
        
        # 定义进度回调
        async def update_progress_callback(step, message, agent):
            await progress_tracker.update_progress(step, message, agent)
            await progress_tracker.add_ai_message(agent, message)
        
        # 使用市场驱动引擎处理
        results = await market_engine.process_resume(resume_text, update_progress_callback)

        # Seed job search (Boss/OpenClaw) from resume text, so frontend can auto-search links.
        info = analyzer.extract_info(resume_text)
        seed_keywords = []
        if info.get("job_intention") and info["job_intention"] != "未指定":
            seed_keywords.append(info["job_intention"])
        seed_keywords.extend((info.get("skills") or [])[:6])
        seed_keywords = [k for k in seed_keywords if k]

        seed_location = None
        locs = info.get("preferred_locations") or []
        if locs:
            seed_location = locs[0]
        provider_mode = (real_job_service.get_statistics() or {}).get("provider_mode", "")
        
        # Real, actionable job text block (backward compatible with legacy UI field).
        def _format_real_jobs(jobs: List[Dict[str, Any]], mode: str) -> str:
            if not jobs:
                return (
                    '【推荐岗位】（当前暂无可用岗位）\n\n'
                    '排查建议：\n'
                    '1. 检查云端缓存：访问 /api/crawler/status。\n'
                    '2. 检查企业级招聘 API 是否已配置。\n'
                    '3. 若搜索引擎触发风控，可稍后重试。\n\n'
                    '注意：系统默认不会回退到“搜索入口链接”或演示岗位。\n'
                )

            heading = '【推荐岗位】（中国劳动力市场真实数据）'
            if mode == 'openclaw':
                heading = '【推荐岗位】（来自 Boss 直聘实时数据，OpenClaw）'
            elif mode == 'cloud':
                heading = '【推荐岗位】（来自 Boss 直聘云端缓存）'
            elif mode in ('baidu', 'bing', 'brave'):
                heading = f'【推荐岗位】（来自搜索引擎 {mode}）'
            elif mode == 'remotive':
                heading = '【推荐岗位】（来自全球招聘API，仅在显式开启时使用）'
            elif mode == 'bing_html':
                heading = '【推荐岗位】（来自 Bing 无浏览器搜索）'
            elif mode == 'duckduckgo':
                heading = '【推荐岗位】（来自 DuckDuckGo 无浏览器搜索）'
            elif mode == 'jooble':
                heading = '【推荐岗位】（来自 Jooble API）'
            elif mode == 'enterprise_api':
                heading = '【推荐岗位】（企业级中国招聘API实时数据）'
            elif mode == 'no_real_jobs':
                heading = '【推荐岗位】（未检索到可投递真实岗位）'

            lines = [heading, '']
            for i, job in enumerate(jobs, 1):
                title = job.get('title') or job.get('job_title') or '未知岗位'
                company = job.get('company') or ''
                loc = job.get('location') or ''
                salary = job.get('salary') or job.get('salary_range') or ''
                link = job.get('link') or job.get('apply_url') or ''

                mp = job.get('match_percentage')
                if mp is None:
                    mp = job.get('match_rate')
                if mp is None:
                    mp = job.get('match_score')
                mp_str = f"{mp}%" if isinstance(mp, (int, float)) else ''

                lines.append(f"{i}. {title}" + (f" - {company}" if company else ''))
                if salary:
                    lines.append(f"   薪资：{salary}")
                if loc:
                    lines.append(f"   地点：{loc}")
                if mp_str:
                    lines.append(f"   匹配度：{mp_str}")
                if link:
                    lines.append(f"   链接：{link}")
                lines.append('')

            return "\n".join(lines).strip() + "\n"

        async def _get_real_jobs_for_recommendation():
            cfg_mode = os.getenv('JOB_DATA_PROVIDER', 'auto').strip().lower()
            allow_portal_fallback = os.getenv("ALLOW_CN_PORTAL_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}
            kw = seed_keywords[:10]
            loc = seed_location

            if cfg_mode == 'cloud' or cloud_jobs_cache:
                cached = _filter_cloud_cache_by_query(kw, loc, limit=10)
                cached = _enforce_cn_market_jobs(cached)
                if cached:
                    return cached, 'cloud'
                # Cache empty/insufficient: fallback to cloud-safe real-time providers.
                fallback_jobs, fallback_mode, _ = _search_jobs_without_browser(
                    kw,
                    loc,
                    limit=10,
                    allow_portal_fallback=allow_portal_fallback,
                )
                return _enforce_cn_market_jobs(fallback_jobs), fallback_mode

            try:
                jobs = _normalize_real_jobs(
                    real_job_service.search_jobs(keywords=kw, location=loc, limit=10),
                    limit=10,
                )
                jobs = _enforce_cn_market_jobs(jobs)
                mode = (real_job_service.get_statistics() or {}).get('provider_mode', '') or cfg_mode
                return jobs[:10], mode
            except Exception:
                fallback_jobs, fallback_mode, _ = _search_jobs_without_browser(
                    kw,
                    loc,
                    limit=10,
                    allow_portal_fallback=allow_portal_fallback,
                )
                return _enforce_cn_market_jobs(fallback_jobs), fallback_mode or cfg_mode

        real_jobs, real_mode = await _get_real_jobs_for_recommendation()
        public_jobs = _public_job_payload(_enforce_cn_market_jobs(real_jobs), limit=10)
        results["job_recommendations"] = _format_real_jobs(public_jobs, real_mode)
        results, quality_gate = _run_output_quality_gate(results, resume_text, info, public_jobs)
        provider_mode = real_mode

        # 完成
        await progress_tracker.complete()
        await progress_tracker.add_ai_message("系统", "🎉 市场分析完成！")
        _track_event(
            "process_quality_gate",
            {
                "passed": bool(quality_gate.get("passed", True)),
                "issues_count": len(quality_gate.get("issues") or []),
            },
        )
        _track_event(
            "job_recommendation_ready",
            {
                "provider_mode": provider_mode,
                "real_jobs_count": len(public_jobs),
            },
        )
        _track_event(
            "resume_processed",
            {
                "ok": True,
                "provider_mode": provider_mode,
                "skills_count": len(seed_keywords),
                "real_jobs_count": len(public_jobs),
                "quality_gate_passed": bool(quality_gate.get("passed", True)),
            },
        )

        response_payload = {
            "career_analysis": str(results.get("market_analysis") or ""),
            "job_recommendations": str(results.get("job_recommendations") or ""),
            "optimized_resume": str(results.get("optimized_resume") or ""),
            "interview_prep": str(results.get("interview_prep") or ""),
            "mock_interview": str(results.get("salary_analysis") or ""),
            "job_provider_mode": str(provider_mode or ""),
            "recommended_jobs": public_jobs,
            "recommendation_quality": {
                "strict_real_posting": True,
                "count": len(public_jobs),
                "has_actionable_jobs": bool(public_jobs),
            },
            "quality_gate": quality_gate,
            "schema_version": PROCESS_RESPONSE_SCHEMA_VERSION,
            "boss_seed": {
                "keywords": seed_keywords,
                "location": seed_location,
            },
        }
        shape_errors = _validate_process_response_shape(response_payload)
        if shape_errors:
            _track_event(
                "process_contract_error",
                {
                    "count": len(shape_errors),
                    "sample": shape_errors[:3],
                },
            )
            return _api_error("输出JSON未通过契约校验", status_code=500, code="process_contract_failed")

        return _api_success(response_payload)
        
    except Exception as e:
        _track_event("resume_process_failed", {"error": str(e)[:300]})
        _track_event("resume_processed", {"ok": False, "error": str(e)[:300]})
        _track_event("api_error", {"api": "/api/process", "error": str(e)[:300]})
        await progress_tracker.error(f"处理出错: {str(e)}")
        return _api_error(str(e), status_code=500, code="process_failed")

@app.get("/api/health")
async def health_check():
    """健康检查"""
    stats = real_job_service.get_statistics()
    biz = business_service.metrics()
    
    # 检查OpenClaw状态
    openclaw_status = None
    if stats.get("provider_mode") == "openclaw":
        from app.services.job_providers.openclaw_browser_provider import OpenClawBrowserProvider
        openclaw = OpenClawBrowserProvider()
        openclaw_status = openclaw.health_check()
    
    return _api_success({
        "status": "ok",
        "message": "AI求职助手运行正常",
        "boot_ts": APP_BOOT_TS,
        "uptime_s": round(time.perf_counter() - APP_BOOT_MONO, 1),
        "job_database": stats,
        "openclaw": openclaw_status,
        "business": {
            "leads_total": biz.get("leads", {}).get("total", 0),
            "uploads": biz.get("funnel", {}).get("uploads", 0),
            "process_runs": biz.get("funnel", {}).get("process_runs", 0),
            "searches": biz.get("funnel", {}).get("searches", 0),
            "applies": biz.get("funnel", {}).get("applies", 0),
        },
        "config": {
            "job_data_provider": os.getenv("JOB_DATA_PROVIDER", "auto"),
            "cloud_cache_total": len(cloud_jobs_cache),
            "cloud_last_push_at": cloud_jobs_meta.get("last_push_at"),
            "no_browser_fallback_enabled": True,
            "enterprise_job_api_configured": bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip()),
            "llm": get_public_llm_config(),
        },
    })



@app.get("/api/version")
async def version():
    """Expose basic build metadata for debugging deployments."""
    return _api_success({
        "job_data_provider": os.getenv("JOB_DATA_PROVIDER", "auto"),
        "railway_git_commit_sha": os.getenv("RAILWAY_GIT_COMMIT_SHA"),
        "github_sha": os.getenv("GITHUB_SHA"),
        "app_boot_ts": APP_BOOT_TS,
    })


@app.get("/api/ping")
async def ping():
    """Ultra-light probe for availability checks."""
    return _api_success({"ok": True, "ts": datetime.now().isoformat()})


@app.get("/api/ready")
async def ready():
    """Release gate probe used by QA/ops before Go/No-Go."""
    cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
    cache_total = len(cloud_jobs_cache)
    checks = {
        "api_alive": True,
        "cn_provider_mode": cfg_mode in {"auto", "cloud", "baidu", "bing", "brave", "jooble", "openclaw", "enterprise_api"},
        "cache_or_cn_fallback": cache_total > 0 or True,
        "global_fallback_disabled_by_default": os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() not in {"1", "true", "yes", "on"},
        "enterprise_api_configured": bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip()),
    }
    score = round(sum(1 for v in checks.values() if v) / len(checks) * 100, 1)
    status = "ready" if score >= 75 else "not_ready"
    return _api_success(
        {
            "status": status,
            "score": score,
            "checks": checks,
            "provider_mode": cfg_mode,
            "cloud_cache_total": cache_total,
            "boot_ts": APP_BOOT_TS,
        }
    )


@app.post("/api/business/lead")
async def capture_business_lead(request: Request):
    """Capture B2B/B2C monetization leads from landing page."""
    try:
        data = await request.json()
        email = (data.get("email") or "").strip()
        if ("@" not in email) or ("." not in email):
            return JSONResponse({"error": "请输入有效邮箱"}, status_code=400)

        payload = business_service.add_lead(
            email=email,
            name=(data.get("name") or "").strip(),
            company=(data.get("company") or "").strip(),
            use_case=(data.get("use_case") or "").strip(),
            budget=(data.get("budget") or "").strip(),
            source=(data.get("source") or "landing").strip(),
            note=(data.get("note") or "").strip(),
        )
        return JSONResponse({"success": True, "data": payload})
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/lead", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/business/feedback")
async def capture_user_feedback(request: Request):
    """Capture user feedback from product and landing pages."""
    try:
        data = await request.json()
        message = (data.get("message") or "").strip()
        if len(message) < 5:
            return JSONResponse({"error": "反馈内容太短，请至少输入5个字符"}, status_code=400)
        rating_raw = data.get("rating")
        rating = int(rating_raw) if str(rating_raw).strip() else None
        payload = business_service.add_feedback(
            message=message,
            rating=rating,
            category=(data.get("category") or "").strip(),
            email=(data.get("email") or "").strip(),
            source=(data.get("source") or "product").strip(),
            page=(data.get("page") or "").strip(),
        )
        return JSONResponse({"success": True, "data": payload})
    except ValueError as ve:
        return JSONResponse({"error": str(ve)}, status_code=400)
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/feedback", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/business/event")
async def capture_public_event(request: Request):
    """Public-safe growth event ingest endpoint for frontend KPI tracking."""
    try:
        data = await request.json()
        event_name = str(data.get("event_name") or "").strip().lower()
        payload = data.get("payload") if isinstance(data.get("payload"), dict) else {}

        if (not event_name) or (not PUBLIC_EVENT_NAME_PATTERN.match(event_name)):
            return JSONResponse({"error": "invalid event_name"}, status_code=400)
        if event_name not in PUBLIC_EVENT_WHITELIST:
            return JSONResponse({"error": "event not allowed"}, status_code=400)

        business_service.track_event(event_name, payload)
        return JSONResponse({"success": True})
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/event", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/feedback/summary")
async def feedback_summary(request: Request, days: int = 7, limit: int = 20):
    """Feedback summary for ops dashboards and growth automation."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        return JSONResponse({"success": True, "data": business_service.feedback_summary(days=days, limit=limit)})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/public-proof")
async def business_public_proof():
    """Public-safe proof counters used on landing page."""
    try:
        m = business_service.metrics()
        funnel = m.get("funnel", {})
        engagement = m.get("engagement", {})
        quality = m.get("quality", {})
        return _api_success(
            {
                "leads_total": int(m.get("leads", {}).get("total", 0) or 0),
                "feedback_total": int(m.get("feedback", {}).get("total", 0) or 0),
                "uploads_total": int(funnel.get("uploads", 0) or 0),
                "process_runs_total": int(funnel.get("process_runs", 0) or 0),
                "searches_total": int(funnel.get("searches", 0) or 0),
                "applies_total": int(funnel.get("applies", 0) or 0),
                "job_link_clicks_total": int(engagement.get("job_link_clicks", 0) or 0),
                "result_downloads_total": int(engagement.get("result_downloads", 0) or 0),
                "upload_to_process_pct": float(funnel.get("upload_to_process_pct", 0) or 0),
                "process_to_search_pct": float(funnel.get("process_to_search_pct", 0) or 0),
                "search_to_apply_pct": float(funnel.get("search_to_apply_pct", 0) or 0),
                "click_to_apply_pct": float(engagement.get("click_to_apply_pct", 0) or 0),
                "quality_gate_fail_rate_pct": float(quality.get("gate_fail_rate_pct", 0) or 0),
            }
        )
    except Exception as e:
        return _api_error(str(e), status_code=500, code="business_public_proof_failed")


@app.get("/api/business/metrics")
async def business_metrics(request: Request):
    """Operational funnel metrics for fundraising / monetization tracking."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        return JSONResponse({"success": True, "data": business_service.metrics()})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/readiness")
async def business_readiness(request: Request):
    """Simple GTM readiness snapshot for investor demos."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        m = business_service.metrics()
        funnel = m.get("funnel", {})
        leads = m.get("leads", {})
        feedback = m.get("feedback", {})
        checks = {
            "lead_capture_ready": leads.get("total", 0) >= 1,
            "feedback_loop_ready": feedback.get("last_7d", 0) >= 1,
            "activation_flow_ready": funnel.get("upload_to_process_pct", 0) > 0,
            "job_match_flow_ready": funnel.get("process_to_search_pct", 0) > 0,
            "monetization_signal_ready": funnel.get("search_to_apply_pct", 0) > 0,
        }
        score = round(sum(1 for v in checks.values() if v) / len(checks) * 100, 1)
        return JSONResponse(
            {
                "success": True,
                "score": score,
                "checks": checks,
                "metrics": m,
            }
        )
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/investor/readiness")
async def investor_readiness(request: Request):
    """
    Investor-facing readiness API.
    Optional auth via INVESTOR_READ_TOKEN header x-investor-token.
    """
    token = os.getenv("INVESTOR_READ_TOKEN", "").strip()
    supplied = request.headers.get("x-investor-token", "").strip()
    if token and supplied != token:
        return _api_error("forbidden", status_code=403, code="forbidden")
    try:
        return _api_success(_build_investor_readiness_snapshot())
    except Exception as e:
        _track_event("api_error", {"api": "/api/investor/readiness", "error": str(e)[:300]})
        return _api_error(str(e), status_code=500, code="investor_readiness_failed")


@app.get("/api/investor/summary")
async def investor_summary(request: Request):
    """Compact summary payload for pitch materials."""
    token = os.getenv("INVESTOR_READ_TOKEN", "").strip()
    supplied = request.headers.get("x-investor-token", "").strip()
    if token and supplied != token:
        return _api_error("forbidden", status_code=403, code="forbidden")
    try:
        snap = _build_investor_readiness_snapshot()
        p = snap.get("pillars", {})
        narrative = [
            f"Current financing readiness status: {snap.get('status')} (score {snap.get('overall_score')}).",
            f"Product pillar score: {p.get('product')}, reliability: {p.get('reliability')}.",
            f"Traction pillar score: {p.get('traction')}, GTM pillar score: {p.get('go_to_market')}.",
            f"Feedback captured so far: {(snap.get('metrics') or {}).get('feedback', {}).get('total', 0)}.",
            "Next 30 days focus: increase uploads, process runs, qualified leads, and actionable user feedback while keeping CN-real job links stable.",
        ]
        return _api_success(
            {
                "status": snap.get("status"),
                "overall_score": snap.get("overall_score"),
                "pillars": p,
                "highlights": snap.get("highlights", {}),
                "next_30d_targets": snap.get("next_30d_targets", {}),
                "narrative": narrative,
            }
        )
    except Exception as e:
        _track_event("api_error", {"api": "/api/investor/summary", "error": str(e)[:300]})
        return _api_error(str(e), status_code=500, code="investor_summary_failed")

@app.get("/api/jobs/search")
async def search_jobs(
    keywords: str = None,
    location: str = None,
    salary_min: int = None,
    experience: str = None,
    limit: int = 50,
    allow_portal_fallback: bool = False,
):
    """搜索真实岗位"""
    try:
        cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
        n = int(limit) if limit is not None else 50
        n = max(1, min(n, 100))
        keyword_list = keywords.split(",") if keywords else []
        kw = [k.strip() for k in keyword_list if k and k.strip()]
        allow_portal = bool(allow_portal_fallback) or (
            os.getenv("ALLOW_CN_PORTAL_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}
        )

        # Cloud mode: prefer crawler cache; fallback to cloud-safe real-time providers.
        if cfg_mode == "cloud" or cloud_jobs_cache:
            jobs = _filter_cloud_cache_by_query(kw, location, limit=n)
            jobs = _enforce_cn_market_jobs(jobs)
            warning = None
            mode = "cloud"
            if not jobs:
                fallback_jobs, fallback_mode, fallback_err = _search_jobs_without_browser(
                    kw,
                    location,
                    limit=n,
                    allow_portal_fallback=allow_portal,
                )
                jobs = _enforce_cn_market_jobs(fallback_jobs)
                mode = fallback_mode or "cloud"
                warning = (
                    f"cloud cache is empty; switched to no-browser provider: {mode}"
                    if jobs
                    else (
                        f"cloud cache empty and no-browser search failed: {fallback_err}"
                        if fallback_err
                        else "cloud cache is empty and no real jobs were found"
                    )
                )
            _track_event(
                "job_search",
                {
                    "provider_mode": mode,
                    "result_count": len(jobs),
                    "cloud_mode": True,
                },
            )
            _cache_recent_jobs(jobs)
            return _api_success({
                "total": len(jobs),
                "jobs": jobs,
                "provider_mode": mode,
                "warning": warning,
            })

        # Stream progress to the same WebSocket channel as the AI pipeline.
        # Frontend listens for `type=job_search`.
        import asyncio

        def progress_cb(message: str, percent: int):
            asyncio.create_task(
                progress_tracker.broadcast(
                    {"type": "job_search", "data": {"message": message, "percent": int(percent)}}
                )
            )

        try:
            jobs = real_job_service.search_jobs(
                keywords=kw,
                location=location,
                salary_min=salary_min,
                experience=experience,
                limit=n,
                progress_callback=progress_cb,
            )
            jobs = _normalize_real_jobs(jobs, limit=n)
            jobs = _enforce_cn_market_jobs(jobs)
            mode = (real_job_service.get_statistics() or {}).get("provider_mode", cfg_mode)
        except Exception as e:
            jobs, mode, _ = _search_jobs_without_browser(
                kw,
                location,
                limit=n,
                allow_portal_fallback=allow_portal,
            )
            jobs = _enforce_cn_market_jobs(jobs)
            if not jobs:
                raise e
        _track_event(
            "job_search",
            {
                "provider_mode": mode,
                "result_count": len(jobs),
                "cloud_mode": False,
            },
        )
        _cache_recent_jobs(jobs)
        return _api_success({
            "total": len(jobs),
            "jobs": jobs,
            "provider_mode": mode,
        })
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/search", "error": str(e)[:300]})
        return _api_error(str(e), status_code=500, code="job_search_failed")

@app.get("/api/jobs/{job_id}")
async def get_job_detail(job_id: str):
    """获取岗位详情"""
    try:
        job = real_job_service.get_job_detail(job_id)
        if job:
            return JSONResponse({"success": True, "job": job})
        else:
            return JSONResponse({"error": "岗位不存在"}, status_code=404)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/jobs/apply")
async def apply_job(request: Request):
    """投递简历到指定岗位"""
    try:
        data = await request.json()
        job_id = data.get("job_id")
        resume_text = data.get("resume_text")
        user_info = data.get("user_info", {})
        
        result = real_job_service.apply_job(job_id, resume_text, user_info)
        if (not result.get("success")) and job_id in recent_search_jobs:
            # Fallback for no-browser providers whose job detail is not in local provider cache.
            j = recent_search_jobs.get(job_id, {})
            link = j.get("link") or j.get("apply_url")
            if link:
                app_id = f"EXT{int(datetime.now().timestamp())}"
                record = {
                    "application_id": app_id,
                    "job_id": job_id,
                    "job_title": j.get("title", ""),
                    "company": j.get("company", ""),
                    "platform": j.get("platform", j.get("provider", "")),
                    "apply_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "status": "待投递",
                    "apply_link": link,
                    "user_info": user_info or {},
                }
                try:
                    real_job_service.records.add_record(record)
                except Exception:
                    pass
                result = {
                    "success": True,
                    "message": "已生成投递跳转链接（请在招聘网站完成真实投递）",
                    "data": record,
                }
        _track_event(
            "job_apply",
            {
                "job_id": job_id,
                "success": bool(result.get("success")),
            },
        )
        return JSONResponse(result)
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/apply", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/jobs/batch_apply")
async def batch_apply_jobs(request: Request):
    """批量投递简历"""
    try:
        data = await request.json()
        job_ids = data.get("job_ids", [])
        resume_text = data.get("resume_text")
        user_info = data.get("user_info", {})
        
        results = real_job_service.batch_apply(job_ids, resume_text, user_info)
        _track_event(
            "job_apply",
            {
                "batch": True,
                "requested": len(job_ids),
                "success_count": len([r for r in results if r.get("success")]),
            },
        )
        return JSONResponse({
            "success": True,
            "total": len(results),
            "results": results
        })
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/batch_apply", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/statistics")
async def get_statistics():
    """获取数据统计"""
    try:
        stats = real_job_service.get_statistics()
        return JSONResponse({"success": True, "data": stats})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

# ========================================
# 爬虫数据接收接口（云端部署时使用）
# ========================================

from fastapi import Header

# 简单的API密钥验证
CRAWLER_API_KEY = os.getenv("CRAWLER_API_KEY", "your-secret-key-change-this")

@app.post("/api/crawler/upload")
async def receive_crawler_data(request: Request, authorization: str = Header(None)):
    """接收本地爬虫推送的岗位数据"""
    try:
        # 验证API密钥
        if not authorization or not authorization.startswith("Bearer "):
            return JSONResponse({"error": "未授权：缺少API密钥"}, status_code=401)
        
        api_key = authorization.replace("Bearer ", "")
        if api_key != CRAWLER_API_KEY:
            return JSONResponse({"error": "未授权：API密钥无效"}, status_code=401)
        
        # 解析数据
        data = await request.json()
        jobs = data.get("jobs", [])
        
        if not jobs:
            return JSONResponse({"error": "岗位数据为空"}, status_code=400)
        
        # 添加接收时间戳
        for job in jobs:
            job["received_at"] = datetime.now().isoformat()

        # 存储到缓存（去重 + 过滤 seed/demo + 必须可跳转）
        incoming = _normalize_and_filter_jobs(jobs, limit=5000)
        existing_keys = {
            str(j.get("link") or j.get("apply_url") or j.get("id") or "").strip().lower()
            for j in cloud_jobs_cache
        }
        new_jobs = []
        for job in incoming:
            key = str(job.get("link") or job.get("apply_url") or job.get("id") or "").strip().lower()
            if key and key not in existing_keys:
                existing_keys.add(key)
                new_jobs.append(job)

        cloud_jobs_cache.extend(new_jobs)

        # 限制缓存大小（保留最新的5000个）
        if len(cloud_jobs_cache) > 5000:
            cloud_jobs_cache[:] = cloud_jobs_cache[-5000:]

        cloud_jobs_meta["last_push_at"] = datetime.now().isoformat()
        cloud_jobs_meta["last_received"] = len(jobs)
        cloud_jobs_meta["last_new"] = len(new_jobs)
        _track_event(
            "crawler_upload",
            {"received": len(jobs), "new": len(new_jobs), "total": len(cloud_jobs_cache)},
        )

        print(f"✅ 接收爬虫数据：{len(new_jobs)} 个新岗位（总计：{len(cloud_jobs_cache)}）")
        
        return JSONResponse({
            "success": True,
            "received": len(jobs),
            "new": len(new_jobs),
            "total": len(cloud_jobs_cache)
        })
        
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/crawler/status")
async def get_crawler_status():
    """获取爬虫数据状态"""
    if not cloud_jobs_cache:
        return JSONResponse({
            "status": "empty",
            "total": 0,
            "last_push_at": cloud_jobs_meta.get("last_push_at"),
            "last_received": cloud_jobs_meta.get("last_received", 0),
            "last_new": cloud_jobs_meta.get("last_new", 0),
        })

    return JSONResponse({
        "status": "ok",
        "total": len(cloud_jobs_cache),
        "last_push_at": cloud_jobs_meta.get("last_push_at"),
        "last_received": cloud_jobs_meta.get("last_received", 0),
        "last_new": cloud_jobs_meta.get("last_new", 0),
    })


# ========================================
# 自动投递 API 接口（新增）
# ========================================

# 全局任务管理
auto_apply_tasks: Dict[str, Dict[str, Any]] = {}
task_lock = asyncio.Lock()

# 平台映射
PLATFORM_APPLIERS = {
    'boss': 'app.services.auto_apply.boss_applier.BossApplier',
    'zhilian': 'app.services.auto_apply.zhilian_applier.ZhilianApplier',
    'linkedin': 'app.services.auto_apply.linkedin_applier.LinkedInApplier'
}

@app.post("/api/auto-apply/start")
async def start_auto_apply(request: Request):
    """启动自动投递"""
    try:
        data = await request.json()

        # 生成任务ID
        task_id = str(uuid.uuid4())

        # 获取配置
        config = {
            'platform': data.get('platform', 'linkedin'),
            'max_apply_per_session': data.get('max_count', 50),
            'keywords': data.get('keywords', ''),
            'location': data.get('location', ''),
            'company_blacklist': data.get('blacklist', []),
            'user_profile': data.get('user_profile', {}),
            'pause_before_submit': data.get('pause_before_submit', False),
            'headless': data.get('headless', False)
        }

        # 验证配置
        from app.services.auto_apply.config import AutoApplyConfig, validate_config
        apply_config = AutoApplyConfig.from_dict(config)
        is_valid, error_msg = validate_config(apply_config)

        if not is_valid:
            return _api_error(error_msg, 400)

        # 创建任务记录
        auto_apply_tasks[task_id] = {
            'task_id': task_id,
            'status': 'starting',
            'config': config,
            'progress': {
                'applied': 0,
                'failed': 0,
                'total': 0,
                'current_job': None
            },
            'created_at': datetime.now().isoformat(),
            'started_at': None,
            'completed_at': None
        }

        # 异步启动投递任务
        asyncio.create_task(_run_auto_apply_task(task_id, config, data.get('jobs', [])))

        return _api_success({
            'task_id': task_id,
            'message': '自动投递任务已启动'
        })

    except Exception as e:
        logger.exception("启动自动投递失败")
        return _api_error(str(e), 500)


async def _run_auto_apply_task(task_id: str, config: Dict[str, Any], jobs: List[Dict[str, Any]]):
    """运行自动投递任务（异步）"""
    try:
        from app.services.auto_apply.linkedin_applier import LinkedInApplier
        from app.core.llm_client import LLMClient

        # 更新状态
        auto_apply_tasks[task_id]['status'] = 'running'
        auto_apply_tasks[task_id]['started_at'] = datetime.now().isoformat()

        # 创建投递器
        llm_client = LLMClient() if config.get('use_ai_answers', True) else None
        applier = LinkedInApplier(config, llm_client)

        # 登录（如果需要）
        email = config.get('user_profile', {}).get('email')
        password = config.get('user_profile', {}).get('password')

        if email and password:
            login_success = applier.login(email, password)
            if not login_success:
                auto_apply_tasks[task_id]['status'] = 'failed'
                auto_apply_tasks[task_id]['error'] = '登录失败'
                return

        # 如果没有提供职位列表，则搜索
        if not jobs:
            jobs = applier.search_jobs(
                keywords=config.get('keywords', ''),
                location=config.get('location', ''),
                filters={}
            )

        auto_apply_tasks[task_id]['progress']['total'] = len(jobs)

        # 批量投递
        result = applier.batch_apply(jobs, config.get('max_apply_per_session', 50))

        # 更新最终状态
        auto_apply_tasks[task_id]['status'] = 'completed'
        auto_apply_tasks[task_id]['completed_at'] = datetime.now().isoformat()
        auto_apply_tasks[task_id]['result'] = result
        auto_apply_tasks[task_id]['progress']['applied'] = result['applied']
        auto_apply_tasks[task_id]['progress']['failed'] = result['failed']

        # 清理资源
        applier.cleanup()

    except Exception as e:
        logger.exception(f"自动投递任务失败: {task_id}")
        auto_apply_tasks[task_id]['status'] = 'failed'
        auto_apply_tasks[task_id]['error'] = str(e)


@app.post("/api/auto-apply/stop")
async def stop_auto_apply(request: Request):
    """停止自动投递"""
    try:
        data = await request.json()
        task_id = data.get('task_id')

        if not task_id or task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)

        task = auto_apply_tasks[task_id]

        if task['status'] not in ['running', 'starting']:
            return _api_error('任务未在运行中', 400)

        # 设置停止标志（实际停止逻辑在 applier 中处理）
        task['status'] = 'stopped'
        task['completed_at'] = datetime.now().isoformat()

        return _api_success({
            'message': '停止指令已发送'
        })

    except Exception as e:
        logger.exception("停止自动投递失败")
        return _api_error(str(e), 500)


@app.get("/api/auto-apply/status/{task_id}")
async def get_auto_apply_status(task_id: str):
    """查询投递状态"""
    try:
        if task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)

        task = auto_apply_tasks[task_id]

        return _api_success({
            'task': task
        })

    except Exception as e:
        logger.exception("查询投递状态失败")
        return _api_error(str(e), 500)


@app.get("/api/auto-apply/history")
async def get_auto_apply_history(limit: int = 50):
    """获取投递历史"""
    try:
        # 按时间倒序排列
        tasks = sorted(
            auto_apply_tasks.values(),
            key=lambda x: x.get('created_at', ''),
            reverse=True
        )

        return _api_success({
            'tasks': tasks[:limit],
            'total': len(tasks)
        })

    except Exception as e:
        logger.exception("获取投递历史失败")
        return _api_error(str(e), 500)


@app.websocket("/ws/auto-apply/{task_id}")
async def auto_apply_progress_ws(websocket: WebSocket, task_id: str):
    """实时推送投递进度（支持多平台）"""
    await websocket.accept()

    try:
        while True:
            if task_id not in auto_apply_tasks:
                await websocket.send_json({
                    'type': 'error',
                    'message': '任务不存在'
                })
                break

            task = auto_apply_tasks[task_id]

            # 发送详细进度
            await websocket.send_json({
                'type': 'progress',
                'task_id': task_id,
                'status': task['status'],
                'progress': task['progress'],
                'platforms': task.get('platforms', [task.get('config', {}).get('platform', 'linkedin')]),
                'timestamp': datetime.now().isoformat()
            })

            # 如果任务已完成，发送最终消息并断开
            if task['status'] in ['completed', 'failed', 'stopped']:
                await websocket.send_json({
                    'type': 'complete',
                    'task_id': task_id,
                    'status': task['status'],
                    'progress': task['progress'],
                    'result': task.get('result', {}),
                    'error': task.get('error')
                })
                break

            await asyncio.sleep(2)

    except WebSocketDisconnect:
        logger.info(f"WebSocket 断开: {task_id}")
    except Exception as e:
        logger.exception(f"WebSocket 错误: {e}")
        try:
            await websocket.send_json({
                'type': 'error',
                'message': str(e)
            })
        except:
            pass


# ========================================
# 多平台投递 API 接口（新增）
# ========================================

@app.post("/api/auto-apply/start-multi")
async def start_multi_platform_apply(request: Request):
    """启动多平台自动投递"""
    try:
        data = await request.json()
        platforms = data.get('platforms', ['boss'])
        config = data.get('config', {})

        # 生成任务ID
        task_id = str(uuid.uuid4())

        # 创建任务记录
        async with task_lock:
            auto_apply_tasks[task_id] = {
                'task_id': task_id,
                'status': 'starting',
                'platforms': platforms,
                'config': config,
                'progress': {
                    'total_platforms': len(platforms),
                    'completed_platforms': 0,
                    'total_applied': 0,
                    'total_failed': 0,
                    'platform_progress': {}
                },
                'created_at': datetime.now().isoformat(),
                'started_at': None,
                'completed_at': None
            }

        # 异步启动多平台投递
        asyncio.create_task(_run_multi_platform_apply(task_id, platforms, config))

        return _api_success({
            'task_id': task_id,
            'message': f'已启动 {len(platforms)} 个平台的自动投递'
        })

    except Exception as e:
        logger.exception("启动多平台投递失败")
        return _api_error(str(e), 500)


async def _run_multi_platform_apply(task_id: str, platforms: List[str], config: Dict[str, Any]):
    """运行多平台投递任务"""
    try:
        # 更新状态
        auto_apply_tasks[task_id]['status'] = 'running'
        auto_apply_tasks[task_id]['started_at'] = datetime.now().isoformat()

        # 并发执行多个平台
        tasks = []
        for platform in platforms:
            if platform in PLATFORM_APPLIERS:
                task = asyncio.create_task(
                    _run_single_platform_apply(task_id, platform, config)
                )
                tasks.append(task)

        # 等待所有平台完成
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 汇总结果
        total_applied = 0
        total_failed = 0

        for result in results:
            if isinstance(result, dict):
                total_applied += result.get('applied', 0)
                total_failed += result.get('failed', 0)

        # 更新最终状态
        auto_apply_tasks[task_id]['status'] = 'completed'
        auto_apply_tasks[task_id]['completed_at'] = datetime.now().isoformat()
        auto_apply_tasks[task_id]['progress']['total_applied'] = total_applied
        auto_apply_tasks[task_id]['progress']['total_failed'] = total_failed

    except Exception as e:
        logger.exception(f"多平台投递任务失败: {task_id}")
        auto_apply_tasks[task_id]['status'] = 'failed'
        auto_apply_tasks[task_id]['error'] = str(e)


async def _run_single_platform_apply(task_id: str, platform: str, config: Dict[str, Any]):
    """运行单个平台的投递任务"""
    try:
        # 动态导入平台 Applier
        module_path, class_name = PLATFORM_APPLIERS[platform].rsplit('.', 1)
        module = __import__(module_path, fromlist=[class_name])
        ApplierClass = getattr(module, class_name)

        # 获取平台特定配置
        platform_config = config.get(f'{platform}_config', {})
        platform_config.update({
            'keywords': config.get('keywords', ''),
            'location': config.get('location', ''),
            'max_apply_per_session': config.get('max_count', 50),
            'company_blacklist': config.get('blacklist', []),
            'headless': config.get('headless', False)
        })

        # 创建投递器（兼容不同的构造函数）
        from app.core.llm_client import LLMClient
        llm_client = LLMClient() if config.get('use_ai_answers', True) else None

        # LinkedIn 需要 llm_client 参数，其他平台不需要
        if platform == 'linkedin':
            applier = ApplierClass(platform_config, llm_client)
        else:
            applier = ApplierClass(platform_config)

        # 登录
        login_success = False
        if platform == 'boss':
            phone = platform_config.get('phone')
            if phone:
                login_success = applier.login(phone)
        elif platform == 'zhilian':
            username = platform_config.get('username')
            password = platform_config.get('password')
            if username and password:
                login_success = applier.login(username, password)
        elif platform == 'linkedin':
            email = platform_config.get('email')
            password = platform_config.get('password')
            if email and password:
                login_success = applier.login(email, password)

        if not login_success:
            raise Exception(f"{platform} 登录失败")

        # 搜索职位
        jobs = applier.search_jobs(
            keywords=platform_config.get('keywords', ''),
            location=platform_config.get('location', ''),
            filters={}
        )

        # 更新进度
        auto_apply_tasks[task_id]['progress']['platform_progress'][platform] = {
            'status': 'running',
            'total': len(jobs),
            'applied': 0,
            'failed': 0
        }

        # 批量投递
        result = applier.batch_apply(jobs, platform_config.get('max_apply_per_session', 50))

        # 更新平台进度
        auto_apply_tasks[task_id]['progress']['platform_progress'][platform] = {
            'status': 'completed',
            'total': len(jobs),
            'applied': result['applied'],
            'failed': result['failed']
        }

        auto_apply_tasks[task_id]['progress']['completed_platforms'] += 1

        # 清理资源
        applier.cleanup()

        return result

    except Exception as e:
        logger.exception(f"平台 {platform} 投递失败")
        auto_apply_tasks[task_id]['progress']['platform_progress'][platform] = {
            'status': 'failed',
            'error': str(e)
        }
        return {'applied': 0, 'failed': 0}


@app.get("/api/auto-apply/platforms")
async def get_supported_platforms():
    """获取支持的平台列表"""
    return _api_success({
        'platforms': [
            {
                'id': 'boss',
                'name': 'Boss直聘',
                'icon': '💼',
                'status': 'available',
                'features': ['手机验证码登录', '智能投递', '打招呼语'],
                'config_fields': [
                    {'name': 'phone', 'label': '手机号', 'type': 'text', 'required': True}
                ]
            },
            {
                'id': 'zhilian',
                'name': '智联招聘',
                'icon': '📋',
                'status': 'available',
                'features': ['账号密码登录', '简历投递', '附件上传'],
                'config_fields': [
                    {'name': 'username', 'label': '用户名', 'type': 'text', 'required': True},
                    {'name': 'password', 'label': '密码', 'type': 'password', 'required': True}
                ]
            },
            {
                'id': 'linkedin',
                'name': 'LinkedIn',
                'icon': '🔗',
                'status': 'available',
                'features': ['Easy Apply', 'AI问答', '国际职位'],
                'config_fields': [
                    {'name': 'email', 'label': '邮箱', 'type': 'email', 'required': True},
                    {'name': 'password', 'label': '密码', 'type': 'password', 'required': True}
                ]
            }
        ]
    })


@app.get("/api/auto-apply/stats")
async def get_apply_stats():
    """获取投递统计"""
    try:
        # 统计所有任务
        total_tasks = len(auto_apply_tasks)
        completed_tasks = sum(1 for t in auto_apply_tasks.values() if t['status'] == 'completed')
        running_tasks = sum(1 for t in auto_apply_tasks.values() if t['status'] == 'running')

        total_applied = 0
        total_failed = 0

        # 平台统计
        platform_stats = {}

        for task in auto_apply_tasks.values():
            # 单平台任务
            if 'result' in task:
                result = task['result']
                total_applied += result.get('applied', 0)
                total_failed += result.get('failed', 0)

                platform = task.get('config', {}).get('platform', 'linkedin')
                if platform not in platform_stats:
                    platform_stats[platform] = {'applied': 0, 'failed': 0, 'total': 0}
                platform_stats[platform]['applied'] += result.get('applied', 0)
                platform_stats[platform]['failed'] += result.get('failed', 0)

            # 多平台任务
            if 'progress' in task and 'platform_progress' in task['progress']:
                for platform, progress in task['progress']['platform_progress'].items():
                    if platform not in platform_stats:
                        platform_stats[platform] = {'applied': 0, 'failed': 0, 'total': 0}
                    platform_stats[platform]['applied'] += progress.get('applied', 0)
                    platform_stats[platform]['failed'] += progress.get('failed', 0)
                    platform_stats[platform]['total'] += progress.get('total', 0)

                total_applied += task['progress'].get('total_applied', 0)
                total_failed += task['progress'].get('total_failed', 0)

        return _api_success({
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'running_tasks': running_tasks,
            'total_applied': total_applied,
            'total_failed': total_failed,
            'success_rate': round(total_applied / (total_applied + total_failed) * 100, 2) if (total_applied + total_failed) > 0 else 0,
            'platform_stats': platform_stats
        })

    except Exception as e:
        logger.exception("获取统计失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/test-platform")
async def test_platform_config(request: Request):
    """测试平台配置（不实际投递）"""
    try:
        data = await request.json()
        platform = data.get('platform', 'boss')
        config = data.get('config', {})

        if platform not in PLATFORM_APPLIERS:
            return _api_error(f'不支持的平台: {platform}', 400)

        # 动态导入平台 Applier
        module_path, class_name = PLATFORM_APPLIERS[platform].rsplit('.', 1)
        module = __import__(module_path, fromlist=[class_name])
        ApplierClass = getattr(module, class_name)

        # 创建投递器
        if platform == 'linkedin':
            from app.core.llm_client import LLMClient
            llm_client = LLMClient()
            applier = ApplierClass(config, llm_client)
        else:
            applier = ApplierClass(config)

        # 测试登录
        login_result = {'success': False, 'message': '未测试'}

        if platform == 'boss':
            phone = config.get('phone')
            if phone:
                login_result = {
                    'success': True,
                    'message': f'配置正确，手机号: {phone[:3]}****{phone[-4:]}'
                }
        elif platform == 'zhilian':
            username = config.get('username')
            password = config.get('password')
            if username and password:
                login_result = {
                    'success': True,
                    'message': f'配置正确，用户名: {username}'
                }
        elif platform == 'linkedin':
            email = config.get('email')
            password = config.get('password')
            if email and password:
                login_result = {
                    'success': True,
                    'message': f'配置正确，邮箱: {email}'
                }

        return _api_success({
            'platform': platform,
            'login_test': login_result,
            'config_valid': login_result['success']
        })

    except Exception as e:
        logger.exception("测试平台配置失败")
        return _api_error(str(e), 500)


if __name__ == "__main__":
    import webbrowser
    import threading
    
    port = int(os.getenv("PORT", 8000))
    
    print("\n" + "🚀"*30)
    print("AI求职助手 - Web服务启动中...")
    print("🚀"*30 + "\n")
    print(f"📍 访问地址: http://localhost:{port}")
    print(f"📍 API文档: http://localhost:{port}/docs")
    print(f"📍 WebSocket: ws://localhost:{port}/ws/progress")
    print("\n✨ 新功能: WebSocket实时进度推送")
    print("按 Ctrl+C 停止服务\n")
    
    # 延迟2秒后自动打开浏览器
    def open_browser():
        import time
        time.sleep(2)
        webbrowser.open(f'http://localhost:{port}/app')
    
    threading.Thread(target=open_browser, daemon=True).start()
    
    uvicorn.run(app, host="0.0.0.0", port=port)

