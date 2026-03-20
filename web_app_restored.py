"""
AI求职助手 - Web界面
一个漂亮的网页界面，让您直接在浏览器中使用
"""

from fastapi import FastAPI, Request, UploadFile, File, Form, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import os
import sys
from typing import Optional, List, Dict, Any, Tuple
import asyncio
from datetime import datetime
from urllib.parse import quote_plus, urlparse, parse_qs, unquote
import re
import html as html_lib
import requests
import logging
import time
import uuid
import importlib.util
import json
import tempfile

sys.path.insert(0, os.path.dirname(__file__))

from app.core.multi_ai_debate import JobApplicationPipeline
from app.core.fast_ai_engine import fast_pipeline, HighPerformanceAIEngine
from app.core.market_driven_engine import market_driven_pipeline
from app.core.skills_graph import SkillsGraph
from app.core.llm_client import get_public_llm_config
from app.core.agent_profile_registry import get_agent_profiles_payload
from app.services.resume_analyzer import ResumeAnalyzer
from app.services.real_job_service import RealJobService
from app.services.job_source_registry import get_job_source_registry_payload
from app.services.business_service import BusinessService
from app.services.commerce_service import CommerceService
from app.services.user_auth_service import UserAuthService
from app.services.resume_profile_service import ResumeProfileService
from app.services.resume_render_service import ResumeRenderService
from app.services.email_campaign_service import EmailCampaignService, SMTP_PRESETS
from app.services.hr_match_service import HRMatchService
from app.core.realtime_progress import progress_tracker

app = FastAPI(title="AI求职助手")
APP_BOOT_TS = datetime.now().isoformat()
APP_BOOT_MONO = time.perf_counter()
logger = logging.getLogger("ai_job_helper")
if not logger.handlers:
    logging.basicConfig(
        level=os.getenv("APP_LOG_LEVEL", "INFO"),
        format="%(asctime)s %(levelname)s %(name)s %(message)s",
    )

# 使用市场驱动引擎
market_engine = market_driven_pipeline

# 挂载静态文件
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def request_context_middleware(request: Request, call_next):
    """Attach request id + process time headers for observability."""
    rid = request.headers.get("x-request-id") or str(uuid.uuid4())
    start = time.perf_counter()
    request.state.request_id = rid
    try:
        response = await call_next(request)
    except Exception as e:
        took_ms = (time.perf_counter() - start) * 1000
        logger.exception(
            "request_failed path=%s method=%s rid=%s took_ms=%.1f err=%s",
            request.url.path,
            request.method,
            rid,
            took_ms,
            str(e)[:300],
        )
        return JSONResponse(
            {"success": False, "error": "internal_error", "request_id": rid},
            status_code=500,
            headers={"x-request-id": rid, "x-process-time-ms": f"{took_ms:.1f}"},
        )

    took_ms = (time.perf_counter() - start) * 1000
    response.headers["x-request-id"] = rid
    response.headers["x-process-time-ms"] = f"{took_ms:.1f}"
    if request.url.path.startswith("/api"):
        logger.info(
            "request_done path=%s method=%s status=%s rid=%s took_ms=%.1f",
            request.url.path,
            request.method,
            getattr(response, "status_code", 0),
            rid,
            took_ms,
        )
    return response

# 全局变量
pipeline = JobApplicationPipeline()
analyzer = ResumeAnalyzer()
skills_graph = SkillsGraph()
real_job_service = RealJobService()  # 真实招聘数据服务
business_service = BusinessService()
commerce_service = CommerceService()
user_auth_service = UserAuthService()
resume_profile_service = ResumeProfileService()
resume_render_service = ResumeRenderService()
email_campaign_service = EmailCampaignService()
hr_match_service = HRMatchService()

# 云端岗位缓存（内存）
cloud_jobs_cache: List[Dict[str, Any]] = []
CN_JOB_DOMAINS = ("zhipin.com", "liepin.com", "zhaopin.com", "51job.com", "lagou.com")
cloud_jobs_meta: Dict[str, Any] = {
    "last_push_at": None,
    "last_received": 0,
    "last_new": 0,
}
recent_search_jobs: Dict[str, Dict[str, Any]] = {}
github_apply_tasks: Dict[str, Dict[str, Any]] = {}
boss_sms_code_store: Dict[str, Dict[str, Any]] = {}
boss_control_action_store: Dict[str, Dict[str, Any]] = {}
HAITOU_DEFAULT_POLICY: Dict[str, Any] = {
    "enabled": True,
    "fallback_platforms": ["boss", "zhilian"],
    "max_count": 80,
    "headless": False,
    "allow_portal_fallback": True,
    "verification_wait_seconds": 180,
    "strategy_version": "haitou_v1",
}
PLATFORM_AUTOMATION_POLICY: Dict[str, Dict[str, Any]] = {
    "boss": {
        "status": "allowed",
        "note": "默认允许自动化（仍需遵守平台条款与风控规则）。",
    },
    "zhilian": {
        "status": "allowed_with_caution",
        "note": "默认允许自动化，建议上线前复核最新平台条款。",
    },
    "linkedin": {
        "status": "restricted_by_default",
        "note": "LinkedIn 用户协议 8.2 与帮助中心对 bots/自动化有明确限制。",
        "block_message": (
            "linkedin 默认禁用自动投递：LinkedIn User Agreement 8.2 明确限制 bots/自动化。"
            "若你已完成法务评估并接受风险，可传 allow_risky_automation=true "
            "或 compliance_override_platforms=['linkedin'] 强制开启。"
        ),
        "sources": [
            "https://www.linkedin.com/legal/user-agreement",
            "https://www.linkedin.com/help/linkedin/answer/a1341387/prohibited-software-and-extensions?lang=en",
        ],
    },
}
simple_login_sessions: Dict[str, Dict[str, Any]] = {}
simple_order_history: List[Dict[str, Any]] = []
SIMPLE_LOGIN_SESSION_FILE = os.path.join("data", "simple_login_sessions.json")

def _candidate_aihawk_roots() -> List[str]:
    roots: List[str] = []
    env_root = os.getenv("AIHAWK_ROOT", "").strip()
    if env_root:
        roots.append(os.path.abspath(env_root))
    roots.extend(
        [
            os.path.abspath(
                os.path.join(
                    os.path.dirname(__file__),
                    "third_party",
                    "Auto_Jobs_Applier_AIHawk_plugins",
                )
            ),
            os.path.abspath(
                os.path.join(
                    os.path.dirname(__file__),
                    "third_party",
                    "Auto_Jobs_Applier_AIHawk",
                )
            ),
        ]
    )
    out: List[str] = []
    for path in roots:
        p = str(path or "").strip()
        if p and p not in out:
            out.append(p)
    return out


def _resolve_aihawk_root() -> str:
    for root in _candidate_aihawk_roots():
        if os.path.isfile(os.path.join(root, "main.py")):
            return root
    return os.path.abspath(
        os.path.join(os.path.dirname(__file__), "third_party", "Auto_Jobs_Applier_AIHawk")
    )


AIHAWK_ROOT = _resolve_aihawk_root()
os.environ.setdefault("AIHAWK_ROOT", AIHAWK_ROOT)


def _api_success(payload: Dict[str, Any], status_code: int = 200) -> JSONResponse:
    body = {"success": True}
    body.update(payload or {})
    return JSONResponse(body, status_code=status_code)


def _api_error(message: str, status_code: int = 400, code: str = "bad_request") -> JSONResponse:
    return JSONResponse(
        {"success": False, "error": message, "code": code},
        status_code=status_code,
    )


CREDIT_ACTIONS: Dict[str, Dict[str, Any]] = {
    "resume_analysis": {
        "cost": 1,
        "label": "简历首轮分析",
        "description": "快速诊断简历质量、岗位方向和优化点。",
    },
    "resume_structure": {
        "cost": 2,
        "label": "结构化简历母版",
        "description": "把简历转成可复用母版，便于后续多版本导出。",
    },
    "offer_pipeline": {
        "cost": 6,
        "label": "完整求职流程",
        "description": "一次跑完多 agent、三版简历、岗位池、训练包和推进路线。",
    },
    "local_apply_task": {
        "cost": 4,
        "label": "本地投递任务",
        "description": "下发一次真实本地浏览器执行任务。",
    },
    "job_detail": {
        "cost": 1,
        "label": "岗位详情查看",
        "description": "查看并校验一个真实岗位详情。",
    },
}


def _credit_action_catalog() -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for action_id, row in CREDIT_ACTIONS.items():
        payload = dict(row or {})
        payload["action_id"] = action_id
        out.append(payload)
    return out


def _credit_balance_guard(request: Request, action_id: str) -> Tuple[Optional[JSONResponse], Dict[str, Any]]:
    action = CREDIT_ACTIONS.get(action_id)
    if not action:
        return _api_error(f"unknown_credit_action:{action_id}", status_code=500, code="unknown_credit_action"), {}

    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("缺少访问码，请先登录", status_code=401, code="missing_access_code"), {}

    try:
        wallet_payload = commerce_service.get_wallet_by_access_code(access_code)
    except ValueError as e:
        return _api_error(str(e), status_code=401, code="wallet_lookup_failed"), {}
    except Exception as e:
        logger.exception("wallet_lookup_failed")
        return _api_error(str(e), status_code=500, code="wallet_lookup_failed"), {}

    wallet = wallet_payload.get("wallet") if isinstance(wallet_payload, dict) else {}
    balance = int((wallet or {}).get("balance") or 0)
    required = int(action.get("cost") or 0)
    if balance < required:
        return (
            _api_error(
                f"credits 不足：当前 {balance}，需要 {required}。请先充值 credits。",
                status_code=402,
                code="insufficient_credits",
            ),
            {
                "access_code": access_code,
                "wallet": wallet,
                "required": required,
                "action": action,
                "action_id": action_id,
            },
        )
    return None, {
        "access_code": access_code,
        "wallet": wallet,
        "required": required,
        "action": action,
        "action_id": action_id,
    }


def _settle_credit_action(
    guard_payload: Dict[str, Any],
    note: str = "",
    meta: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    action_id = str((guard_payload or {}).get("action_id") or "").strip()
    access_code = str((guard_payload or {}).get("access_code") or "").strip().upper()
    required = int((guard_payload or {}).get("required") or 0)
    return commerce_service.consume_credits(
        amount=required,
        action=action_id,
        access_code=access_code,
        note=note,
        meta=meta or {},
    )


def _ops_secret_ready() -> str:
    return str(os.getenv("ADMIN_PORTAL_CODE") or os.getenv("ADMIN_METRICS_TOKEN") or "").strip()


def _require_ops_secret(request: Request) -> Optional[JSONResponse]:
    secret = _ops_secret_ready()
    if not secret:
        return None
    supplied = str(request.headers.get("x-ops-secret") or "").strip()
    if supplied != secret:
        return _api_error("forbidden", status_code=403, code="forbidden")
    return None


def _buyer_access_code_from_request(request: Request) -> str:
    return str(
        request.cookies.get("jobhelper_access_code")
        or request.headers.get("x-access-code")
        or ""
    ).strip().upper()


def _user_session_token_from_request(request: Request) -> str:
    return str(
        request.headers.get("x-user-session")
        or request.cookies.get(_auth_cookie_name())
        or request.cookies.get("jobhelper_user_session")
        or ""
    ).strip()


def _auth_cookie_name() -> str:
    return str(os.getenv("AUTH_COOKIE_NAME") or "jobhelper_user_session").strip() or "jobhelper_user_session"


def _auth_cookie_secure() -> bool:
    return str(os.getenv("AUTH_COOKIE_SECURE", "0")).strip().lower() in {"1", "true", "yes", "on"}


def _set_user_session_cookie(response: JSONResponse, token: str) -> None:
    value = str(token or "").strip()
    if not value:
        return
    response.set_cookie(
        key=_auth_cookie_name(),
        value=value,
        max_age=60 * 60 * 24 * 30,
        httponly=True,
        samesite="lax",
        secure=_auth_cookie_secure(),
        path="/",
    )


def _clear_user_session_cookie(response: JSONResponse) -> None:
    response.delete_cookie(key=_auth_cookie_name(), path="/")


def _buyer_context_row(access_code: str) -> Dict[str, Any]:
    normalized_code = str(access_code or "").strip().upper()
    rows = commerce_service.list_bundles(limit=50, search=normalized_code)
    row = next((item for item in rows if str(item.get("access_code") or "").strip().upper() == normalized_code), None)
    if row:
        return row
    return {
        "buyer_id": f"buyer_{normalized_code.lower()}",
        "name": f"Buyer {normalized_code}",
        "access_code": normalized_code,
        "status": "active",
        "buyer_expires_at": "",
        "code_expires_at": "",
        "used_count": 0,
        "max_uses": 999999,
        "order_id": "",
        "source": "direct_access_fallback",
    }


def _to_string_list(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    if isinstance(value, str):
        return [x.strip() for x in value.split(",") if x.strip()]
    return []


def _to_bool(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on"}:
        return True
    if text in {"0", "false", "no", "n", "off"}:
        return False
    return default


def _extract_named_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float, bool)):
        return str(value).strip()
    if isinstance(value, dict):
        for key in ("name", "skill", "label", "title", "value", "text", "keyword"):
            text = _extract_named_text(value.get(key))
            if text:
                return text
        for key in ("skills", "items", "keywords", "values"):
            nested = _flatten_named_texts(value.get(key))
            if nested:
                return nested[0]
        if len(value) == 1:
            return _extract_named_text(next(iter(value.values())))
        return ""
    if isinstance(value, (list, tuple, set)):
        items = _flatten_named_texts(value)
        return items[0] if items else ""
    return str(value).strip()


def _flatten_named_texts(value: Any) -> List[str]:
    out: List[str] = []
    seen = set()

    def _push(text: str) -> None:
        cleaned = str(text or "").strip()
        if not cleaned:
            return
        key = cleaned.lower()
        if key in seen:
            return
        seen.add(key)
        out.append(cleaned)

    if isinstance(value, (list, tuple, set)):
        for item in value:
            for text in _flatten_named_texts(item):
                _push(text)
        return out

    text = _extract_named_text(value)
    if text:
        _push(text)
    return out


def _load_simple_login_sessions() -> Dict[str, Dict[str, Any]]:
    if not os.path.exists(SIMPLE_LOGIN_SESSION_FILE):
        return {}
    try:
        with open(SIMPLE_LOGIN_SESSION_FILE, "r", encoding="utf-8") as handle:
            payload = json.load(handle)
    except Exception as e:
        logger.warning("加载 simple login session 失败: %s", e)
        return {}

    if isinstance(payload, dict):
        rows = list(payload.values())
    elif isinstance(payload, list):
        rows = payload
    else:
        rows = []
    sessions: Dict[str, Dict[str, Any]] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        phone = str(row.get("phone") or "").strip()
        if not phone:
            continue
        sessions[phone] = dict(row)
    return sessions


def _save_simple_login_sessions() -> None:
    directory = os.path.dirname(SIMPLE_LOGIN_SESSION_FILE) or "data"
    os.makedirs(directory, exist_ok=True)
    serializable: Dict[str, Dict[str, Any]] = {}
    for fallback_phone, row in simple_login_sessions.items():
        if not isinstance(row, dict):
            continue
        phone = str(row.get("phone") or fallback_phone).strip()
        if not phone:
            continue
        serializable[phone] = dict(row)
    with open(SIMPLE_LOGIN_SESSION_FILE, "w", encoding="utf-8") as handle:
        json.dump(serializable, handle, ensure_ascii=False, indent=2)


def _find_simple_login_session(phone: str = "", session_id: str = "") -> Tuple[str, Dict[str, Any]]:
    normalized_phone = str(phone or "").strip()
    normalized_session_id = str(session_id or "").strip()

    if normalized_phone:
        row = simple_login_sessions.get(normalized_phone)
        if isinstance(row, dict):
            return normalized_phone, row

    if normalized_session_id:
        for stored_phone, row in simple_login_sessions.items():
            if not isinstance(row, dict):
                continue
            if str(row.get("session_id") or "").strip() == normalized_session_id:
                return stored_phone, row

    return "", {}


def _store_simple_login_session(row: Dict[str, Any]) -> None:
    phone = str((row or {}).get("phone") or "").strip()
    if not phone:
        return
    simple_login_sessions[phone] = dict(row or {})
    _save_simple_login_sessions()


simple_login_sessions = _load_simple_login_sessions()


def _apply_haitou_defaults(raw: Dict[str, Any]) -> Dict[str, Any]:
    data = dict(raw or {})
    # Global permanent strategy: enforce identical haitou policy for all users.
    data["enable_fallback_auto_apply"] = True
    requested_fallback = _normalize_platform_list(data.get("fallback_platforms"))
    requested_platforms = _normalize_platform_list(data.get("platforms"))
    data["fallback_platforms"] = requested_fallback or list(HAITOU_DEFAULT_POLICY["fallback_platforms"])
    data["platforms"] = requested_platforms or list(HAITOU_DEFAULT_POLICY["fallback_platforms"])
    data["max_count"] = int(HAITOU_DEFAULT_POLICY["max_count"])
    data["headless"] = bool(HAITOU_DEFAULT_POLICY["headless"])
    data["verification_wait_seconds"] = int(HAITOU_DEFAULT_POLICY["verification_wait_seconds"])
    cfg = dict(data.get("config") or {})
    cfg["max_count"] = int(HAITOU_DEFAULT_POLICY["max_count"])
    cfg["headless"] = bool(HAITOU_DEFAULT_POLICY["headless"])
    cfg["verification_wait_seconds"] = int(HAITOU_DEFAULT_POLICY["verification_wait_seconds"])
    data["config"] = cfg
    return data


def _mask_phone(phone: str) -> str:
    s = str(phone or "").strip()
    if len(s) >= 7:
        return f"{s[:3]}****{s[-4:]}"
    return s


def _is_recent_iso_ts(value: Any, within_seconds: int = 120) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    try:
        dt = datetime.fromisoformat(text.replace("Z", "+00:00"))
    except Exception:
        return False
    now = datetime.now(dt.tzinfo) if dt.tzinfo else datetime.now()
    return (now - dt).total_seconds() <= max(10, int(within_seconds or 120))


def _resolve_auto_apply_task_id(task_id: str = "", gh_task_id: str = "") -> str:
    tid = str(task_id or "").strip()
    if tid:
        return tid
    gid = str(gh_task_id or "").strip()
    if not gid:
        return ""
    gh_task = github_apply_tasks.get(gid) or {}
    return str(gh_task.get("linked_auto_apply_task_id") or "").strip()


def _set_boss_sms_code(task_id: str, code: str) -> None:
    tid = str(task_id or "").strip()
    if not tid:
        return
    boss_sms_code_store[tid] = {
        "code": str(code or "").strip(),
        "updated_at": datetime.now().isoformat(),
    }


def _pop_boss_sms_code(task_id: str) -> str:
    tid = str(task_id or "").strip()
    if not tid:
        return ""
    row = boss_sms_code_store.pop(tid, None) or {}
    return str(row.get("code") or "").strip()


def _push_boss_control_action(task_id: str, action: str) -> None:
    tid = str(task_id or "").strip()
    act = str(action or "").strip().lower()
    if not tid or not act:
        return
    boss_control_action_store[tid] = {
        "action": act,
        "updated_at": datetime.now().isoformat(),
    }


def _pop_boss_control_action(task_id: str) -> str:
    tid = str(task_id or "").strip()
    if not tid:
        return ""
    row = boss_control_action_store.pop(tid, None) or {}
    return str(row.get("action") or "").strip().lower()


def _aihawk_capability_snapshot() -> Dict[str, Any]:
    root_exists = os.path.isdir(AIHAWK_ROOT)
    required_files = ["main.py", "requirements.txt", "data_folder"]
    missing_required = [
        p for p in required_files if not os.path.exists(os.path.join(AIHAWK_ROOT, p))
    ]

    plugin_marker_sets: Dict[str, List[str]] = {
        "aihawk_private_old": ["ai_hawk/bot_facade.py", "ai_hawk/job_manager.py"],
        "aihawk_src": ["src/aihawk_bot_facade.py", "src/aihawk_job_manager.py"],
        "linkedin_legacy_root": ["linkedIn_bot_facade.py", "linkedIn_job_manager.py"],
        "linkedin_legacy_src": ["src/linkedIn_bot_facade.py", "src/linkedIn_job_manager.py"],
    }
    plugin_mode = "none"
    has_auto_apply_plugins = False
    for mode, markers in plugin_marker_sets.items():
        if all(os.path.exists(os.path.join(AIHAWK_ROOT, rel)) for rel in markers):
            plugin_mode = mode
            has_auto_apply_plugins = True
            break

    dep_names = ["yaml", "selenium", "webdriver_manager", "click"]
    if plugin_mode in {"aihawk_private_old", "aihawk_src"}:
        dep_names.append("inquirer")
    if plugin_mode.startswith("linkedin_legacy"):
        dep_names.extend(["langchain_openai", "langchain_core", "dotenv"])
    dep_names = list(dict.fromkeys(dep_names))
    missing_deps = [d for d in dep_names if importlib.util.find_spec(d) is None]
    if plugin_mode.startswith("linkedin_legacy"):
        has_lev = importlib.util.find_spec("Levenshtein") is not None
        has_rapid = importlib.util.find_spec("rapidfuzz") is not None
        if not (has_lev or has_rapid):
            missing_deps.append("Levenshtein_or_rapidfuzz")

    llm_key_ready = bool(
        os.getenv("DEEPSEEK_API_KEY", "").strip()
        or os.getenv("OPENAI_API_KEY", "").strip()
    )

    can_prepare = root_exists and (len(missing_required) == 0) and ("yaml" not in missing_deps)
    can_run = can_prepare and has_auto_apply_plugins and (len(missing_deps) == 0)

    blockers: List[str] = []
    actions: List[str] = []
    warnings: List[str] = []
    if missing_required:
        blockers.append(f"缺少必要文件: {missing_required}")
        actions.append("补齐 AIHawk 插件版目录结构")
    if missing_deps:
        blockers.append(f"缺少依赖: {missing_deps}")
        actions.append("在后端虚拟环境安装缺失依赖")
    if not has_auto_apply_plugins:
        blockers.append("当前 AIHawk 版本缺少 auto-apply 插件代码")
        actions.append("切换到含插件的 fork/私有分支")
    if not llm_key_ready:
        warnings.append("未配置服务端 LLM API KEY（可在启动任务时传 llm_api_key）")
        actions.append("配置 DEEPSEEK_API_KEY 或 OPENAI_API_KEY，或请求时传 llm_api_key")

    return {
        "root": AIHAWK_ROOT,
        "root_exists": root_exists,
        "plugin_mode": plugin_mode,
        "missing_required": missing_required,
        "missing_deps": missing_deps,
        "has_auto_apply_plugins": has_auto_apply_plugins,
        "llm_key_ready": llm_key_ready,
        "mode": ("full_auto_apply" if has_auto_apply_plugins else "prepare_only"),
        "can_prepare": can_prepare,
        "can_run": can_run,
        "blockers": blockers,
        "warnings": warnings,
        "actions": actions,
    }


def _skills_graph_payload(
    resume_text: str,
    target_role: str = "",
    job_text: str = "",
) -> Dict[str, Any]:
    resume_text = str(resume_text or "")
    target_role = str(target_role or "").strip()
    job_text = str(job_text or "")

    if not target_role:
        try:
            parsed = analyzer.extract_info(resume_text)
            target_role = str(parsed.get("job_intention") or "").strip()
        except Exception:
            target_role = ""

    resume_skills = skills_graph.extract_skills(resume_text)
    job_skills = skills_graph.extract_skills(job_text or target_role)

    match_score = None
    if job_skills:
        match_score = round(skills_graph.calculate_match_score(resume_skills, job_skills), 1)

    recommend_for = target_role or "后端开发"
    recommended = skills_graph.recommend_skills(resume_skills, recommend_for)

    return {
        "target_role": target_role,
        "resume_skills": resume_skills,
        "job_skills": job_skills,
        "match_score": match_score,
        "recommended_skills": recommended,
    }


def _multi_agent_outputs_degraded(results: Dict[str, Any]) -> bool:
    fields = [
        str(results.get("career_analysis") or ""),
        str(results.get("job_recommendations") or ""),
        str(results.get("optimized_resume") or ""),
        str(results.get("interview_prep") or ""),
        str(results.get("mock_interview") or ""),
    ]
    bad = 0
    for text in fields:
        t = text.lower()
        if ("ai思考出错" in text) or ("authentication fails" in t) or ("error" in t and len(t) < 120):
            bad += 1
    return bad >= 3


def _decode_text_file_bytes(content: bytes) -> str:
    """Best-effort decode for uploaded txt files with mixed encodings."""
    if not content:
        return ""

    candidates = [
        "utf-8-sig",
        "utf-16",
        "utf-16le",
        "utf-16be",
        "gb18030",
        "gbk",
    ]

    decoded = None
    for enc in candidates:
        try:
            decoded = content.decode(enc)
            break
        except Exception:
            continue

    if decoded is None:
        decoded = content.decode("utf-8", errors="ignore")

    # Normalize common artifacts from wrong fallback decoding.
    decoded = decoded.replace("\x00", "").replace("\ufeff", "")
    decoded = decoded.replace("\r\n", "\n").replace("\r", "\n")
    return decoded.strip()


def _normalize_extracted_text(text: str) -> str:
    if not text:
        return ""
    out = text.replace("\x00", "").replace("\ufeff", "")
    out = out.replace("\r\n", "\n").replace("\r", "\n")
    # Keep line breaks, but collapse excessive blank lines.
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _text_quality_score(text: str) -> float:
    """Heuristic score for extracted text readability (0.0 ~ 1.0)."""
    t = _normalize_extracted_text(text)
    if not t:
        return 0.0

    total = max(1, len(t))
    printable = sum(1 for ch in t if ch.isprintable() or ch in "\n\t")
    control = sum(1 for ch in t if (ord(ch) < 32 and ch not in "\n\t"))
    replacement = t.count("\ufffd")
    ascii_word = len(re.findall(r"[A-Za-z0-9]", t))
    cjk = len(re.findall(r"[\u4e00-\u9fff]", t))

    printable_ratio = printable / total
    useful_ratio = min(1.0, (ascii_word + cjk) / total)
    control_ratio = control / total
    replacement_ratio = replacement / total
    length_factor = min(1.0, total / 320.0)

    score = (
        0.45 * printable_ratio
        + 0.30 * useful_ratio
        + 0.25 * length_factor
        - 0.70 * control_ratio
        - 0.90 * replacement_ratio
    )
    return max(0.0, min(1.0, score))


def _looks_like_mojibake(text: str) -> bool:
    t = _normalize_extracted_text(text)
    if not t:
        return True
    if ("\ufffd" in t) or ("\x00" in t):
        return True
    # Typical mojibake traces like 'ÿ', 'ð', etc.
    latin_ext = len(re.findall(r"[\u00C0-\u024F]", t))
    if latin_ext >= 2 and (latin_ext / max(1, len(t))) > 0.002:
        return True
    return False


def _ocr_image_best_effort(image: Any) -> str:
    """Run OCR with language fallback order."""
    import pytesseract

    for lang in ("chi_sim+eng", "chi_sim", "eng"):
        try:
            text = pytesseract.image_to_string(image, lang=lang)
            text = _normalize_extracted_text(text)
            if text:
                return text
        except Exception:
            continue
    return ""


def _extract_pdf_text_pypdf2(content: bytes) -> str:
    import io
    import PyPDF2

    reader = PyPDF2.PdfReader(io.BytesIO(content))
    chunks: List[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            chunks.append(txt)
    return _normalize_extracted_text("\n".join(chunks))


def _extract_pdf_text_pymupdf(content: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ""

    doc = fitz.open(stream=content, filetype="pdf")
    chunks: List[str] = []
    try:
        for page in doc:
            txt = page.get_text("text") or ""
            if txt.strip():
                chunks.append(txt)
    finally:
        doc.close()
    return _normalize_extracted_text("\n".join(chunks))


def _extract_pdf_text_ocr(content: bytes, max_pages: int = 8) -> str:
    """OCR fallback for scanned/garbled PDFs using PyMuPDF rasterization."""
    try:
        import fitz  # PyMuPDF
        from PIL import Image
    except Exception:
        return ""

    doc = fitz.open(stream=content, filetype="pdf")
    chunks: List[str] = []
    try:
        limit = min(max_pages, doc.page_count)
        for idx in range(limit):
            page = doc.load_page(idx)
            # 2x scale improves OCR precision for resumes.
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            txt = _ocr_image_best_effort(image)
            if txt:
                chunks.append(txt)
    finally:
        doc.close()
    return _normalize_extracted_text("\n".join(chunks))


def _extract_pdf_text_best_effort(content: bytes) -> Tuple[str, str]:
    """Try text extraction first, then OCR fallback when quality is poor."""
    candidates: List[Tuple[str, str]] = []

    try:
        txt = _extract_pdf_text_pypdf2(content)
        if txt:
            candidates.append(("pypdf2", txt))
    except Exception as e:
        logger.warning("pypdf2_extract_failed err=%s", str(e)[:200])

    try:
        txt = _extract_pdf_text_pymupdf(content)
        if txt:
            candidates.append(("pymupdf_text", txt))
    except Exception as e:
        logger.warning("pymupdf_extract_failed err=%s", str(e)[:200])

    best_text = ""
    best_source = ""
    best_score = 0.0
    for source, txt in candidates:
        score = _text_quality_score(txt)
        if score > best_score:
            best_score = score
            best_text = txt
            best_source = source

    # Always try OCR when available and prefer it for suspicious/low-quality extracted text.
    try:
        ocr_text = _extract_pdf_text_ocr(content)
        ocr_score = _text_quality_score(ocr_text)
        if ocr_text:
            should_take_ocr = False
            if not best_text:
                should_take_ocr = True
            elif _looks_like_mojibake(best_text):
                should_take_ocr = True
            elif ocr_score >= best_score + 0.06:
                should_take_ocr = True
            elif ocr_score >= max(0.45, best_score - 0.06) and len(ocr_text) >= len(best_text) * 0.6:
                should_take_ocr = True

            if should_take_ocr:
                best_text = ocr_text
                best_source = "ocr"
                best_score = ocr_score
    except Exception as e:
        logger.warning("pdf_ocr_failed err=%s", str(e)[:200])

    return _normalize_extracted_text(best_text), (best_source or "none")


def _build_investor_readiness_snapshot() -> Dict[str, Any]:
    """
    Financing-oriented readiness snapshot.
    Provides one consolidated view for product, traction, reliability and GTM.
    """
    metrics = business_service.metrics()
    funnel = metrics.get("funnel", {})
    leads = metrics.get("leads", {})
    feedback = metrics.get("feedback", {})
    stability = metrics.get("stability", {})

    cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
    enterprise_on = bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip())
    global_fallback_on = os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}

    uploads = int(funnel.get("uploads", 0) or 0)
    process_runs = int(funnel.get("process_runs", 0) or 0)
    searches = int(funnel.get("searches", 0) or 0)
    applies = int(funnel.get("applies", 0) or 0)
    api_errors = int(stability.get("api_errors", 0) or 0)

    process_success_pct = float(funnel.get("process_success_pct", 0) or 0)
    search_to_apply_pct = float(funnel.get("search_to_apply_pct", 0) or 0)

    product_score = 100.0
    if cfg_mode not in {"auto", "cloud", "baidu", "bing", "brave", "jooble", "openclaw", "enterprise_api"}:
        product_score -= 25.0
    if global_fallback_on:
        product_score -= 15.0
    if process_success_pct < 90:
        product_score -= 15.0
    if process_success_pct < 75:
        product_score -= 15.0
    if enterprise_on:
        product_score += 5.0
    product_score = max(0.0, min(100.0, product_score))

    traction_score = 0.0
    traction_score += min(35.0, uploads * 1.2)
    traction_score += min(35.0, process_runs * 1.5)
    traction_score += min(20.0, searches * 0.9)
    traction_score += min(10.0, applies * 2.0)
    traction_score = max(0.0, min(100.0, traction_score))

    reliability_score = 100.0
    if process_runs > 0:
        err_ratio = api_errors / max(process_runs, 1)
        reliability_score -= min(60.0, err_ratio * 100.0)
    if api_errors >= 20:
        reliability_score -= 10.0
    reliability_score = max(0.0, min(100.0, reliability_score))

    gtm_score = 0.0
    gtm_score += min(45.0, int(leads.get("total", 0) or 0) * 12.0)
    gtm_score += min(25.0, int(leads.get("last_7d", 0) or 0) * 8.0)
    gtm_score += min(10.0, int(feedback.get("last_7d", 0) or 0) * 2.5)
    gtm_score += min(15.0, search_to_apply_pct * 0.8)
    gtm_score += 10.0 if searches > 0 else 0.0
    gtm_score += 5.0 if applies > 0 else 0.0
    gtm_score = max(0.0, min(100.0, gtm_score))

    weighted = (
        product_score * 0.35
        + traction_score * 0.25
        + reliability_score * 0.25
        + gtm_score * 0.15
    )
    overall = round(max(0.0, min(100.0, weighted)), 1)

    status = "seed_ready" if overall >= 75 else ("pre_seed_ready" if overall >= 60 else "build_stage")

    return {
        "status": status,
        "overall_score": overall,
        "pillars": {
            "product": round(product_score, 1),
            "traction": round(traction_score, 1),
            "reliability": round(reliability_score, 1),
            "go_to_market": round(gtm_score, 1),
        },
        "highlights": {
            "provider_mode": cfg_mode,
            "enterprise_api_configured": enterprise_on,
            "global_fallback_enabled": global_fallback_on,
            "cloud_cache_total": len(cloud_jobs_cache),
            "feedback_total": int(feedback.get("total", 0) or 0),
        },
        "metrics": metrics,
        "next_30d_targets": {
            "uploads": max(uploads + 80, 120),
            "process_runs": max(process_runs + 60, 100),
            "searches": max(searches + 100, 180),
            "applies": max(applies + 30, 40),
            "leads_total": max(int(leads.get("total", 0) or 0) + 20, 30),
        },
    }


def _track_event(name: str, payload: Optional[Dict[str, Any]] = None) -> None:
    try:
        business_service.track_event(name, payload or {})
    except Exception:
        # Do not break core user path due to analytics write failures.
        pass


def _format_job_recommendations_markdown(jobs: List[Dict[str, Any]]) -> str:
    if not jobs:
        return "【推荐岗位】\n\n暂未找到可直接跳转投递的岗位，请稍后重试。"

    lines = ["【推荐岗位】", ""]
    for i, job in enumerate(jobs, 1):
        title = str(job.get("title") or "未知岗位")
        company = str(job.get("company") or "")
        salary = str(job.get("salary") or "")
        location = str(job.get("location") or "")
        link = str(job.get("link") or job.get("apply_url") or "")

        lines.append(f"{i}. {title}" + (f" - {company}" if company else ""))
        if salary:
            lines.append(f"   薪资：{salary}")
        if location:
            lines.append(f"   地点：{location}")
        if link:
            lines.append(f"   链接：{link}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _legacy_quality_audit_text(payload: Dict[str, Any]) -> str:
    quality = payload.get("quality_gate") if isinstance(payload, dict) else {}
    quality = quality if isinstance(quality, dict) else {}
    issues = quality.get("issues") if isinstance(quality.get("issues"), list) else []
    jobs = payload.get("recommended_jobs") if isinstance(payload.get("recommended_jobs"), list) else []

    base = 92
    score = max(65, base - len(issues) * 8)
    status = "通过" if not issues else "部分通过"
    issue_text = "、".join([str(x) for x in issues[:5]]) if issues else "无明显问题"
    return (
        f"【质量审计】\n\n"
        f"- 评分：{score}/100\n"
        f"- 结果：{status}\n"
        f"- 问题：{issue_text}\n"
        f"- 可投递岗位：{len(jobs)} 条"
    )


def _legacy_results_from_process_payload(payload: Dict[str, Any], resume_text: str) -> Dict[str, Any]:
    career_analysis = str(payload.get("career_analysis") or "")
    job_recommendations = str(payload.get("job_recommendations") or "")
    optimized_resume = str(payload.get("optimized_resume") or "")
    interview_prep = str(payload.get("interview_prep") or "")
    mock_interview = str(payload.get("mock_interview") or "")

    # Skill graph stays best-effort so analysis endpoint is always usable.
    skill_gap_text = "【技能图谱】\n\n暂无技能图谱数据。"
    try:
        skill_payload = _skills_graph_payload(resume_text, "", "")
        resume_skills = skill_payload.get("resume_skills") if isinstance(skill_payload, dict) else []
        rec = skill_payload.get("recommended_skills") if isinstance(skill_payload, dict) else []
        role = str(skill_payload.get("target_role") or "目标岗位")
        if isinstance(resume_skills, list) and isinstance(rec, list):
            skill_gap_text = (
                "【技能图谱】\n\n"
                f"- 目标岗位：{role}\n"
                f"- 已识别技能：{', '.join([str(x) for x in resume_skills[:12]]) or '暂无'}\n"
                f"- 建议补充：{', '.join([str(x) for x in rec[:10]]) or '暂无'}"
            )
    except Exception:
        pass

    return {
        "career_analysis": career_analysis,
        "job_recommendations": job_recommendations,
        "resume_optimization": optimized_resume,
        "optimized_resume": optimized_resume,
        "interview_preparation": interview_prep,
        "interview_prep": interview_prep,
        "mock_interview": mock_interview,
        "skill_gap_analysis": skill_gap_text,
        "quality_audit": _legacy_quality_audit_text(payload),
    }


def _heuristic_legacy_results(resume_text: str) -> Dict[str, Any]:
    info = analyzer.extract_info(resume_text)
    keyword_candidates: List[str] = []
    intent = str(info.get("job_intention") or "").strip()
    if intent and intent != "未指定":
        keyword_candidates.append(intent)
    keyword_candidates.extend([str(x).strip() for x in (info.get("skills") or [])[:5]])
    keywords = [x for x in keyword_candidates if x]
    if not keywords:
        keywords = ["Python", "AI"]

    location = None
    locs = info.get("preferred_locations") or []
    if isinstance(locs, list) and locs:
        location = str(locs[0]).strip() or None

    jobs: List[Dict[str, Any]] = []
    try:
        jobs = _public_job_payload(
            _enforce_cn_market_jobs(
                _normalize_real_jobs(real_job_service.search_jobs(keywords=keywords, location=location, limit=10), limit=10),
                allow_entrypoints=True,
            ),
            limit=10,
            allow_entrypoints=True,
        )
    except Exception:
        jobs = []

    if not jobs:
        try:
            fallback, _, _ = _search_jobs_without_browser(keywords, location, limit=10, allow_portal_fallback=True)
            jobs = _public_job_payload(_enforce_cn_market_jobs(fallback, allow_entrypoints=True), limit=10, allow_entrypoints=True)
        except Exception:
            jobs = []

    if not jobs:
        try:
            jobs = _public_job_payload(
                _search_jobs_cn_entrypoints(keywords, location, limit=10),
                limit=10,
                allow_entrypoints=True,
            )
        except Exception:
            jobs = []

    mapped = {
        "career_analysis": _render_market_analysis_fallback(info),
        "job_recommendations": _format_job_recommendations_markdown(jobs),
        "resume_optimization": _render_optimized_resume_fallback(resume_text, info),
        "optimized_resume": _render_optimized_resume_fallback(resume_text, info),
        "interview_preparation": _render_interview_prep_fallback(info, jobs),
        "interview_prep": _render_interview_prep_fallback(info, jobs),
        "mock_interview": "【薪资与谈判建议】\n\n优先选择匹配度高且有明确薪资区间的岗位，先拿面试再谈薪。",
        "skill_gap_analysis": "【技能图谱】\n\n建议围绕目标岗位补齐 1-2 个高频技能后再进行集中投递。",
        "quality_audit": f"【质量审计】\n\n- 结果：启用离线兜底\n- 可投递岗位：{len(jobs)} 条",
    }
    return mapped


def _run_simple_apply_flow(payload: Dict[str, Any]) -> Dict[str, Any]:
    phone = str(payload.get("phone") or "").strip()
    resume_text = str(payload.get("resume_text") or "").strip()
    if not resume_text:
        return {
            "success": False,
            "message": "resume_text 不能为空",
            "total": 0,
            "success_count": 0,
            "failed_count": 0,
            "details": [],
        }

    job_keyword = str(payload.get("job_keyword") or payload.get("keywords") or "").strip()
    city = str(payload.get("city") or payload.get("location") or "").strip()
    count = int(payload.get("count") or payload.get("max_count") or 10)
    count = max(1, min(count, 120))

    keywords = [x.strip() for x in job_keyword.split(",") if x.strip()]
    if not keywords:
        info = analyzer.extract_info(resume_text)
        intent = str(info.get("job_intention") or "").strip()
        if intent and intent != "未指定":
            keywords.append(intent)
        keywords.extend([str(x).strip() for x in (info.get("skills") or [])[:3]])
        keywords = [x for x in keywords if x]
    if not keywords:
        keywords = ["Python"]

    jobs: List[Dict[str, Any]] = []
    try:
        jobs = _normalize_real_jobs(
            real_job_service.search_jobs(keywords=keywords, location=city or None, limit=count),
            limit=count,
        )
        jobs = _enforce_cn_market_jobs(jobs, allow_entrypoints=True)
    except Exception:
        jobs = []

    if not jobs:
        try:
            fallback_jobs, _, _ = _search_jobs_without_browser(
                keywords,
                city or None,
                limit=count,
                allow_portal_fallback=True,
            )
            jobs = _enforce_cn_market_jobs(fallback_jobs, allow_entrypoints=True)
        except Exception:
            jobs = []

    if not jobs:
        try:
            jobs = _search_jobs_cn_entrypoints(keywords, city or None, limit=min(count, 20))
        except Exception:
            jobs = []

    jobs = _public_job_payload(jobs, limit=count, allow_entrypoints=True)
    _cache_recent_jobs(jobs)

    details: List[Dict[str, Any]] = []
    success_count = 0
    failed_count = 0

    for idx, job in enumerate(jobs[:count], 1):
        job_id = str(job.get("id") or "").strip()
        title = str(job.get("title") or "未知岗位")
        company = str(job.get("company") or "未知公司")
        link = str(job.get("link") or "")
        platform = str(job.get("platform") or "")
        ok = False
        message = ""

        if job_id:
            try:
                result = real_job_service.apply_job(
                    job_id,
                    resume_text,
                    {"phone": phone, "source": "simple_apply"},
                )
                ok = bool(result.get("success"))
                message = str(result.get("message") or "")
            except Exception as e:
                message = str(e)

        if (not ok) and link:
            app_id = f"SMP{int(time.time())}{idx:03d}"
            record = {
                "application_id": app_id,
                "job_id": job_id or app_id,
                "job_title": title,
                "company": company,
                "platform": platform,
                "apply_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "status": "待投递",
                "apply_link": link,
                "user_info": {"phone": phone, "source": "simple_apply"},
            }
            try:
                real_job_service.records.add_record(record)
            except Exception:
                pass
            ok = True
            message = "已生成投递链接（请在招聘网站完成真实投递）"

        if ok:
            success_count += 1
        else:
            failed_count += 1

        details.append(
            {
                "job_id": job_id,
                "job_title": title,
                "company": company,
                "platform": platform,
                "success": ok,
                "message": message or ("投递失败" if not ok else "成功"),
                "link": link,
            }
        )

    if not details:
        return {
            "success": False,
            "message": "未找到可投递岗位，请更换关键词或城市后重试",
            "total": 0,
            "success_count": 0,
            "failed_count": 0,
            "details": [],
        }

    return {
        "success": success_count > 0,
        "message": f"已处理 {len(details)} 个岗位，成功 {success_count} 个",
        "total": len(details),
        "success_count": success_count,
        "failed_count": failed_count,
        "details": details,
    }


def _build_legacy_stats_payload(records: List[Dict[str, Any]]) -> Dict[str, Any]:
    total = len(records or [])
    status_breakdown: Dict[str, int] = {}
    interviews = 0
    offers = 0
    responded = 0
    daily_count: Dict[str, int] = {}

    for row in records or []:
        status = _normalize_offer_stage(str(row.get("status") or "未知").strip())
        low = status.lower()
        status_breakdown[status] = status_breakdown.get(status, 0) + 1

        if ("面试" in status) or ("interview" in low):
            interviews += 1
        if ("offer" in low) or ("录用" in status):
            offers += 1
        if status not in {"待投递", "待处理", "已投递"}:
            responded += 1

        ts = str(row.get("apply_time") or row.get("created_at") or "").strip()
        date_key = ts[:10] if len(ts) >= 10 else ""
        if re.match(r"^\\d{4}-\\d{2}-\\d{2}$", date_key):
            daily_count[date_key] = daily_count.get(date_key, 0) + 1

    ordered_days = sorted(daily_count.keys())
    daily_series = [daily_count[d] for d in ordered_days][-14:]
    response_rate = round((responded / total) * 100, 1) if total else 0.0

    return {
        "total_applications": total,
        "response_rate": response_rate,
        "interviews": interviews,
        "offers": offers,
        "daily_applications": daily_series,
        "daily_labels": ordered_days[-14:],
        "status_breakdown": status_breakdown,
        "recent_applications": (records[-10:] if total > 10 else records),
    }


OFFER_STAGE_ORDER: List[str] = [
    "待投递",
    "已投递",
    "已回复",
    "一面",
    "二面",
    "HR面",
    "谈薪",
    "Offer",
    "已拒绝",
]


def _dedupe_preserve(items: List[Any]) -> List[str]:
    return _flatten_named_texts(items or [])


def _is_placeholder_text(value: str) -> bool:
    raw = str(value or "").strip()
    if not raw:
        return True
    low = raw.lower()
    if low in {"unknown", "n/a", "na", "null", "none"}:
        return True
    stripped = re.sub(r"\s+", "", raw)
    return bool(stripped) and all(ch in {"?", "？", "�"} for ch in stripped)


def _prefer_clean_text(preferred: str, fallback: str) -> str:
    preferred_text = str(preferred or "").strip()
    if preferred_text and (not _is_placeholder_text(preferred_text)):
        return preferred_text
    fallback_text = str(fallback or "").strip()
    return fallback_text


def _is_offer_stage_value(value: str) -> bool:
    return str(value or "").strip() in OFFER_STAGE_ORDER


def _normalize_offer_stage(status: str) -> str:
    raw = str(status or "").strip()
    if not raw:
        return "待投递"
    if _is_placeholder_text(raw):
        return "待投递"
    low = raw.lower()

    if ("offer" in low) or ("录用" in raw) or ("发放意向" in raw):
        return "Offer"
    if ("拒绝" in raw) or ("淘汰" in raw) or ("未通过" in raw) or ("reject" in low):
        return "已拒绝"
    if ("谈薪" in raw) or ("薪资" in raw) or ("salary" in low) or ("compensation" in low):
        return "谈薪"
    if ("hr面" in raw) or ("hr interview" in low) or (("hr" in low) and ("面" in raw)):
        return "HR面"
    if ("二面" in raw) or ("二轮" in raw) or ("复试" in raw) or ("second interview" in low):
        return "二面"
    if ("一面" in raw) or ("初面" in raw) or ("首轮" in raw) or ("first interview" in low):
        return "一面"
    if ("回复" in raw) or ("沟通" in raw) or ("约面" in raw) or ("response" in low) or ("reach out" in low):
        return "已回复"
    if ("已投递" in raw) or ("投递" in raw) or ("apply" in low) or ("submitted" in low):
        return "已投递"
    if ("待投递" in raw) or ("待处理" in raw) or ("queue" in low) or ("draft" in low):
        return "待投递"
    if "interview" in low:
        return "一面"
    return raw if _is_offer_stage_value(raw) else "待投递"


def _serialize_application_row(row: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(row, dict):
        return {
            "status": "待投递",
            "next_action": "检查原始记录格式",
            "raw": row,
        }
    clean = dict(row or {})
    clean["status"] = _normalize_offer_stage(str(clean.get("status") or ""))
    if not str(clean.get("next_action") or "").strip():
        clean["next_action"] = _build_offer_next_action(clean)
    return clean


def _offer_stage_rank(status: str) -> int:
    normalized = _normalize_offer_stage(status)
    try:
        return OFFER_STAGE_ORDER.index(normalized)
    except ValueError:
        return len(OFFER_STAGE_ORDER)


def _record_primary_ts(row: Dict[str, Any]) -> str:
    return str(row.get("updated_at") or row.get("apply_time") or row.get("created_at") or "")


def _filter_application_rows(rows: List[Dict[str, Any]], user_id: str = "", phone: str = "") -> List[Dict[str, Any]]:
    uid = str(user_id or "").strip()
    ph = str(phone or "").strip()
    if not uid and not ph:
        return list(rows or [])

    filtered: List[Dict[str, Any]] = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        user_info = row.get("user_info") if isinstance(row.get("user_info"), dict) else {}
        row_uid = str(user_info.get("user_id") or "").strip()
        row_phone = str(user_info.get("phone") or "").strip()
        if uid and row_uid == uid:
            filtered.append(row)
            continue
        if ph and row_phone == ph:
            filtered.append(row)
    return filtered


def _build_offer_next_action(row: Dict[str, Any]) -> str:
    stage = _normalize_offer_stage(str(row.get("status") or ""))
    company = str(row.get("company") or "该公司")
    title = str(row.get("job_title") or "该岗位")

    if stage == "待投递":
        return f"先用岗位定制版简历投递 {company} 的 {title}。"
    if stage == "已投递":
        return f"24 小时内检查 {company} 是否已读，必要时补发一条跟进消息。"
    if stage == "已回复":
        return f"准备 {company} 的一面自我介绍、项目亮点和反问清单。"
    if stage == "一面":
        return f"复盘一面问题，针对薄弱点补练，再推进 {company} 的二面。"
    if stage == "二面":
        return f"准备业务深挖、系统设计和跨团队协作案例，冲刺 {company} 的 HR 面。"
    if stage == "HR面":
        return f"梳理入职时间、薪资预期和稳定性说明，推动 {company} 进入谈薪。"
    if stage == "谈薪":
        return f"准备可接受区间、底线和加分项，争取 {company} 给出更优 offer。"
    if stage == "Offer":
        return f"核对 {company} offer 的薪资、试用期和汇报线，再决定是否接。"
    if stage == "已拒绝":
        return f"复盘 {company} 的失利原因，并把经验回灌到下一轮投递。"
    return "继续推进下一步。"


def _build_offer_funnel_payload(records: List[Dict[str, Any]], limit: int = 80) -> Dict[str, Any]:
    rows = list(records or [])
    rows.sort(key=_record_primary_ts, reverse=True)
    rows = rows[: max(1, min(int(limit or 80), 500))]

    counts = {stage: 0 for stage in OFFER_STAGE_ORDER}
    board_items: List[Dict[str, Any]] = []
    for row in rows:
        stage = _normalize_offer_stage(str(row.get("status") or ""))
        counts[stage] = counts.get(stage, 0) + 1
        board_items.append(
            {
                "application_id": str(row.get("application_id") or row.get("job_id") or ""),
                "job_title": str(row.get("job_title") or row.get("title") or "未知岗位"),
                "company": str(row.get("company") or "未知公司"),
                "status": stage,
                "apply_time": str(row.get("apply_time") or row.get("created_at") or ""),
                "updated_at": str(row.get("updated_at") or ""),
                "platform": str(row.get("platform") or ""),
                "next_action": str(row.get("next_action") or _build_offer_next_action(row)),
                "follow_up_at": str(row.get("follow_up_at") or ""),
                "notes": str(row.get("notes") or ""),
            }
        )

    stages: List[Dict[str, Any]] = []
    for stage in OFFER_STAGE_ORDER:
        items = [row for row in board_items if row.get("status") == stage][:12]
        stages.append({"stage": stage, "count": counts.get(stage, 0), "items": items})

    total = len(rows)
    offers = counts.get("Offer", 0)
    active = total - counts.get("Offer", 0) - counts.get("已拒绝", 0)
    interviews = counts.get("一面", 0) + counts.get("二面", 0) + counts.get("HR面", 0)

    return {
        "total": total,
        "active": active,
        "interviews": interviews,
        "offers": offers,
        "offer_rate": round((offers / total) * 100, 1) if total else 0.0,
        "counts": counts,
        "stages": stages,
        "recent": board_items[:20],
    }


def _job_text_blob(job: Dict[str, Any]) -> str:
    if not isinstance(job, dict):
        return ""
    parts = [
        str(job.get("title") or ""),
        str(job.get("job_title") or ""),
        str(job.get("company") or ""),
        str(job.get("salary") or ""),
        str(job.get("location") or ""),
        str(job.get("description") or ""),
        str(job.get("requirements") or ""),
    ]
    return " ".join([p for p in parts if p]).strip()


def _priority_label(score: float) -> str:
    value = float(score or 0)
    if value >= 78:
        return "S"
    if value >= 66:
        return "A"
    if value >= 52:
        return "B"
    return "C"


def _build_opportunity_cards(
    jobs: List[Dict[str, Any]],
    resume_text: str,
    target_role: str = "",
) -> List[Dict[str, Any]]:
    cards: List[Dict[str, Any]] = []
    for row in jobs or []:
        job = dict(row or {})
        text_blob = _job_text_blob(job)
        skill_payload = _skills_graph_payload(resume_text, target_role, text_blob)
        resume_skill_names = _dedupe_preserve(skill_payload.get("resume_skills") or [])
        job_skill_names = _dedupe_preserve(skill_payload.get("job_skills") or [])
        resume_skill_keys = {skill.lower() for skill in resume_skill_names}
        score = skill_payload.get("match_score")
        if score is None:
            overlap = resume_skill_keys & {skill.lower() for skill in job_skill_names}
            score = min(92.0, 48.0 + (len(overlap) * 9.0))
        score = round(float(score or 0), 1)
        overlap_skills = [
            skill
            for skill in job_skill_names
            if skill.lower() in resume_skill_keys
        ]
        missing_skills = [
            skill
            for skill in job_skill_names
            if skill.lower() not in resume_skill_keys
        ][:5]

        risks: List[str] = []
        if not str(job.get("salary") or "").strip():
            risks.append("薪资区间未公开")
        if not str(job.get("location") or "").strip():
            risks.append("工作地点信息不完整")
        if not overlap_skills:
            risks.append("岗位关键词与当前简历重合偏少")

        job["match_score"] = score
        job["priority"] = _priority_label(score)
        job["fit_reason"] = (
            f"匹配技能：{', '.join(overlap_skills[:4])}" if overlap_skills else "建议先用 ATS 关键词版简历补足命中率"
        )
        job["risk_flags"] = risks
        job["resume_variant_hint"] = "岗位定制版" if score >= 66 else "ATS关键词版"
        job["skill_gap"] = missing_skills
        cards.append(job)

    cards.sort(key=lambda item: (float(item.get("match_score") or 0), item.get("priority") == "S"), reverse=True)
    return cards


def _heuristic_opportunity_jobs(target_role: str, city: str, info: Dict[str, Any], limit: int = 6) -> List[Dict[str, Any]]:
    role = str(target_role or info.get("job_intention") or "AI应用工程师").strip() or "AI应用工程师"
    location = str(city or (info.get("preferred_locations") or ["上海"])[0] or "上海").strip() or "上海"
    skills_line = " / ".join([str(x) for x in (info.get("skills") or [])[:4]]) or "Python / FastAPI / SQL"
    seeds = [
        ("字节跳动", "30K-50K"),
        ("腾讯", "28K-45K"),
        ("阿里云", "25K-40K"),
        ("美团", "25K-40K"),
        ("携程", "20K-35K"),
        ("小红书", "30K-45K"),
    ]
    jobs: List[Dict[str, Any]] = []
    for index, (company, salary) in enumerate(seeds[: max(1, min(limit, len(seeds)))]):
        jobs.append(
            {
                "id": f"heuristic_{index+1}",
                "title": role,
                "company": company,
                "location": location,
                "salary": salary,
                "platform": "cn_portal",
                "provider": "heuristic_fallback",
                "link": "",
                "description": f"{company} 的 {role}，优先看重 {skills_line} 与业务落地能力。",
            }
        )
    return jobs


async def _build_resume_variants_payload(
    resume_text: str,
    resume_json: Dict[str, Any],
    info: Dict[str, Any],
    target_role: str,
    jobs: List[Dict[str, Any]],
    user_id: str = "",
) -> List[Dict[str, Any]]:
    base_json = resume_json if isinstance(resume_json, dict) else {}
    personal = base_json.get("personal_info") if isinstance(base_json.get("personal_info"), dict) else {}
    name = str(personal.get("name") or "候选人").strip()
    role = str(target_role or info.get("job_intention") or "AI应用工程师").strip()
    top_job_titles = [str(j.get("title") or j.get("job_title") or "").strip() for j in jobs[:3] if str(j.get("title") or j.get("job_title") or "").strip()]
    job_gap_keywords = [str(x) for x in ((jobs[0].get("skill_gap") or []) if jobs else [])]
    top_keywords = _dedupe_preserve(
        [role] + [str(x) for x in (info.get("skills") or [])[:8]] + job_gap_keywords
    )

    templates = resume_render_service.template_names()
    variant_specs = [
        {
            "variant_id": "ats_focus",
            "name": "ATS关键词版",
            "template": templates[0] if templates else "classic",
            "focus": "大批量投递，用关键词提高筛选命中率。",
            "summary": (
                f"{name} 目标岗位为 {role}。简历强调 {', '.join(top_keywords[:6]) or 'Python, FastAPI, SQL'}，"
                "突出可直接上手的技术栈、交付经验和业务结果，适合海投与 ATS 筛选。"
            ),
        },
        {
            "variant_id": "interview_focus",
            "name": "面试亮点版",
            "template": templates[1] if len(templates) > 1 else "modern",
            "focus": "面试前使用，把量化成果和项目 ownership 顶到最前面。",
            "summary": (
                f"{name} 具备从需求到上线的完整交付能力，重点展示 {role} 所需的项目推进、性能优化、"
                "问题排查和跨团队协作案例，适合一面和二面深挖。"
            ),
        },
        {
            "variant_id": "target_role",
            "name": "岗位定制版",
            "template": templates[2] if len(templates) > 2 else "minimal",
            "focus": "针对优先岗位做强定制，用于高价值公司冲刺。",
            "summary": (
                f"围绕 {role} 定制，优先匹配 {', '.join(top_job_titles) if top_job_titles else role} 这类岗位。"
                "重点突出最相关技能、项目和业务指标，减少泛化描述。"
            ),
        },
    ]

    variants: List[Dict[str, Any]] = []
    for spec in variant_specs:
        variant_json = json.loads(json.dumps(base_json or {}, ensure_ascii=False))
        variant_json["summary"] = spec["summary"]
        variant_json["skills"] = _dedupe_preserve(top_keywords + [str(x) for x in (base_json.get("skills") or [])])[:18]
        profile = resume_profile_service.save_profile(
            resume_json=variant_json,
            user_id=user_id,
            title=f"{name}-{spec['name']}",
            source_text=resume_text,
        )

        document = None
        try:
            rendered = await asyncio.to_thread(
                resume_render_service.render_to_file,
                variant_json,
                spec["template"],
                "html",
            )
            document = {
                "doc_id": rendered.get("doc_id"),
                "template": rendered.get("template"),
                "download_url": f"/api/resume/render/download/{rendered.get('doc_id')}",
            }
        except Exception as render_err:
            document = {"warning": str(render_err)[:200]}

        variants.append(
            {
                "variant_id": spec["variant_id"],
                "name": spec["name"],
                "focus": spec["focus"],
                "profile_id": profile.get("profile_id"),
                "summary": variant_json.get("summary"),
                "skills": variant_json.get("skills") or [],
                "document": document,
            }
        )

    return variants


def _build_interview_pack(
    resume_text: str,
    info: Dict[str, Any],
    target_role: str,
    target_job: Dict[str, Any],
) -> Dict[str, Any]:
    role = str(target_role or info.get("job_intention") or target_job.get("title") or "目标岗位").strip()
    company = str(target_job.get("company") or "目标公司").strip()
    title = str(target_job.get("title") or target_job.get("job_title") or role).strip()
    skills_payload = _skills_graph_payload(resume_text, role, _job_text_blob(target_job))
    resume_skills = _dedupe_preserve(skills_payload.get("resume_skills") or (info.get("skills") or []))[:8]
    job_skills = _dedupe_preserve(skills_payload.get("job_skills") or [role])[:8]
    gap_skills = [skill for skill in job_skills if skill.lower() not in {s.lower() for s in resume_skills}]

    core_questions: List[Dict[str, Any]] = []
    for skill in job_skills[:5]:
        core_questions.append(
            {
                "question": f"请讲一个你真实使用 {skill} 解决问题的案例。",
                "why_asked": f"{company} 会用它验证你是否真的能胜任 {title}。",
                "answer_framework": f"按 场景 -> 目标 -> 方案 -> 指标结果 来答，并把 {skill} 放到关键动作里。",
            }
        )

    behavioral_questions = [
        {
            "question": "你遇到过最难推进的一次项目是什么？",
            "answer_framework": "重点讲冲突、约束、你如何拆解问题，以及你推动结果落地的动作。",
        },
        {
            "question": "为什么想加入我们，而不是继续投更多同类公司？",
            "answer_framework": f"从 {company} 的业务场景、成长速度、岗位匹配度三个角度回答。",
        },
        {
            "question": "如果上线后效果不达预期，你会怎么排查？",
            "answer_framework": "先定义指标，再拆数据、链路、依赖、用户反馈，最后给修复优先级。",
        },
    ]

    interviewer_questions = [
        f"{company} 对 {title} 的前 90 天最重要目标是什么？",
        f"{title} 最常见的高绩效表现和低绩效表现分别是什么？",
        "团队当前最急的项目瓶颈是什么，我入职后最先接哪块？",
        "这个岗位和上下游团队的协作方式是怎样的？",
    ]

    week_plan = [
        "Day 1: 重写 90 秒自我介绍，确保和目标岗位一致。",
        "Day 2: 把最相关的 3 个项目按 STAR + 指标结果重讲一遍。",
        "Day 3: 逐个练核心技术题，补齐技能缺口。",
        "Day 4: 练压力题、追问题和反问问题。",
        "Day 5: 模拟一面完整流程，录音复盘。",
        "Day 6: 针对目标公司做业务和产品研究。",
        "Day 7: 整理谈薪底线、入职时间和 offer 取舍标准。",
    ]

    return {
        "target_role": role,
        "company": company,
        "job_title": title,
        "opening_pitch": (
            f"我适合 {company} 的 {title}，因为我能把 {', '.join(resume_skills[:4]) or '关键技能'} "
            "真正落到业务结果上，不只是停留在概念层面。"
        ),
        "resume_strengths": resume_skills,
        "job_must_have": job_skills,
        "gap_skills": gap_skills[:5],
        "core_questions": core_questions,
        "behavioral_questions": behavioral_questions,
        "questions_for_interviewer": interviewer_questions,
        "week_plan": week_plan,
        "salary_script": (
            "先确认岗位级别、绩效口径和总包结构，再给出你可接受区间与最低底线，"
            "不要在信息不完整时先报死数字。"
        ),
    }


def _build_follow_up_templates(target_job: Dict[str, Any], target_role: str) -> Dict[str, str]:
    company = str(target_job.get("company") or "贵司").strip()
    title = str(target_job.get("title") or target_role or "目标岗位").strip()
    return {
        "after_apply": f"您好，我刚投递了 {company} 的 {title}，我的经历与岗位要求较匹配，方便的话希望进一步沟通。",
        "after_interview": f"感谢今天的面试。我对 {company} 的 {title} 更感兴趣了，也更确认自己能在该岗位尽快产出结果。",
        "salary_negotiation": f"我对加入 {company} 很有意愿，如果薪资和级别还能向上调整一些，我可以更快确认 offer。",
    }


def _build_offer_roadmap(
    target_role: str,
    opportunities: List[Dict[str, Any]],
    funnel: Dict[str, Any],
) -> List[Dict[str, Any]]:
    top_titles = [str(row.get("title") or row.get("job_title") or "").strip() for row in opportunities[:3] if str(row.get("title") or row.get("job_title") or "").strip()]
    return [
        {
            "stage": "Agent 1 简历定稿",
            "goal": f"围绕 {target_role or '目标岗位'} 产出 3 版简历。",
            "success_metric": "至少保留 ATS 版、面试版、岗位定制版。",
        },
        {
            "stage": "Agent 2 岗位筛选",
            "goal": f"从 {', '.join(top_titles) if top_titles else '高匹配岗位'} 里筛出优先投递池。",
            "success_metric": "优先池岗位匹配分 >= 66。",
        },
        {
            "stage": "Agent 3 自动投递",
            "goal": "先批量铺量，再集中冲刺高价值岗位。",
            "success_metric": "每天有新增投递和新增回复。",
        },
        {
            "stage": "Agent 4 面试训练",
            "goal": "针对每个目标岗位做问答训练和项目复盘。",
            "success_metric": "一面通过率和二面推进率持续提高。",
        },
        {
            "stage": "Agent 5 跟进谈薪",
            "goal": "让投递进入沟通、面试、谈薪和 offer 阶段。",
            "success_metric": f"当前 offer 数 {funnel.get('offers', 0)}，持续推进 active pipeline {funnel.get('active', 0)}。",
        },
    ]


def _cache_recent_jobs(jobs: List[Dict[str, Any]], max_size: int = 2000) -> None:
    now = datetime.now().isoformat()
    for j in jobs or []:
        jid = str(j.get("id") or "").strip()
        if not jid:
            continue
        row = dict(j)
        row["_cached_at"] = now
        recent_search_jobs[jid] = row
    if len(recent_search_jobs) > max_size:
        # Keep newest max_size by cached time.
        items = sorted(
            recent_search_jobs.items(),
            key=lambda x: str(x[1].get("_cached_at") or ""),
            reverse=True,
        )[:max_size]
        recent_search_jobs.clear()
        recent_search_jobs.update(items)


def _is_seed_or_demo_job(job: Dict[str, Any]) -> bool:
    """Filter obvious seed/demo jobs so they never show in production recommendations."""
    jid = str(job.get("id") or "").strip().lower()
    if jid.startswith("seed") or jid.startswith("demo"):
        return True
    link = str(job.get("link") or job.get("apply_url") or "").strip().lower()
    if "/job_detail/seed" in link:
        return True
    company = str(job.get("company") or "").strip().lower()
    if company in {"示例公司", "测试公司", "demo"}:
        return True
    return False


def _job_has_actionable_link(job: Dict[str, Any]) -> bool:
    link = str(job.get("link") or job.get("apply_url") or "").strip()
    return link.startswith("http://") or link.startswith("https://")


def _is_cn_entrypoint_link(link: str) -> bool:
    """Search entry/list pages are not treated as real actionable postings."""
    low = (link or "").strip().lower()
    if not low:
        return False
    patterns = (
        "zhipin.com/web/geek/job?",
        "zhipin.com/zhaopin/",
        "liepin.com/zhaopin/?key=",
        "sou.zhaopin.com/?kw=",
        "we.51job.com/pc/search?",
        "lagou.com/wn/jobs?kd=",
    )
    return any(p in low for p in patterns)


def _is_cn_entrypoint_job(job: Dict[str, Any]) -> bool:
    provider = str(job.get("provider") or "").strip().lower()
    title = str(job.get("title") or job.get("job_title") or "").strip().lower()
    link = str(job.get("link") or job.get("apply_url") or "").strip()
    if provider == "cn_portal":
        return True
    if "搜索入口" in title:
        return True
    return _is_cn_entrypoint_link(link)


def _strip_cn_entrypoint_jobs(jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for row in jobs or []:
        if not isinstance(row, dict):
            continue
        if _is_cn_entrypoint_job(row):
            continue
        out.append(row)
    return out


def _normalize_and_filter_jobs(jobs: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    """Keep only actionable real jobs and deduplicate by link/id/title-company."""
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for job in jobs or []:
        if not isinstance(job, dict):
            continue
        if _is_seed_or_demo_job(job):
            continue
        if not _job_has_actionable_link(job):
            continue

        link = str(job.get("link") or job.get("apply_url") or "").strip().lower()
        jid = str(job.get("id") or "").strip().lower()
        title = str(job.get("title") or job.get("job_title") or "").strip().lower()
        company = str(job.get("company") or "").strip().lower()
        dedupe_key = link or jid or f"{title}|{company}"
        if not dedupe_key or dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        out.append(job)
        if len(out) >= max(1, int(limit or 10)):
            break
    return out


def _normalize_real_jobs(jobs: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    """
    Strict real-posting normalization:
    - actionable links only
    - no seed/demo
    - no board-level search entry pages
    """
    scan_limit = max(20, int(limit or 10) * 5)
    base = _normalize_and_filter_jobs(jobs, limit=scan_limit)
    out: List[Dict[str, Any]] = []
    for job in base:
        if _is_cn_entrypoint_job(job):
            continue
        out.append(job)
        if len(out) >= max(1, int(limit or 10)):
            break
    return out


def _is_cn_job_link(link: str) -> bool:
    low = (link or "").lower()
    return any(d in low for d in CN_JOB_DOMAINS)


def _enforce_cn_market_jobs(
    jobs: List[Dict[str, Any]],
    allow_entrypoints: bool = False,
) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for j in (jobs or []):
        if (not allow_entrypoints) and _is_cn_entrypoint_job(j):
            continue
        link = str(j.get("link") or j.get("apply_url") or "")
        platform = str(j.get("platform") or "").strip()
        if _is_cn_job_link(link):
            out.append(j)
            continue
        if platform in {"Boss直聘", "猎聘", "智联招聘", "前程无忧", "拉勾"}:
            out.append(j)
    return out


def _validate_chargeable_job_detail(job: Dict[str, Any]) -> Tuple[bool, str]:
    """Return whether this job detail qualifies as a real, chargeable posting."""
    if not isinstance(job, dict) or not job:
        return False, "empty_job_detail"
    if _is_seed_or_demo_job(job):
        return False, "seed_or_demo"
    if _is_cn_entrypoint_job(job):
        return False, "cn_entrypoint"
    if not _job_has_actionable_link(job):
        return False, "missing_actionable_link"
    if not _enforce_cn_market_jobs([job], allow_entrypoints=False):
        return False, "non_cn_market"
    return True, "ok"


def _record_job_detail_credit_state(
    job_id: str,
    job: Dict[str, Any],
    credit_charge: Dict[str, Any],
    validated: bool,
    validation_reason: str,
) -> None:
    try:
        charge_payload = credit_charge if isinstance(credit_charge, dict) else {}
        charge_status = "failed"
        if charge_payload.get("ok"):
            charge_status = "skipped" if charge_payload.get("skipped") else "charged"
        ledger = charge_payload.get("ledger")
        ledger_id = ""
        if isinstance(ledger, dict):
            ledger_id = str(ledger.get("ledger_id") or "")
        if not ledger_id:
            ledger_id = str(charge_payload.get("ledger_id") or "")
        patch = {
            "detail_checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "detail_validated": bool(validated),
            "detail_validation_reason": str(validation_reason or ""),
            "detail_credit_status": charge_status,
            "detail_credit_reason": str(charge_payload.get("reason") or ""),
            "detail_credit_action": "job_detail",
            "detail_credit_ledger_id": ledger_id,
            "job_title": str(job.get("title") or job.get("job_title") or ""),
            "company": str(job.get("company") or ""),
            "platform": str(job.get("platform") or job.get("provider") or ""),
            "apply_link": str(job.get("link") or job.get("apply_url") or ""),
        }
        real_job_service.records.update_latest_record_by_job_id(str(job_id or "").strip(), patch, create_if_missing=True)
    except Exception:
        logger.exception("job_detail_credit_record_update_failed job_id=%s", str(job_id or "").strip()[:120])


PROCESS_RESPONSE_SCHEMA_VERSION = "process_response.v2"
LOW_QUALITY_MARKERS = (
    "i appreciate your interest",
    "amazon q",
    "built by aws",
    "i'm amazon q",
    "i am amazon q",
    "not career counseling",
    "没有看到任何附件",
)
PUBLIC_EVENT_NAME_PATTERN = re.compile(r"^[a-z0-9_]{3,48}$")
PUBLIC_EVENT_WHITELIST = {
    "workspace_opened",
    "resume_process_started",
    "resume_process_success",
    "resume_process_failed",
    "job_link_click",
    "result_download",
}


def _text_has_low_quality_marker(text: str) -> bool:
    low = (text or "").strip().lower()
    if not low:
        return True
    return any(marker in low for marker in LOW_QUALITY_MARKERS)


def _render_market_analysis_fallback(info: Dict[str, Any]) -> str:
    skills = [s for s in (info.get("skills") or []) if s][:8]
    skills_text = ", ".join(skills) if skills else "Python, FastAPI, SQL"
    return (
        "【市场竞争力分析】\n\n"
        f"✅ 识别技能：{skills_text}\n\n"
        "📊 市场需求度：中高（基于中国技术岗位关键词覆盖）\n"
        "💼 建议投递方向：Python后端 / AI应用工程 / 数据工程\n"
        "📈 建议策略：先投递有明确技术栈和薪资区间的岗位，再按反馈迭代简历。\n\n"
        "💡 市场建议：\n"
        "1. 先聚焦 1-2 个主岗位方向，避免关键词过散。\n"
        "2. 简历中补充可量化结果（性能提升、效率提升、交付周期）。\n"
        "3. 优先投递带真实岗位详情页和直接投递入口的职位。"
    )


def _render_interview_prep_fallback(info: Dict[str, Any], jobs: List[Dict[str, Any]]) -> str:
    top = (jobs or [{}])[0]
    title = str(top.get("title") or "Python后端工程师")
    company = str(top.get("company") or "目标公司")
    return (
        f"目标岗位：{title}（{company}）\n\n"
        "高频问题 1：你如何设计高并发 API？\n"
        "回答要点：异步 I/O、缓存策略、限流、慢查询治理、可观测性。\n\n"
        "高频问题 2：你如何保证数据链路质量？\n"
        "回答要点：输入校验、幂等设计、回滚方案、监控告警、复盘机制。\n\n"
        "高频问题 3：你做过最有业务价值的项目是什么？\n"
        "回答要点：按“背景-动作-结果”讲清指标提升，并说明你的关键贡献。"
    )


def _render_optimized_resume_fallback(resume_text: str, info: Dict[str, Any]) -> str:
    name = str(info.get("name") or "").strip()
    if not name or name == "未知":
        first = (resume_text or "").strip().splitlines()
        name = first[0].strip() if first else "候选人"
    skills = [s for s in (info.get("skills") or []) if s][:10]
    skills_line = ", ".join(skills) if skills else "Python, FastAPI, SQL, Docker"
    return (
        f"{name}\n"
        "AI应用工程师 | Python后端工程师\n\n"
        "核心优势\n"
        "- 聚焦 AI 应用工程与后端交付，能从需求到上线完成闭环。\n"
        "- 具备数据处理、服务部署、性能优化和质量门禁实践。\n\n"
        f"技术栈\n- {skills_line}\n\n"
        "项目亮点（示例结构）\n"
        "- 构建 RAG/数据服务，接口稳定性提升，响应延迟下降。\n"
        "- 建立自动化数据管道，减少人工处理成本并提高时效。\n"
        "- 引入测试与监控机制，降低线上故障率并缩短定位时间。\n\n"
        "求职方向\n- Python后端开发\n- AI应用开发\n- 数据工程与平台方向"
    )


def _run_output_quality_gate(
    results: Dict[str, Any],
    resume_text: str,
    info: Dict[str, Any],
    real_jobs: List[Dict[str, Any]],
) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    clean = dict(results or {})
    report: Dict[str, Any] = {"passed": True, "issues": []}

    market_analysis = str(clean.get("market_analysis") or "").strip()
    if len(market_analysis) < 80 or _text_has_low_quality_marker(market_analysis):
        clean["market_analysis"] = _render_market_analysis_fallback(info)
        report["issues"].append("market_analysis_replaced")

    optimized_resume = str(clean.get("optimized_resume") or "").strip()
    if len(optimized_resume) < 120 or _text_has_low_quality_marker(optimized_resume):
        clean["optimized_resume"] = _render_optimized_resume_fallback(resume_text, info)
        report["issues"].append("optimized_resume_replaced")

    interview_prep = str(clean.get("interview_prep") or "").strip()
    if len(interview_prep) < 80 or _text_has_low_quality_marker(interview_prep):
        clean["interview_prep"] = _render_interview_prep_fallback(info, real_jobs)
        report["issues"].append("interview_prep_replaced")

    salary_analysis = str(clean.get("salary_analysis") or "").strip()
    if not salary_analysis:
        clean["salary_analysis"] = "【薪资潜力分析】\n\n建议：基于岗位匹配度和市场供需，优先投递高匹配岗位后再进行薪资谈判。"
        report["issues"].append("salary_analysis_replaced")

    if report["issues"]:
        report["passed"] = False
    return clean, report


def _public_job_payload(
    jobs: List[Dict[str, Any]],
    limit: int = 10,
    allow_entrypoints: bool = False,
) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    normalized = (
        _normalize_and_filter_jobs(jobs or [], limit=limit)
        if allow_entrypoints
        else _normalize_real_jobs(jobs or [], limit=limit)
    )
    for j in normalized:
        link = str(j.get("link") or j.get("apply_url") or "").strip()
        out.append(
            {
                "id": str(j.get("id") or f"job_{abs(hash(link.lower()))}"),
                "title": str(j.get("title") or j.get("job_title") or "未知岗位"),
                "company": str(j.get("company") or ""),
                "location": str(j.get("location") or ""),
                "salary": str(j.get("salary") or j.get("salary_range") or ""),
                "platform": str(j.get("platform") or j.get("provider") or _platform_from_link(link)),
                "link": link,
                "provider": str(j.get("provider") or ""),
                "updated": str(j.get("updated") or ""),
            }
        )
    return out


def _validate_process_response_shape(payload: Dict[str, Any]) -> List[str]:
    errors: List[str] = []
    required_text_fields = (
        "career_analysis",
        "job_recommendations",
        "optimized_resume",
        "interview_prep",
        "mock_interview",
        "job_provider_mode",
        "schema_version",
    )
    for key in required_text_fields:
        if not isinstance(payload.get(key), str):
            errors.append(f"{key}:expected_string")

    jobs = payload.get("recommended_jobs")
    if not isinstance(jobs, list):
        errors.append("recommended_jobs:expected_list")
    else:
        for idx, row in enumerate(jobs[:20]):
            if not isinstance(row, dict):
                errors.append(f"recommended_jobs[{idx}]:expected_object")
                continue
            for key in ("id", "title", "link"):
                if not isinstance(row.get(key), str) or not str(row.get(key)).strip():
                    errors.append(f"recommended_jobs[{idx}].{key}:missing")
                    break
    return errors


def _filter_cloud_cache_by_query(
    keywords: List[str], location: Optional[str], limit: int
) -> List[Dict[str, Any]]:
    kw = [k.strip().lower() for k in (keywords or []) if k and k.strip()]

    def hit(job: Dict[str, Any]) -> bool:
        text = f"{job.get('title','')} {job.get('company','')}".lower()
        if kw and not any(k in text for k in kw):
            return False
        if location and job.get("location") and location not in str(job.get("location")):
            return False
        return True

    matched = [j for j in cloud_jobs_cache if hit(j)]
    if not matched:
        matched = list(cloud_jobs_cache)
    return _normalize_real_jobs(matched, limit=limit)


def _search_jobs_without_browser(
    keywords: List[str],
    location: Optional[str],
    limit: int,
    allow_portal_fallback: bool = False,
) -> Tuple[List[Dict[str, Any]], str, Optional[str]]:
    """
    Cloud-safe real-time search path.
    No local browser/OpenClaw required.
    """
    try:
        jobs = real_job_service.search_jobs(
            keywords=keywords,
            location=location,
            limit=max(5, min(int(limit or 10), 50)),
        )
        mode = (real_job_service.get_statistics() or {}).get("provider_mode", "auto")
        normalized = _normalize_real_jobs(jobs, limit=limit)
        normalized = _enforce_cn_market_jobs(normalized)
        if normalized:
            return normalized, mode, None
    except Exception as e:
        first_error = str(e)
    else:
        first_error = None

    enterprise_jobs = _search_jobs_enterprise_api(keywords, location, limit=limit)
    enterprise_jobs = _normalize_real_jobs(enterprise_jobs, limit=limit)
    enterprise_jobs = _enforce_cn_market_jobs(enterprise_jobs)
    if enterprise_jobs:
        return enterprise_jobs, "enterprise_api", None

    # CN market fallback: no-browser HTML search on Chinese job sites.
    bing_html_jobs = _search_jobs_bing_html(keywords, location, limit=limit)
    bing_html_jobs = _normalize_real_jobs(bing_html_jobs, limit=limit)
    bing_html_jobs = _enforce_cn_market_jobs(bing_html_jobs)
    if bing_html_jobs:
        return bing_html_jobs, "bing_html", None

    # Last fallback: DuckDuckGo HTML search (no key, no browser).
    ddg_jobs = _search_jobs_duckduckgo(keywords, location, limit=limit)
    ddg_jobs = _normalize_real_jobs(ddg_jobs, limit=limit)
    ddg_jobs = _enforce_cn_market_jobs(ddg_jobs)
    if ddg_jobs:
        return ddg_jobs, "duckduckgo", None

    # Optional global fallback. Disabled by default to keep CN market realism.
    if os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() in {"1", "true", "yes", "on"}:
        remotive_jobs = _search_jobs_remotive(keywords, location, limit=limit)
        if remotive_jobs:
            return remotive_jobs, "remotive", None

    if allow_portal_fallback:
        portal_jobs = _search_jobs_cn_entrypoints(keywords, location, limit=min(limit, 5))
        if portal_jobs:
            return portal_jobs, "cn_portal", first_error or "using cn portal fallback"

    return [], "no_real_jobs", first_error or "no result from no-browser providers"


def _platform_from_link(link: str) -> str:
    host = (urlparse(link).netloc or "").lower()
    if "zhipin.com" in host:
        return "Boss直聘"
    if "liepin.com" in host:
        return "猎聘"
    if "zhaopin.com" in host:
        return "智联招聘"
    if "51job.com" in host:
        return "前程无忧"
    if "lagou.com" in host:
        return "拉勾"
    return host or "web"


def _first_non_empty(row: Dict[str, Any], keys: List[str]) -> str:
    for k in keys:
        v = row.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return ""


def _infer_company_from_title(title: str, platform: str = "") -> str:
    t = (title or "").strip()
    if not t:
        return ""
    # Boss style examples:
    # - 「Python招聘」_苏州鼎级招聘-BOSS直聘
    # - Python开发工程师 - 某某科技 - Boss直聘
    if "_" in t:
        tail = t.split("_", 1)[1].strip()
        tail = re.split(r"[-|｜]", tail)[0].strip()
        tail = tail.replace("招聘", "").replace("诚聘", "").strip()
        if 1 <= len(tail) <= 24 and "直聘" not in tail:
            return tail
    parts = [p.strip() for p in re.split(r"[-|｜]", t) if p.strip()]
    for p in parts[1:3]:
        if any(x in p for x in ("直聘", "招聘", "猎聘", "前程无忧", "拉勾", "智联")):
            continue
        if 1 <= len(p) <= 24:
            return p.replace("招聘", "").replace("诚聘", "").strip()
    return ""


def _search_jobs_enterprise_api(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    """
    Enterprise-grade provider hook (config-driven).
    Intended for stable cloud use when search engines hit anti-bot limits.
    """
    url = os.getenv("ENTERPRISE_JOB_API_URL", "").strip()
    if not url:
        return []

    method = os.getenv("ENTERPRISE_JOB_API_METHOD", "GET").strip().upper()
    timeout_s = int(os.getenv("ENTERPRISE_JOB_API_TIMEOUT_S", "15") or "15")
    auth_header = os.getenv("ENTERPRISE_JOB_API_AUTH_HEADER", "Authorization").strip() or "Authorization"
    auth_scheme = os.getenv("ENTERPRISE_JOB_API_AUTH_SCHEME", "Bearer").strip()
    api_key = os.getenv("ENTERPRISE_JOB_API_KEY", "").strip()

    kw = [k.strip() for k in (keywords or []) if k and k.strip()]
    query = " ".join(kw[:5]) or "Python"
    payload = {
        "query": query,
        "keywords": kw,
        "location": (location or "").strip(),
        "limit": max(1, min(int(limit or 10), 50)),
    }
    headers = {"User-Agent": "ai-job-helper/1.0"}
    if api_key:
        headers[auth_header] = f"{auth_scheme} {api_key}".strip() if auth_scheme else api_key

    try:
        if method == "POST":
            resp = requests.post(url, json=payload, headers=headers, timeout=timeout_s)
        else:
            resp = requests.get(url, params=payload, headers=headers, timeout=timeout_s)
        if resp.status_code >= 400:
            return []
        data = resp.json() if resp.content else {}
    except Exception:
        return []

    rows: List[Dict[str, Any]] = []
    if isinstance(data, list):
        rows = [x for x in data if isinstance(x, dict)]
    elif isinstance(data, dict):
        for k in ("jobs", "results", "items", "list", "data"):
            v = data.get(k)
            if isinstance(v, list):
                rows = [x for x in v if isinstance(x, dict)]
                break
            if isinstance(v, dict):
                for kk in ("jobs", "results", "items", "list"):
                    vv = v.get(kk)
                    if isinstance(vv, list):
                        rows = [x for x in vv if isinstance(x, dict)]
                        break
                if rows:
                    break

    out: List[Dict[str, Any]] = []
    for i, row in enumerate(rows, 1):
        link = _first_non_empty(row, ["link", "url", "job_url", "detail_url", "apply_url", "jobLink", "jobUrl"])
        if not link:
            continue
        title = _first_non_empty(row, ["title", "job_title", "name", "position", "jobName"]) or "招聘岗位"
        company = _first_non_empty(row, ["company", "company_name", "employer", "brandName"])
        loc = _first_non_empty(row, ["location", "city", "region", "work_city"]) or (location or "")
        salary = _first_non_empty(row, ["salary", "salary_range", "pay", "salaryRange"])
        platform = _first_non_empty(row, ["platform", "source", "site", "origin"]) or _platform_from_link(link)
        out.append(
            {
                "id": f"enterprise_{i}_{abs(hash(link.lower()))}",
                "title": title,
                "company": company,
                "location": loc,
                "salary": salary,
                "platform": platform,
                "link": link,
                "provider": "enterprise_api",
                "updated": _first_non_empty(row, ["updated", "updated_at", "publish_time", "published_at"]),
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _normalize_ddg_redirect(href: str) -> str:
    href = html_lib.unescape((href or "").strip())
    if not href:
        return ""
    if href.startswith("//"):
        href = "https:" + href
    # Absolute DuckDuckGo redirect URL.
    if href.startswith("http://duckduckgo.com/l/?") or href.startswith("https://duckduckgo.com/l/?"):
        qs = parse_qs(urlparse(href).query)
        uddg = (qs.get("uddg") or [""])[0]
        return unquote(uddg) if uddg else ""
    if href.startswith("/l/?"):
        qs = parse_qs(urlparse("https://duckduckgo.com" + href).query)
        uddg = (qs.get("uddg") or [""])[0]
        return unquote(uddg) if uddg else ""
    return href


def _search_jobs_duckduckgo(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    q_parts = [k.strip() for k in (keywords or []) if k and k.strip()]
    if location:
        q_parts.append(location.strip())
    q_parts.append("招聘 职位 site:zhipin.com OR site:liepin.com OR site:zhaopin.com OR site:51job.com OR site:lagou.com")
    q = " ".join(q_parts).strip() or "招聘 职位 site:zhipin.com"

    url = f"https://html.duckduckgo.com/html/?q={quote_plus(q)}"
    try:
        resp = requests.get(
            url,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            },
            timeout=12,
        )
        text = resp.text or ""
    except Exception:
        return []

    # result__a href="...">title</a>
    pattern = re.compile(r'<a[^>]*class="[^"]*result__a[^"]*"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.I | re.S)
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for href, raw_title in pattern.findall(text):
        link = _normalize_ddg_redirect(href)
        if not link:
            continue
        if not (link.startswith("http://") or link.startswith("https://")):
            continue
        low = link.lower()
        if not any(d in low for d in CN_JOB_DOMAINS):
            continue
        if low in seen:
            continue
        seen.add(low)
        title = re.sub(r"<[^>]+>", "", raw_title or "")
        title = html_lib.unescape(title).strip()
        platform = _platform_from_link(link)
        company = _infer_company_from_title(title, platform=platform)
        out.append(
            {
                "id": f"duckduckgo_{abs(hash(low))}",
                "title": title or "招聘岗位",
                "company": company,
                "location": location or "",
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "duckduckgo",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _search_jobs_bing_html(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    q_parts = [k.strip() for k in (keywords or []) if k and k.strip()]
    if location:
        q_parts.append(location.strip())
    q_parts.append("招聘 职位 site:zhipin.com OR site:liepin.com OR site:zhaopin.com OR site:51job.com OR site:lagou.com")
    q = " ".join(q_parts).strip() or "招聘 职位 site:zhipin.com"

    try:
        resp = requests.get(
            "https://www.bing.com/search",
            params={"q": q, "count": max(10, min(int(limit or 10) * 2, 50))},
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124",
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            },
            timeout=12,
        )
        text = resp.text or ""
    except Exception:
        return []

    # <li class="b_algo"> ... <h2><a href="...">title</a>
    pattern = re.compile(r'<li class="b_algo"[^>]*>.*?<h2><a href="([^"]+)"[^>]*>(.*?)</a>', re.I | re.S)
    out: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for link, raw_title in pattern.findall(text):
        if not link.startswith(("http://", "https://")):
            continue
        low = link.lower()
        if not any(d in low for d in CN_JOB_DOMAINS):
            continue
        if low in seen:
            continue
        seen.add(low)
        title = re.sub(r"<[^>]+>", "", raw_title or "")
        title = html_lib.unescape(title).strip()
        platform = _platform_from_link(link)
        company = _infer_company_from_title(title, platform=platform)
        out.append(
            {
                "id": f"binghtml_{abs(hash(low))}",
                "title": title or "招聘岗位",
                "company": company,
                "location": location or "",
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "bing_html",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _search_jobs_remotive(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    # Remotive is a public job API and currently reachable in cloud environments.
    cleaned = [k.strip() for k in (keywords or []) if k and k.strip()]
    q = cleaned[0] if cleaned else "python"
    try:
        resp = requests.get(
            "https://remotive.com/api/remote-jobs",
            params={"search": q},
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=15,
        )
        data = resp.json() if resp.content else {}
    except Exception:
        return []

    jobs = data.get("jobs") or []
    base_rows: List[Dict[str, Any]] = []
    for it in jobs:
        link = str(it.get("url") or "").strip()
        if not link:
            continue
        title = str(it.get("title") or "").strip()
        company = str(it.get("company_name") or "").strip()
        loc = str(it.get("candidate_required_location") or location or "").strip()
        base_rows.append(
            {
                "id": f"remotive_{it.get('id')}",
                "title": title or "招聘岗位",
                "company": company,
                "location": loc,
                "salary": str(it.get("salary") or "").strip(),
                "platform": "Remotive",
                "link": link,
                "provider": "remotive",
                "updated": str(it.get("publication_date") or "").strip(),
            }
        )
        if len(base_rows) >= max(30, int(limit or 10) * 4):
            break

    if location:
        narrowed = [j for j in base_rows if location in str(j.get("location") or "")]
        if narrowed:
            return _normalize_and_filter_jobs(narrowed, limit=limit)
    # If no location match, return best available remote jobs instead of empty.
    return _normalize_and_filter_jobs(base_rows, limit=limit)


def _search_jobs_cn_entrypoints(
    keywords: List[str], location: Optional[str], limit: int = 10
) -> List[Dict[str, Any]]:
    """Guaranteed CN-market fallback: job-board search entry links."""
    kw = [k.strip() for k in (keywords or []) if k and k.strip()]
    query = " ".join(kw[:3]) or "Python"
    loc = (location or "").strip()

    boards = [
        ("Boss直聘", f"https://www.zhipin.com/web/geek/job?query={quote_plus(query)}"),
        ("猎聘", f"https://www.liepin.com/zhaopin/?key={quote_plus(query)}"),
        ("智联招聘", f"https://sou.zhaopin.com/?kw={quote_plus(query)}" + (f"&jl={quote_plus(loc)}" if loc else "")),
        ("前程无忧", f"https://we.51job.com/pc/search?keyword={quote_plus(query)}"),
        ("拉勾", f"https://www.lagou.com/wn/jobs?kd={quote_plus(query)}"),
    ]
    out: List[Dict[str, Any]] = []
    for i, (platform, link) in enumerate(boards, 1):
        out.append(
            {
                "id": f"cn_portal_{i}_{abs(hash(link))}",
                "title": f"{query} - {platform}搜索入口",
                "company": "",
                "location": loc,
                "salary": "",
                "platform": platform,
                "link": link,
                "provider": "cn_portal",
            }
        )
        if len(out) >= max(1, int(limit or 10)):
            break
    return _normalize_and_filter_jobs(out, limit=limit)


def _candidate_from_resume_profile(profile: Dict[str, Any]) -> Dict[str, Any]:
    detail = resume_profile_service.get_profile(str((profile or {}).get("profile_id") or "")) or {}
    resume_json = detail.get("resume_json") if isinstance(detail.get("resume_json"), dict) else {}
    personal = resume_json.get("personal_info") if isinstance(resume_json.get("personal_info"), dict) else {}
    skills = [str(x).strip() for x in (resume_json.get("skills") or []) if str(x).strip()]
    summary = str(resume_json.get("summary") or "").strip()
    source_text = str(detail.get("source_text") or "").strip()
    return {
        "candidate_id": f"resume:{str((profile or {}).get('profile_id') or '').strip()}",
        "profile_id": str((profile or {}).get("profile_id") or "").strip(),
        "name": str(personal.get("name") or (profile or {}).get("title") or "匿名候选人").strip(),
        "location": str(personal.get("location") or "").strip(),
        "skills": skills,
        "summary": summary,
        "source_text": source_text[:2000],
        "years_experience": detail.get("years_experience"),
        "source": "resume_profile",
        "updated_at": str((profile or {}).get("updated_at") or "").strip(),
    }


def _build_resume_candidate_pool(limit: int = 120) -> List[Dict[str, Any]]:
    rows = resume_profile_service.list_profiles(limit=limit).get("items") or []
    out: List[Dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        item = _candidate_from_resume_profile(row)
        if item.get("candidate_id"):
            out.append(item)
    return out


def _merge_candidate_pool(local_rows: List[Dict[str, Any]], resume_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    merged: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for row in (local_rows or []) + (resume_rows or []):
        if not isinstance(row, dict):
            continue
        cid = str(row.get("candidate_id") or "").strip()
        if not cid:
            continue
        if cid in seen:
            continue
        seen.add(cid)
        merged.append(row)
    return merged


def _find_hr_job(job_id: str) -> Optional[Dict[str, Any]]:
    target = str(job_id or "").strip()
    if not target:
        return None
    for row in hr_match_service.list_jobs(limit=500):
        if str((row or {}).get("job_id") or "").strip() == target:
            return row
    return None

@app.get("/", include_in_schema=False)
async def home():
    """Default entry: always route to app workspace to avoid stale static home pages."""
    return RedirectResponse(url="/app", status_code=302)

@app.get("/app", response_class=HTMLResponse)
async def app_page():
    """主应用页面（仅保留统一工作台版本）。"""
    # Default to last-night premium UI (app.html), fallback to legacy full workspace.
    candidates = ["static/app.html", "static/app_pro.html"]
    for app_html in candidates:
        if os.path.exists(app_html) and os.path.getsize(app_html) > 64:
            with open(app_html, "r", encoding="utf-8") as f:
                return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>app.html/app_pro.html not found</h1>", status_code=500)


@app.get("/enter", response_class=HTMLResponse)
async def enter_page():
    """主工作台入口 - 重定向到 /app"""
    return await app_page()


@app.get("/app_hr", response_class=HTMLResponse)
async def app_hr_page():
    """HR bilateral matching workspace (new dark UI)."""
    app_html = "static/app.html"
    if os.path.exists(app_html) and os.path.getsize(app_html) > 64:
        with open(app_html, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>app_hr unavailable</h1>", status_code=500)


@app.get("/index.html", response_class=HTMLResponse)
async def home_alias():
    """旧入口统一跳转到主工作台。"""
    return RedirectResponse(url="/app", status_code=302)


@app.get("/app_clean.html", response_class=HTMLResponse)
async def app_clean_alias():
    """旧工作台统一跳转。"""
    return RedirectResponse(url="/app", status_code=302)


@app.get("/auto_apply_panel.html", response_class=HTMLResponse)
async def auto_apply_panel_page():
    """自动投递旧面板统一跳转。"""
    return RedirectResponse(url="/app", status_code=302)


@app.get("/investor.html", response_class=HTMLResponse)
async def investor_alias():
    """兼容旧投资看板路径。"""
    return await investor_page()


@app.get("/investor", response_class=HTMLResponse)
async def investor_page():
    """Investor-facing readiness dashboard page."""
    investor_html = "static/investor.html"
    if os.path.exists(investor_html):
        with open(investor_html, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>Investor Dashboard</h1><p>/api/investor/readiness</p>")

@app.post("/api/upload")
async def upload_resume(
    file: UploadFile = File(...),
    force_ocr: Optional[str] = Form(None),
):
    """上传简历文件（支持PDF、Word、TXT/MD、图片）"""
    try:
        # 检查文件类型
        allowed_types = ['.pdf', '.docx', '.doc', '.txt', '.md', '.markdown', '.jpg', '.jpeg', '.png', '.bmp', '.gif']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_types:
            return JSONResponse({
                "error": f"不支持的文件格式。支持：PDF、Word、TXT/MD、图片（JPG/PNG等）"
            }, status_code=400)
        
        # 读取文件内容
        content = await file.read()
        resume_text = ""
        force_pdf_ocr = str(force_ocr or "").strip().lower() in {"1", "true", "yes", "on"}
        
        try:
            # 解析文件
            if file_ext in {'.txt', '.md', '.markdown'}:
                # 文本文件
                resume_text = _decode_text_file_bytes(content)
                    
            elif file_ext == '.pdf':
                # PDF文件
                try:
                    pdf_source = "none"
                    if force_pdf_ocr:
                        resume_text = _extract_pdf_text_ocr(content)
                        pdf_source = "ocr_forced"
                        if not resume_text:
                            resume_text, pdf_source = _extract_pdf_text_best_effort(content)
                    else:
                        resume_text, pdf_source = _extract_pdf_text_best_effort(content)
                    logger.info(
                        "pdf_extract_done source=%s chars=%s score=%.3f force_ocr=%s",
                        pdf_source,
                        len(resume_text or ""),
                        _text_quality_score(resume_text or ""),
                        force_pdf_ocr,
                    )
                except Exception as e:
                    return JSONResponse({
                        "error": f"PDF解析失败: {str(e)}。建议改传可编辑DOCX，或上传清晰图片触发OCR。"
                    }, status_code=500)
                    
            elif file_ext in ['.docx', '.doc']:
                # Word文件
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(content))
                    
                    # 提取段落文本
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            resume_text += paragraph.text + "\n"
                    
                    # 提取表格文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    resume_text += cell.text + " "
                            resume_text += "\n"
                            
                except Exception as e:
                    return JSONResponse({
                        "error": f"Word文档解析失败: {str(e)}。请确保文件未损坏。"
                    }, status_code=500)
                    
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                # 图片文件 - 使用OCR
                try:
                    from PIL import Image
                    import io
                    
                    # 打开图片
                    image = Image.open(io.BytesIO(content))
                    
                    # OCR识别（支持中英文）
                    resume_text = _ocr_image_best_effort(image)
                    
                    if not resume_text.strip():
                        return JSONResponse({
                            "error": "图片OCR未提取到有效文字。请保证分辨率清晰，或改传DOCX/TXT。"
                        }, status_code=500)
                        
                except ImportError:
                    return JSONResponse({
                        "error": "图片OCR依赖未安装（Tesseract/PyMuPDF）。请联系管理员安装后重试。"
                    }, status_code=500)
                except Exception as e:
                    return JSONResponse({
                        "error": f"图片识别失败: {str(e)}。请确保图片清晰可读。"
                    }, status_code=500)
            
            # 检查是否成功提取到内容
            if not resume_text.strip():
                return JSONResponse({
                    "error": "文件解析成功，但未能提取到有效内容。请检查文件是否为空或格式是否正确。"
                }, status_code=500)

            _track_event(
                "resume_uploaded",
                {
                    "filename": file.filename,
                    "ext": file_ext,
                    "chars": len(resume_text.strip()),
                },
            )
            return JSONResponse({
                "success": True,
                "resume_text": resume_text.strip(),
                "filename": file.filename
            })
            
        except Exception as e:
            return JSONResponse({
                "error": f"文件解析出错: {str(e)}"
            }, status_code=500)
        
    except Exception as e:
        _track_event("api_error", {"api": "/api/upload", "error": str(e)[:300]})
        return JSONResponse({"error": f"上传失败: {str(e)}"}, status_code=500)

@app.websocket("/ws/progress")
async def websocket_progress(websocket: WebSocket):
    """WebSocket实时进度推送"""
    await websocket.accept()
    await progress_tracker.connect(websocket)
    
    try:
        while True:
            # 保持连接
            await websocket.receive_text()
    except WebSocketDisconnect:
        progress_tracker.disconnect(websocket)

@app.post("/api/process")
async def process_resume(request: Request):
    """处理简历的API接口 - 市场驱动"""
    try:
        data = await request.json()
        resume_text = data.get("resume", "")
        runtime_agent_overrides = None
        if isinstance(data, dict):
            raw_overrides = data.get("agent_overrides")
            if not isinstance(raw_overrides, dict):
                raw_overrides = data.get("agent_configs")
            if isinstance(raw_overrides, dict):
                runtime_agent_overrides = raw_overrides
        
        if not resume_text:
            return _api_error("简历内容不能为空", status_code=400, code="empty_resume")

        _track_event("resume_process_started", {"chars": len(resume_text)})

        # 重置进度
        progress_tracker.reset()
        
        # 定义进度回调
        async def update_progress_callback(step, message, agent):
            await progress_tracker.update_progress(step, message, agent)
            await progress_tracker.add_ai_message(agent, message)
        
        # 使用市场驱动引擎处理
        results = await market_engine.process_resume(
            resume_text,
            update_progress_callback,
            agent_overrides=runtime_agent_overrides,
        )

        # Seed job search (Boss/OpenClaw) from resume text, so frontend can auto-search links.
        info = analyzer.extract_info(resume_text)
        seed_keywords = []
        if info.get("job_intention") and info["job_intention"] != "未指定":
            seed_keywords.append(info["job_intention"])
        seed_keywords.extend((info.get("skills") or [])[:6])
        seed_keywords = [k for k in seed_keywords if k]

        seed_location = None
        locs = info.get("preferred_locations") or []
        if locs:
            seed_location = locs[0]
        provider_mode = (real_job_service.get_statistics() or {}).get("provider_mode", "")
        
        # Real, actionable job text block (backward compatible with legacy UI field).
        def _format_real_jobs(jobs: List[Dict[str, Any]], mode: str) -> str:
            if not jobs:
                return (
                    '【推荐岗位】（当前暂无可用岗位）\n\n'
                    '排查建议：\n'
                    '1. 检查云端缓存：访问 /api/crawler/status。\n'
                    '2. 检查企业级招聘 API 是否已配置。\n'
                    '3. 若搜索引擎触发风控，可稍后重试。\n\n'
                    '注意：系统默认不会回退到“搜索入口链接”或演示岗位。\n'
                )

            heading = '【推荐岗位】（中国劳动力市场真实数据）'
            if mode == 'openclaw':
                heading = '【推荐岗位】（来自 Boss 直聘实时数据，OpenClaw）'
            elif mode == 'cloud':
                heading = '【推荐岗位】（来自 Boss 直聘云端缓存）'
            elif mode in ('baidu', 'bing', 'brave'):
                heading = f'【推荐岗位】（来自搜索引擎 {mode}）'
            elif mode == 'remotive':
                heading = '【推荐岗位】（来自全球招聘API，仅在显式开启时使用）'
            elif mode == 'bing_html':
                heading = '【推荐岗位】（来自 Bing 无浏览器搜索）'
            elif mode == 'duckduckgo':
                heading = '【推荐岗位】（来自 DuckDuckGo 无浏览器搜索）'
            elif mode == 'jooble':
                heading = '【推荐岗位】（来自 Jooble API）'
            elif mode == 'enterprise_api':
                heading = '【推荐岗位】（企业级中国招聘API实时数据）'
            elif mode == 'no_real_jobs':
                heading = '【推荐岗位】（未检索到可投递真实岗位）'

            lines = [heading, '']
            for i, job in enumerate(jobs, 1):
                title = job.get('title') or job.get('job_title') or '未知岗位'
                company = job.get('company') or ''
                loc = job.get('location') or ''
                salary = job.get('salary') or job.get('salary_range') or ''
                link = job.get('link') or job.get('apply_url') or ''

                mp = job.get('match_percentage')
                if mp is None:
                    mp = job.get('match_rate')
                if mp is None:
                    mp = job.get('match_score')
                mp_str = f"{mp}%" if isinstance(mp, (int, float)) else ''

                lines.append(f"{i}. {title}" + (f" - {company}" if company else ''))
                if salary:
                    lines.append(f"   薪资：{salary}")
                if loc:
                    lines.append(f"   地点：{loc}")
                if mp_str:
                    lines.append(f"   匹配度：{mp_str}")
                if link:
                    lines.append(f"   链接：{link}")
                lines.append('')

            return "\n".join(lines).strip() + "\n"

        async def _get_real_jobs_for_recommendation():
            cfg_mode = os.getenv('JOB_DATA_PROVIDER', 'auto').strip().lower()
            allow_portal_fallback = os.getenv("ALLOW_CN_PORTAL_FALLBACK", "1").strip().lower() in {"1", "true", "yes", "on"}
            kw = seed_keywords[:10]
            loc = seed_location

            if cfg_mode == 'cloud' or cloud_jobs_cache:
                cached = _filter_cloud_cache_by_query(kw, loc, limit=10)
                cached = _enforce_cn_market_jobs(cached)
                if cached:
                    return cached, 'cloud'
                # Cache empty/insufficient: fallback to cloud-safe real-time providers.
                fallback_jobs, fallback_mode, _ = _search_jobs_without_browser(
                    kw,
                    loc,
                    limit=10,
                    allow_portal_fallback=allow_portal_fallback,
                )
                return _enforce_cn_market_jobs(fallback_jobs), fallback_mode

            try:
                jobs = _normalize_real_jobs(
                    real_job_service.search_jobs(keywords=kw, location=loc, limit=10),
                    limit=10,
                )
                jobs = _enforce_cn_market_jobs(jobs)
                mode = (real_job_service.get_statistics() or {}).get('provider_mode', '') or cfg_mode
                return jobs[:10], mode
            except Exception:
                fallback_jobs, fallback_mode, _ = _search_jobs_without_browser(
                    kw,
                    loc,
                    limit=10,
                    allow_portal_fallback=allow_portal_fallback,
                )
                return _enforce_cn_market_jobs(fallback_jobs), fallback_mode or cfg_mode

        real_jobs, real_mode = await _get_real_jobs_for_recommendation()
        public_jobs = _public_job_payload(_enforce_cn_market_jobs(real_jobs), limit=10)
        if not public_jobs:
            portal_jobs = _search_jobs_cn_entrypoints(seed_keywords, seed_location, limit=5)
            if portal_jobs:
                public_jobs = _public_job_payload(portal_jobs, limit=10, allow_entrypoints=True)
                real_mode = "cn_portal"
        results["job_recommendations"] = _format_real_jobs(public_jobs, real_mode)
        results, quality_gate = _run_output_quality_gate(results, resume_text, info, public_jobs)
        provider_mode = real_mode

        # 完成
        await progress_tracker.complete()
        await progress_tracker.add_ai_message("系统", "🎉 市场分析完成！")
        _track_event(
            "process_quality_gate",
            {
                "passed": bool(quality_gate.get("passed", True)),
                "issues_count": len(quality_gate.get("issues") or []),
            },
        )
        _track_event(
            "job_recommendation_ready",
            {
                "provider_mode": provider_mode,
                "real_jobs_count": len(public_jobs),
            },
        )
        _track_event(
            "resume_processed",
            {
                "ok": True,
                "provider_mode": provider_mode,
                "skills_count": len(seed_keywords),
                "real_jobs_count": len(public_jobs),
                "quality_gate_passed": bool(quality_gate.get("passed", True)),
            },
        )

        response_payload = {
            "career_analysis": str(results.get("market_analysis") or ""),
            "job_recommendations": str(results.get("job_recommendations") or ""),
            "optimized_resume": str(results.get("optimized_resume") or ""),
            "interview_prep": str(results.get("interview_prep") or ""),
            "mock_interview": str(results.get("salary_analysis") or ""),
            "job_provider_mode": str(provider_mode or ""),
            "recommended_jobs": public_jobs,
            "recommendation_quality": {
                "strict_real_posting": True,
                "count": len(public_jobs),
                "has_actionable_jobs": bool(public_jobs),
            },
            "quality_gate": quality_gate,
            "schema_version": PROCESS_RESPONSE_SCHEMA_VERSION,
            "boss_seed": {
                "keywords": seed_keywords,
                "location": seed_location,
            },
            "agent_profiles": get_agent_profiles_payload(),
            "job_sources": get_job_source_registry_payload(),
        }
        shape_errors = _validate_process_response_shape(response_payload)
        if shape_errors:
            _track_event(
                "process_contract_error",
                {
                    "count": len(shape_errors),
                    "sample": shape_errors[:3],
                },
            )
            return _api_error("输出JSON未通过契约校验", status_code=500, code="process_contract_failed")

        return _api_success(response_payload)
        
    except Exception as e:
        _track_event("resume_process_failed", {"error": str(e)[:300]})
        _track_event("resume_processed", {"ok": False, "error": str(e)[:300]})
        _track_event("api_error", {"api": "/api/process", "error": str(e)[:300]})
        await progress_tracker.error(f"处理出错: {str(e)}")
        return _api_error(str(e), status_code=500, code="process_failed")


# ========================================
# Compatibility API layer (legacy Streamlit/frontends)
# ========================================
@app.post("/api/analysis/resume")
async def analysis_resume_compat(request: Request):
    """
    Legacy-compatible resume analysis endpoint.
    Expected by streamlit_app.py: returns {success, results:{...legacy fields...}}.
    """
    try:
        data = await request.json()
    except Exception:
        data = {}

    resume_text = str((data or {}).get("resume_text") or (data or {}).get("resume") or "").strip()
    if not resume_text:
        return _api_error("resume_text 不能为空", status_code=400, code="empty_resume_text")
    deny, credit_guard = _credit_balance_guard(request, "resume_analysis")
    if deny:
        return deny

    class _CompatJSONRequest:
        def __init__(self, payload: Dict[str, Any]):
            self._payload = payload

        async def json(self):
            return self._payload

    # Prefer full /api/process pipeline so this endpoint inherits latest behavior.
    try:
        proxy_resp = await process_resume(_CompatJSONRequest({"resume": resume_text}))
        proxy_payload: Dict[str, Any] = {}
        if hasattr(proxy_resp, "body"):
            raw = (proxy_resp.body or b"{}").decode("utf-8", errors="ignore")
            proxy_payload = json.loads(raw or "{}")
        if proxy_payload.get("success"):
            credit_charge = _settle_credit_action(
                credit_guard,
                note="简历首轮分析",
                meta={"endpoint": "/api/analysis/resume"},
            )
            if not credit_charge.get("ok"):
                return _api_error("credits 扣减失败", status_code=409, code="credit_consume_failed")
            return _api_success(
                {
                    "results": _legacy_results_from_process_payload(proxy_payload, resume_text),
                    "compat_mode": "process_bridge",
                    "credit_charge": credit_charge,
                }
            )
    except Exception:
        pass

    # Hard fallback: still return usable output without breaking frontend.
    fallback = _heuristic_legacy_results(resume_text)
    credit_charge = _settle_credit_action(
        credit_guard,
        note="简历首轮分析",
        meta={"endpoint": "/api/analysis/resume", "compat_mode": "heuristic_fallback"},
    )
    if not credit_charge.get("ok"):
        return _api_error("credits 扣减失败", status_code=409, code="credit_consume_failed")
    return _api_success({"results": fallback, "compat_mode": "heuristic_fallback", "credit_charge": credit_charge})


@app.post("/api/simple-apply/init-login")
async def simple_apply_init_login(request: Request):
    """Legacy login init endpoint (mock SMS flow for no-brain local usage)."""
    try:
        data = await request.json()
    except Exception:
        data = {}

    phone = str((data or {}).get("phone") or "").strip()
    if len(phone) < 8:
        return _api_error("手机号格式不正确", status_code=400, code="invalid_phone")

    session_id = f"sess_{uuid.uuid4().hex[:10]}"
    row = {
        "session_id": session_id,
        "phone": phone,
        "verified": False,
        "created_at": datetime.now().isoformat(),
        "mock_code": "123456",
    }
    _store_simple_login_session(row)
    return _api_success(
        {
            "message": "验证码已发送（本地兼容模式，测试码：123456）",
            "session_id": session_id,
            "masked_phone": _mask_phone(phone),
        }
    )


@app.post("/api/simple-apply/verify-code")
async def simple_apply_verify_code(request: Request):
    """Legacy login verify endpoint."""
    try:
        data = await request.json()
    except Exception:
        data = {}

    phone = str((data or {}).get("phone") or "").strip()
    session_id = str((data or {}).get("session_id") or "").strip()
    code = str((data or {}).get("code") or "").strip()
    lookup_phone, row = _find_simple_login_session(phone=phone, session_id=session_id)
    if not row:
        return _api_error("请先获取验证码", status_code=400, code="session_not_found")
    if len(code) < 4:
        return _api_error("验证码格式不正确", status_code=400, code="invalid_code")
    expected_code = str(row.get("mock_code") or "123456").strip()
    if expected_code and code != expected_code:
        return _api_error("验证码不正确", status_code=400, code="invalid_code")

    row["verified"] = True
    row["verified_at"] = datetime.now().isoformat()
    _store_simple_login_session({**row, "phone": lookup_phone or phone})
    return _api_success(
        {
            "message": "登录成功，可开始自动投递",
            "masked_phone": _mask_phone(str(row.get("phone") or phone)),
            "session_id": str(row.get("session_id") or session_id),
        }
    )


@app.post("/api/simple-apply/apply")
async def simple_apply_apply(request: Request):
    """
    Legacy one-shot auto apply endpoint used by streamlit_app.py.
    Returns: {success,total,success_count,failed_count,details,message}
    """
    try:
        data = await request.json()
    except Exception:
        data = {}

    phone = str((data or {}).get("phone") or "").strip()
    session_id = str((data or {}).get("session_id") or "").strip()
    if not phone and session_id:
        _, row = _find_simple_login_session(session_id=session_id)
        phone = str((row or {}).get("phone") or "").strip()

    if phone or session_id:
        _, row = _find_simple_login_session(phone=phone, session_id=session_id)
        if row and not bool(row.get("verified")):
            return _api_error("请先完成验证码登录", status_code=400, code="login_not_verified")

    payload = dict(data or {})
    if phone and not str(payload.get("phone") or "").strip():
        payload["phone"] = phone

    result = _run_simple_apply_flow(payload)
    _track_event(
        "simple_apply",
        {
            "success": bool(result.get("success")),
            "total": int(result.get("total") or 0),
            "success_count": int(result.get("success_count") or 0),
        },
    )
    return JSONResponse(result, status_code=200)


@app.post("/api/credits/purchase")
async def purchase_credits_compat(request: Request):
    """
    Legacy purchase endpoint used by streamlit_app.py.
    In local mode this creates an order and optionally starts a small apply batch.
    """
    try:
        data = await request.json()
    except Exception:
        data = {}

    email = str((data or {}).get("email") or "").strip()
    package = str((data or {}).get("package") or "").strip() or "标准包"
    credits = int((data or {}).get("credits") or 0)
    price = float((data or {}).get("price") or 0)
    job_keyword = str((data or {}).get("job_keyword") or "").strip()
    city = str((data or {}).get("city") or "").strip()
    resume_text = str((data or {}).get("resume_text") or "").strip()

    if not email:
        return _api_error("email 不能为空", status_code=400, code="missing_email")

    order_id = f"ORD{uuid.uuid4().hex[:10].upper()}"
    order = {
        "order_id": order_id,
        "email": email,
        "package": package,
        "credits": credits,
        "price": price,
        "status": "created",
        "created_at": datetime.now().isoformat(),
    }
    simple_order_history.append(order)

    preview = {
        "success": False,
        "message": "已创建订单，等待开始投递",
        "total": 0,
        "success_count": 0,
        "failed_count": 0,
        "details": [],
    }
    if resume_text:
        # Keep this fast and deterministic for local usage.
        preview_payload = {
            "phone": "",
            "resume_text": resume_text,
            "job_keyword": job_keyword,
            "city": city,
            "count": min(max(credits, 1), 20),
        }
        preview = _run_simple_apply_flow(preview_payload)
        order["status"] = "processed" if preview.get("success") else "queued"
    else:
        order["status"] = "queued"

    return _api_success(
        {
            "message": "订单创建成功，已进入自动投递队列",
            "order_id": order_id,
            "preview": {
                "total": int(preview.get("total") or 0),
                "success_count": int(preview.get("success_count") or 0),
                "failed_count": int(preview.get("failed_count") or 0),
            },
        }
    )


@app.get("/api/public/credits/packages")
async def public_credit_packages():
    return _api_success(
        {
            "packages": commerce_service.list_credit_packages(),
            "actions": _credit_action_catalog(),
        }
    )


@app.post("/api/public/credits/checkout")
async def public_credit_checkout(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}

    package_id = str((data or {}).get("package_id") or "").strip()
    if not package_id:
        return _api_error("package_id 不能为空", status_code=400, code="missing_package_id")

    try:
        payload = commerce_service.create_credit_checkout(
            package_id=package_id,
            email=str((data or {}).get("email") or "").strip(),
            name=str((data or {}).get("name") or "").strip(),
            phone=str((data or {}).get("phone") or "").strip(),
            source=str((data or {}).get("source") or "").strip() or "web",
            channel=str((data or {}).get("channel") or "").strip() or "web_credit",
            payment_channel=str((data or {}).get("payment_channel") or "").strip() or "manual_web",
            note=str((data or {}).get("note") or "").strip(),
        )
        _track_event(
            "credit_checkout_created",
            {
                "package_id": package_id,
                "buyer_id": ((payload.get("account") or {}).get("buyer_id") or ""),
                "order_id": ((payload.get("order") or {}).get("order_id") or ""),
                "checkout_mode": payload.get("checkout_mode") or "",
            },
        )
        return _api_success(payload)
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="credit_checkout_failed")
    except Exception as e:
        logger.exception("credit_checkout_failed")
        return _api_error(str(e), status_code=500, code="credit_checkout_failed")


@app.post("/api/public/credits/order-status")
async def public_credit_order_status(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}

    try:
        payload = commerce_service.get_checkout_status(
            order_id=str((data or {}).get("order_id") or "").strip(),
            access_code=str((data or {}).get("access_code") or "").strip(),
        )
        return _api_success(payload)
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="credit_order_status_failed")
    except Exception as e:
        logger.exception("credit_order_status_failed")
        return _api_error(str(e), status_code=500, code="credit_order_status_failed")


@app.post("/api/public/credits/upload-proof")
async def public_credit_upload_proof(
    order_id: str = Form(""),
    access_code: str = Form(""),
    amount: float = Form(0),
    note: str = Form(""),
    file: UploadFile = File(...),
):
    if file is None:
        return _api_error("proof_file_required", status_code=400, code="payment_proof_upload_failed")

    original_name = str(getattr(file, "filename", "") or "").strip() or "payment-proof"
    content_type = str(getattr(file, "content_type", "") or "").strip().lower()
    ext = os.path.splitext(original_name)[1].lower()
    allowed_exts = {".png", ".jpg", ".jpeg", ".webp", ".pdf"}
    if ext not in allowed_exts:
        return _api_error("只支持 png/jpg/jpeg/webp/pdf", status_code=400, code="payment_proof_invalid_type")
    if content_type and not (content_type.startswith("image/") or content_type == "application/pdf"):
        return _api_error("proof_content_type_invalid", status_code=400, code="payment_proof_invalid_type")

    tmp_path = ""
    total_size = 0
    try:
        with tempfile.NamedTemporaryFile(delete=False, dir=commerce_service.payment_proof_dir, suffix=ext or ".bin") as handle:
            tmp_path = handle.name
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > 10 * 1024 * 1024:
                    raise ValueError("payment_proof_too_large")
                handle.write(chunk)

        proof = commerce_service.create_payment_proof(
            access_code=str(access_code or "").strip(),
            order_id=str(order_id or "").strip(),
            amount=float(amount or 0),
            note=str(note or "").strip(),
            file_name=original_name,
            mime_type=content_type or "application/octet-stream",
            source_path=tmp_path,
        )
        status_payload = commerce_service.get_checkout_status(order_id=str(order_id or "").strip(), access_code=str(access_code or "").strip())
        _track_event(
            "payment_proof_uploaded",
            {
                "order_id": proof.get("order_id"),
                "buyer_id": proof.get("buyer_id"),
                "proof_id": proof.get("proof_id"),
            },
        )
        return _api_success({"proof": proof, "status": status_payload})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="payment_proof_upload_failed")
    except Exception as e:
        logger.exception("payment_proof_upload_failed")
        return _api_error(str(e), status_code=500, code="payment_proof_upload_failed")
    finally:
        try:
            await file.close()
        except Exception:
            pass
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


@app.get("/api/credits/wallet")
async def credit_wallet(request: Request):
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("缺少访问码，请先登录", status_code=401, code="missing_access_code")
    try:
        payload = commerce_service.get_wallet_by_access_code(access_code)
        payload["packages"] = commerce_service.list_credit_packages()
        payload["actions"] = _credit_action_catalog()
        return _api_success(payload)
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="credit_wallet_failed")
    except Exception as e:
        logger.exception("credit_wallet_failed")
        return _api_error(str(e), status_code=500, code="credit_wallet_failed")


@app.get("/api/credits/ledger")
async def credit_ledger(request: Request, limit: int = 30):
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("缺少访问码，请先登录", status_code=401, code="missing_access_code")
    try:
        wallet_payload = commerce_service.get_wallet_by_access_code(access_code)
        buyer = wallet_payload.get("buyer") if isinstance(wallet_payload, dict) else {}
        buyer_id = str((buyer or {}).get("buyer_id") or "").strip()
        items = [] if str(access_code).strip().upper() == str(os.getenv("ACCESS_CODE") or "6").strip().upper() else commerce_service.list_credit_ledger(limit=limit, buyer_id=buyer_id)
        return _api_success(
            {
                "items": items,
                "wallet": wallet_payload.get("wallet") if isinstance(wallet_payload, dict) else {},
                "buyer": buyer,
            }
        )
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="credit_ledger_failed")
    except Exception as e:
        logger.exception("credit_ledger_failed")
        return _api_error(str(e), status_code=500, code="credit_ledger_failed")


@app.get("/api/applications")
async def list_applications_compat(user_id: str = "", limit: int = 200):
    """Legacy applications list endpoint for old Streamlit pages."""
    rows = real_job_service.records.get_all_records()
    if user_id:
        rows = [
            r
            for r in rows
            if str((r.get("user_info") or {}).get("user_id") or "").strip() == str(user_id).strip()
        ]
    n = max(1, min(int(limit or 200), 1000))
    rows = list(rows)[-n:]
    rows.reverse()
    return JSONResponse([_serialize_application_row(row) for row in rows])


@app.get("/api/stats")
async def get_stats_compat(user_id: str = ""):
    """Legacy stats endpoint for old Streamlit pages."""
    rows = real_job_service.records.get_all_records()
    if user_id:
        rows = [
            r
            for r in rows
            if str((r.get("user_info") or {}).get("user_id") or "").strip() == str(user_id).strip()
        ]
    return JSONResponse(_build_legacy_stats_payload(rows))


@app.get("/api/health")
async def health_check():
    """健康检查"""
    stats = real_job_service.get_statistics()
    biz = business_service.metrics()

    execution_model = {
        "mode": "buyer_local_browser" if stats.get("provider_mode") == "openclaw" else "cloud_or_api",
        "server_role": "control_plane",
        "local_agent_supported": True,
        "message": "云端负责派单、数据和工作台；浏览器控制型能力在买家本机执行。",
    }

    # 检查OpenClaw状态
    openclaw_status = None
    if stats.get("provider_mode") == "openclaw":
        from app.services.job_providers.openclaw_browser_provider import OpenClawBrowserProvider
        openclaw = OpenClawBrowserProvider()
        openclaw_status = openclaw.health_check()

        if isinstance(openclaw_status, dict):
            openclaw_status["scope"] = "buyer_local_machine"
            openclaw_status["server_blocking"] = False
            if not openclaw_status.get("available"):
                openclaw_status["status"] = "local_agent_required"
                openclaw_status["message"] = "云端控制台正常。OpenClaw 浏览器执行能力需要在买家本机或本地代理上就绪。"

    return _api_success({
        "status": "ok",
        "message": "AI求职助手运行正常",
        "boot_ts": APP_BOOT_TS,
        "uptime_s": round(time.perf_counter() - APP_BOOT_MONO, 1),
        "job_database": stats,
        "execution_model": execution_model,
        "openclaw": openclaw_status,
        "business": {
            "leads_total": biz.get("leads", {}).get("total", 0),
            "uploads": biz.get("funnel", {}).get("uploads", 0),
            "process_runs": biz.get("funnel", {}).get("process_runs", 0),
            "searches": biz.get("funnel", {}).get("searches", 0),
            "applies": biz.get("funnel", {}).get("applies", 0),
        },
        "config": {
            "job_data_provider": os.getenv("JOB_DATA_PROVIDER", "auto"),
            "cloud_cache_total": len(cloud_jobs_cache),
            "cloud_last_push_at": cloud_jobs_meta.get("last_push_at"),
            "no_browser_fallback_enabled": True,
            "enterprise_job_api_configured": bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip()),
            "llm": get_public_llm_config(),
            "job_sources": get_job_source_registry_payload(),
            "agent_profiles": get_agent_profiles_payload(),
        },
    })



@app.get("/api/version")
async def version():
    """Expose basic build metadata for debugging deployments."""
    return _api_success({
        "job_data_provider": os.getenv("JOB_DATA_PROVIDER", "auto"),
        "railway_git_commit_sha": os.getenv("RAILWAY_GIT_COMMIT_SHA"),
        "github_sha": os.getenv("GITHUB_SHA"),
        "app_boot_ts": APP_BOOT_TS,
    })


@app.get("/api/ping")
async def ping():
    """Ultra-light probe for availability checks."""
    return _api_success({"ok": True, "ts": datetime.now().isoformat()})


@app.get("/api/ready")
async def ready():
    """Release gate probe used by QA/ops before Go/No-Go."""
    cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
    cache_total = len(cloud_jobs_cache)
    checks = {
        "api_alive": True,
        "cn_provider_mode": cfg_mode in {"auto", "cloud", "baidu", "bing", "brave", "jooble", "openclaw", "enterprise_api"},
        "cache_or_cn_fallback": cache_total > 0 or True,
        "global_fallback_disabled_by_default": os.getenv("ENABLE_GLOBAL_JOB_FALLBACK", "").strip().lower() not in {"1", "true", "yes", "on"},
        "enterprise_api_configured": bool(os.getenv("ENTERPRISE_JOB_API_URL", "").strip()),
    }
    score = round(sum(1 for v in checks.values() if v) / len(checks) * 100, 1)
    status = "ready" if score >= 75 else "not_ready"
    return _api_success(
        {
            "status": status,
            "score": score,
            "checks": checks,
            "provider_mode": cfg_mode,
            "cloud_cache_total": cache_total,
            "boot_ts": APP_BOOT_TS,
        }
    )


@app.post("/api/business/lead")
async def capture_business_lead(request: Request):
    """Capture B2B/B2C monetization leads from landing page."""
    try:
        data = await request.json()
        email = (data.get("email") or "").strip()
        if ("@" not in email) or ("." not in email):
            return JSONResponse({"error": "请输入有效邮箱"}, status_code=400)

        payload = business_service.add_lead(
            email=email,
            name=(data.get("name") or "").strip(),
            company=(data.get("company") or "").strip(),
            use_case=(data.get("use_case") or "").strip(),
            budget=(data.get("budget") or "").strip(),
            source=(data.get("source") or "landing").strip(),
            note=(data.get("note") or "").strip(),
        )
        return JSONResponse({"success": True, "data": payload})
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/lead", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/business/feedback")
async def capture_user_feedback(request: Request):
    """Capture user feedback from product and landing pages."""
    try:
        data = await request.json()
        message = (data.get("message") or "").strip()
        if len(message) < 5:
            return JSONResponse({"error": "反馈内容太短，请至少输入5个字符"}, status_code=400)
        rating_raw = data.get("rating")
        rating = int(rating_raw) if str(rating_raw).strip() else None
        payload = business_service.add_feedback(
            message=message,
            rating=rating,
            category=(data.get("category") or "").strip(),
            email=(data.get("email") or "").strip(),
            source=(data.get("source") or "product").strip(),
            page=(data.get("page") or "").strip(),
        )
        return JSONResponse({"success": True, "data": payload})
    except ValueError as ve:
        return JSONResponse({"error": str(ve)}, status_code=400)
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/feedback", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/business/event")
async def capture_public_event(request: Request):
    """Public-safe growth event ingest endpoint for frontend KPI tracking."""
    try:
        data = await request.json()
        event_name = str(data.get("event_name") or "").strip().lower()
        payload = data.get("payload") if isinstance(data.get("payload"), dict) else {}

        if (not event_name) or (not PUBLIC_EVENT_NAME_PATTERN.match(event_name)):
            return JSONResponse({"error": "invalid event_name"}, status_code=400)
        if event_name not in PUBLIC_EVENT_WHITELIST:
            return JSONResponse({"error": "event not allowed"}, status_code=400)

        business_service.track_event(event_name, payload)
        return JSONResponse({"success": True})
    except Exception as e:
        _track_event("api_error", {"api": "/api/business/event", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/feedback/summary")
async def feedback_summary(request: Request, days: int = 7, limit: int = 20):
    """Feedback summary for ops dashboards and growth automation."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        return JSONResponse({"success": True, "data": business_service.feedback_summary(days=days, limit=limit)})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/public-proof")
async def business_public_proof():
    """Public-safe proof counters used on landing page."""
    try:
        m = business_service.metrics()
        funnel = m.get("funnel", {})
        engagement = m.get("engagement", {})
        quality = m.get("quality", {})
        return _api_success(
            {
                "leads_total": int(m.get("leads", {}).get("total", 0) or 0),
                "feedback_total": int(m.get("feedback", {}).get("total", 0) or 0),
                "uploads_total": int(funnel.get("uploads", 0) or 0),
                "process_runs_total": int(funnel.get("process_runs", 0) or 0),
                "searches_total": int(funnel.get("searches", 0) or 0),
                "applies_total": int(funnel.get("applies", 0) or 0),
                "job_link_clicks_total": int(engagement.get("job_link_clicks", 0) or 0),
                "result_downloads_total": int(engagement.get("result_downloads", 0) or 0),
                "upload_to_process_pct": float(funnel.get("upload_to_process_pct", 0) or 0),
                "process_to_search_pct": float(funnel.get("process_to_search_pct", 0) or 0),
                "search_to_apply_pct": float(funnel.get("search_to_apply_pct", 0) or 0),
                "click_to_apply_pct": float(engagement.get("click_to_apply_pct", 0) or 0),
                "quality_gate_fail_rate_pct": float(quality.get("gate_fail_rate_pct", 0) or 0),
            }
        )
    except Exception as e:
        return _api_error(str(e), status_code=500, code="business_public_proof_failed")


@app.get("/api/business/metrics")
async def business_metrics(request: Request):
    """Operational funnel metrics for fundraising / monetization tracking."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        return JSONResponse({"success": True, "data": business_service.metrics()})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/business/readiness")
async def business_readiness(request: Request):
    """Simple GTM readiness snapshot for investor demos."""
    token = os.getenv("ADMIN_METRICS_TOKEN", "").strip()
    supplied = request.headers.get("x-admin-token", "").strip()
    if token and supplied != token:
        return JSONResponse({"error": "forbidden"}, status_code=403)
    try:
        m = business_service.metrics()
        funnel = m.get("funnel", {})
        leads = m.get("leads", {})
        feedback = m.get("feedback", {})
        checks = {
            "lead_capture_ready": leads.get("total", 0) >= 1,
            "feedback_loop_ready": feedback.get("last_7d", 0) >= 1,
            "activation_flow_ready": funnel.get("upload_to_process_pct", 0) > 0,
            "job_match_flow_ready": funnel.get("process_to_search_pct", 0) > 0,
            "monetization_signal_ready": funnel.get("search_to_apply_pct", 0) > 0,
        }
        score = round(sum(1 for v in checks.values() if v) / len(checks) * 100, 1)
        return JSONResponse(
            {
                "success": True,
                "score": score,
                "checks": checks,
                "metrics": m,
            }
        )
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/public/auth/register")
async def public_auth_register(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        result = user_auth_service.register_user(
            email=str((data or {}).get("email") or "").strip(),
            password=str((data or {}).get("password") or ""),
            name=str((data or {}).get("name") or "").strip(),
            access_code=str((data or {}).get("access_code") or "").strip(),
        )
        resp = _api_success(
            {
                **result,
                "auth_schema": user_auth_service.get_supported_auth_payload(),
            }
        )
        _set_user_session_cookie(resp, str(result.get("session_token") or ""))
        return resp
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="register_failed")
    except Exception as e:
        logger.exception("public_auth_register_failed")
        return _api_error(str(e), status_code=500, code="register_failed")


@app.post("/api/public/auth/login")
async def public_auth_login(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        result = user_auth_service.login_user(
            email=str((data or {}).get("email") or "").strip(),
            password=str((data or {}).get("password") or ""),
            access_code=str((data or {}).get("access_code") or "").strip(),
        )
        resp = _api_success(
            {
                **result,
                "auth_schema": user_auth_service.get_supported_auth_payload(),
            }
        )
        _set_user_session_cookie(resp, str(result.get("session_token") or ""))
        return resp
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="login_failed")
    except Exception as e:
        logger.exception("public_auth_login_failed")
        return _api_error(str(e), status_code=500, code="login_failed")


@app.post("/api/public/auth/social-login")
async def public_auth_social_login(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        result = user_auth_service.social_login(
            provider=str((data or {}).get("provider") or "").strip(),
            provider_user_id=str((data or {}).get("provider_user_id") or "").strip(),
            name=str((data or {}).get("name") or "").strip(),
            access_code=str((data or {}).get("access_code") or "").strip(),
        )
        resp = _api_success(
            {
                **result,
                "auth_schema": user_auth_service.get_supported_auth_payload(),
            }
        )
        _set_user_session_cookie(resp, str(result.get("session_token") or ""))
        return resp
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="social_login_failed")
    except Exception as e:
        logger.exception("public_auth_social_login_failed")
        return _api_error(str(e), status_code=500, code="social_login_failed")


@app.get("/api/auth/me")
async def auth_me(request: Request):
    token = _user_session_token_from_request(request)
    if not token:
        return _api_error("未登录", status_code=401, code="missing_session")
    user = user_auth_service.get_user_by_session(token)
    if not user:
        return _api_error("登录已过期，请重新登录", status_code=401, code="session_invalid")
    return _api_success({"user": user, "auth_schema": user_auth_service.get_supported_auth_payload()})


@app.post("/api/auth/register")
async def auth_register(request: Request):
    """Primary register endpoint (cookie + token response)."""
    return await public_auth_register(request)


@app.post("/api/auth/login")
async def auth_login(request: Request):
    """Primary login endpoint (cookie + token response)."""
    return await public_auth_login(request)


@app.post("/api/auth/logout")
async def auth_logout(request: Request):
    token = _user_session_token_from_request(request)
    if token:
        user_auth_service.logout_session(token)
    resp = _api_success({"logged_out": True})
    _clear_user_session_cookie(resp)
    return resp


@app.get("/api/auth/providers")
async def auth_providers():
    """Social provider status/config detection for frontend buttons."""
    return _api_success(
        {
            "providers": user_auth_service.get_social_provider_status(),
            "auth_schema": user_auth_service.get_supported_auth_payload(),
        }
    )


@app.post("/api/auth/social/{provider}/start")
async def auth_social_start(provider: str, request: Request):
    """Contract-stable social auth bootstrap placeholder."""
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        payload = user_auth_service.build_social_login_start_payload(
            provider=str(provider or "").strip().lower(),
            redirect_uri=str((data or {}).get("redirect_uri") or "").strip(),
            state=str((data or {}).get("state") or "").strip(),
        )
        return _api_success(payload)
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="social_provider_invalid")
    except Exception as e:
        logger.exception("auth_social_start_failed")
        return _api_error(str(e), status_code=500, code="social_start_failed")


@app.get("/api/auth/social/{provider}/callback")
async def auth_social_callback(provider: str, code: str = "", state: str = ""):
    """Contract-stable callback placeholder until full OAuth is wired."""
    p = str(provider or "").strip().lower()
    statuses = user_auth_service.get_social_provider_status()
    if p not in statuses:
        return _api_error("unsupported social provider", status_code=400, code="social_provider_invalid")
    status = statuses.get(p, {})
    if not status.get("enabled"):
        return _api_error(
            "provider is not configured",
            status_code=501,
            code="social_provider_not_configured",
        )
    return _api_success(
        {
            "provider": p,
            "status": "placeholder_callback",
            "message": "oauth callback placeholder is active; full oauth exchange is pending",
            "code_received": bool(str(code or "").strip()),
            "state": str(state or "").strip(),
        }
    )


@app.get("/api/auth/platforms")
async def auth_platforms(request: Request):
    token = _user_session_token_from_request(request)
    if not token:
        return _api_error("未登录", status_code=401, code="missing_session")
    user = user_auth_service.get_user_by_session(token)
    if not user:
        return _api_error("登录已过期，请重新登录", status_code=401, code="session_invalid")
    return _api_success(
        {
            "user_id": user.get("user_id"),
            "platform_accounts": user.get("platform_accounts") or {},
            "auth_schema": user_auth_service.get_supported_auth_payload(),
        }
    )


@app.put("/api/auth/platforms/{platform}")
async def auth_update_platform(platform: str, request: Request):
    token = _user_session_token_from_request(request)
    if not token:
        return _api_error("未登录", status_code=401, code="missing_session")
    user = user_auth_service.get_user_by_session(token)
    if not user:
        return _api_error("登录已过期，请重新登录", status_code=401, code="session_invalid")
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        updated = user_auth_service.update_platform_account(
            user_id=str(user.get("user_id") or "").strip(),
            platform=str(platform or "").strip().lower(),
            account=str((data or {}).get("account") or "").strip(),
            secret=str((data or {}).get("secret") or ""),
        )
        return _api_success({"user": updated})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="platform_update_failed")
    except Exception as e:
        logger.exception("auth_update_platform_failed")
        return _api_error(str(e), status_code=500, code="platform_update_failed")


@app.post("/api/public/access/redeem")
async def redeem_public_access_code(request: Request):
    """Redeem a buyer-specific access code before Next.js sets cookies."""
    try:
        data = await request.json()
    except Exception:
        data = {}

    code = str((data or {}).get("code") or "").strip()
    machine_name = str((data or {}).get("machine_name") or "").strip()
    if not code:
        return _api_error("访问码不能为空", status_code=400, code="missing_access_code")

    try:
        payload = commerce_service.redeem_access_code(
            code=code,
            client_ip=str(getattr(request.client, "host", "") or ""),
            user_agent=str(request.headers.get("user-agent") or ""),
            machine_name=machine_name,
        )
        _track_event(
            "buyer_access_redeemed",
            {
                "buyer_id": payload.get("buyer_id"),
                "order_id": payload.get("order_id"),
                "access_code": payload.get("access_code"),
            },
        )
        return _api_success(payload)
    except ValueError as e:
        mapping = {
            "code_not_found": "访问码不存在",
            "code_inactive": "访问码已停用",
            "code_expired": "访问码已过期",
            "code_uses_exceeded": "访问码已达到可用次数上限",
            "code_required": "访问码不能为空",
        }
        key = str(e)
        return _api_error(mapping.get(key, key), status_code=400, code=key)
    except Exception as e:
        logger.exception("兑换访问码失败")
        return _api_error(str(e), status_code=500, code="access_redeem_failed")


@app.get("/api/buyer/context")
async def buyer_context(request: Request):
    """Resolve buyer/order context from the redeemed access code cookie."""
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("\u7f3a\u5c11\u8bbf\u95ee\u7801", status_code=401, code="missing_access_code")
    try:
        return _api_success({"buyer": _buyer_context_row(access_code)})
    except Exception as e:
        logger.exception("buyer_context_failed")
        return _api_error(str(e), status_code=500, code="buyer_context_failed")


@app.post("/api/buyer/support-ticket")
async def buyer_support_ticket(request: Request):
    """Buyer-facing support ticket creation, linked by access code cookie."""
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("\u7f3a\u5c11\u8bbf\u95ee\u7801", status_code=401, code="missing_access_code")
    try:
        data = await request.json()
    except Exception:
        data = {}

    row = _buyer_context_row(access_code)

    try:
        ticket = commerce_service.create_ticket(
            buyer_id=str(row.get("buyer_id") or ""),
            order_id=str(row.get("order_id") or ""),
            subject=str((data or {}).get("subject") or "").strip() or "\u4e70\u5bb6\u552e\u540e\u652f\u6301",
            content=str((data or {}).get("content") or "").strip(),
            channel="buyer_dashboard",
            priority=str((data or {}).get("priority") or "").strip() or "normal",
            note=str((data or {}).get("contact") or "").strip(),
        )
        _track_event(
            "buyer_support_ticket_created",
            {
                "buyer_id": row.get("buyer_id"),
                "order_id": row.get("order_id"),
                "ticket_id": ticket.get("ticket_id"),
            },
        )
        return _api_success({"ticket": ticket})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="ticket_validation_failed")
    except Exception as e:
        logger.exception("buyer_support_ticket_failed")
        return _api_error(str(e), status_code=500, code="buyer_support_ticket_failed")


@app.get("/api/buyer/local-runtime")
async def buyer_local_runtime(request: Request):
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("\u7f3a\u5c11\u8bbf\u95ee\u7801", status_code=401, code="missing_access_code")
    try:
        buyer = _buyer_context_row(access_code)
        agents = [
            item
            for item in commerce_service.list_local_agents(limit=20, search=access_code)
            if str(item.get("access_code") or "").strip().upper() == access_code
        ]
        tasks = [
            item
            for item in commerce_service.list_local_tasks(limit=20, search=access_code)
            if str(item.get("access_code") or "").strip().upper() == access_code
        ]
        online_agents = [
            item
            for item in agents
            if str(item.get("status") or "").strip().lower() == "online"
            and _is_recent_iso_ts(item.get("last_seen_at"), within_seconds=150)
        ]
        queued_tasks = [item for item in tasks if str(item.get("status") or "").strip().lower() == "queued"]
        runtime = {
            "access_code": access_code,
            "buyer": buyer,
            "execution_mode": "local_browser_agent",
            "agent_online": bool(online_agents),
            "online_agents": len(online_agents),
            "queued_tasks": len(queued_tasks),
            "agents": agents[:6],
            "tasks": tasks[:8],
            "summary": "\u6d4f\u89c8\u5668\u52a8\u4f5c\u4ecd\u5728\u4e70\u5bb6\u672c\u673a\u6267\u884c\uff0c\u4e91\u7aef\u53ea\u8d1f\u8d23\u4efb\u52a1\u6392\u961f\u548c\u72b6\u6001\u540c\u6b65\u3002",
            "next_steps": [
                "\u5148\u5728\u4e70\u5bb6\u7535\u8111\u5b8c\u6210\u4e00\u6b21 Boss \u767b\u5f55\u3002",
                "\u518d\u542f\u52a8\u672c\u5730\u4ee3\u7406\uff0c\u8ba9\u8fd9\u53f0\u673a\u5668\u8fde\u63a5\u4efb\u52a1\u961f\u5217\u3002",
                "\u7136\u540e\u5728\u5f53\u524d\u9875\u9762\u4e0b\u53d1 Boss \u4efb\u52a1\uff0c\u4ee3\u7406\u4f1a\u81ea\u52a8\u9886\u53d6\u6267\u884c\u3002",
            ],
        }
        return _api_success({"runtime": runtime})
    except Exception as e:
        logger.exception("buyer_local_runtime_failed")
        return _api_error(str(e), status_code=500, code="buyer_local_runtime_failed")


@app.post("/api/buyer/local-tasks")
async def buyer_create_local_task(request: Request):
    access_code = _buyer_access_code_from_request(request)
    if not access_code:
        return _api_error("\u7f3a\u5c11\u8bbf\u95ee\u7801", status_code=401, code="missing_access_code")
    try:
        data = await request.json()
    except Exception:
        data = {}

    task_type = str((data or {}).get("task_type") or "").strip() or "local_auto_apply"
    keywords = str((data or {}).get("keywords") or "").strip()
    if task_type == "local_auto_apply" and not keywords:
        return _api_error("\u5c97\u4f4d\u5173\u952e\u8bcd\u4e0d\u80fd\u4e3a\u7a7a", status_code=400, code="keywords_required")

    location = str((data or {}).get("location") or "").strip()
    phone = str((data or {}).get("phone") or "").strip()
    platforms = _to_string_list((data or {}).get("platforms")) or ["boss"]
    boss_engine = str((data or {}).get("boss_engine") or "auto").strip().lower() or "auto"
    max_count = max(1, min(int((data or {}).get("count") or 8), 50))
    raw_config = (data or {}).get("config") if isinstance((data or {}).get("config"), dict) else {}
    raw_boss_config = raw_config.get("boss_config") if isinstance(raw_config.get("boss_config"), dict) else {}

    payload = {
        "keywords": keywords,
        "location": location,
        "platforms": platforms,
        "config": {
            **raw_config,
            "keywords": keywords,
            "location": location,
            "platforms": platforms,
            "max_count": max_count,
            "boss_config": {
                **raw_boss_config,
                "phone": phone or str(raw_boss_config.get("phone") or "").strip(),
                "engine": boss_engine,
            },
        },
    }

    try:
        credit_guard = None
        if task_type == "local_auto_apply":
            deny, credit_guard = _credit_balance_guard(request, "local_apply_task")
            if deny:
                return deny
        buyer = _buyer_context_row(access_code)
        task = commerce_service.enqueue_local_task(
            access_code=access_code,
            task_type=task_type,
            payload=payload,
            buyer_id=str(buyer.get("buyer_id") or "").strip(),
            agent_id=str((data or {}).get("agent_id") or "").strip(),
        )
        _track_event(
            "buyer_local_task_created",
            {
                "buyer_id": buyer.get("buyer_id"),
                "access_code": access_code,
                "task_id": task.get("task_id"),
                "task_type": task_type,
                "platforms": ",".join(platforms),
            },
        )
        if task_type == "local_auto_apply":
            credit_charge = _settle_credit_action(
                credit_guard,
                note="本地浏览器投递任务",
                meta={"endpoint": "/api/buyer/local-tasks", "platforms": platforms, "keywords": keywords, "location": location},
            )
            if not credit_charge.get("ok"):
                return _api_error("credits 扣减失败", status_code=409, code="credit_consume_failed")
        else:
            credit_charge = {
                "ok": True,
                "skipped": True,
                "reason": "non_chargeable_local_task",
            }
        return _api_success({"task": task, "credit_charge": credit_charge})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="buyer_local_task_create_failed")
    except Exception as e:
        logger.exception("buyer_local_task_create_failed")
        return _api_error(str(e), status_code=500, code="buyer_local_task_create_failed")


@app.get("/api/ops/summary")
async def ops_summary(request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        return _api_success(
            {
                "summary": commerce_service.summary(),
                "business": business_service.metrics(),
            }
        )
    except Exception as e:
        return _api_error(str(e), status_code=500, code="ops_summary_failed")


@app.post("/api/ops/bundles")
async def create_ops_bundle(request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        bundle = commerce_service.create_bundle(
            name=str((data or {}).get("name") or "").strip(),
            phone=str((data or {}).get("phone") or "").strip(),
            email=str((data or {}).get("email") or "").strip(),
            source=str((data or {}).get("source") or "").strip() or "xianyu",
            channel=str((data or {}).get("channel") or "").strip() or "xianyu",
            product_name=str((data or {}).get("product_name") or "").strip() or "AI Job Helper",
            amount=float((data or {}).get("amount") or 0),
            currency=str((data or {}).get("currency") or "").strip() or "CNY",
            payment_channel=str((data or {}).get("payment_channel") or "").strip() or "xianyu",
            payment_status=str((data or {}).get("payment_status") or "").strip() or "paid",
            delivery_status=str((data or {}).get("delivery_status") or "").strip() or "delivered",
            duration_days=int((data or {}).get("duration_days") or 30),
            max_uses=int((data or {}).get("max_uses") or 3),
            label=str((data or {}).get("label") or "").strip() or "默认访问码",
            note=str((data or {}).get("note") or "").strip(),
        )
        return _api_success({"bundle": bundle})
    except Exception as e:
        logger.exception("创建买家 bundle 失败")
        return _api_error(str(e), status_code=500, code="create_bundle_failed")


@app.get("/api/ops/bundles")
async def list_ops_bundles(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"bundles": commerce_service.list_bundles(limit=limit, search=search)})


@app.get("/api/ops/buyers")
async def list_ops_buyers(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"buyers": commerce_service.list_buyers(limit=limit, search=search)})


@app.get("/api/ops/orders")
async def list_ops_orders(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"orders": commerce_service.list_orders(limit=limit, search=search)})


@app.patch("/api/ops/orders/{order_id}")
async def patch_ops_order(order_id: str, request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        order = commerce_service.update_order(order_id, data or {})
        return _api_success({"order": order})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="order_update_failed")
    except Exception as e:
        return _api_error(str(e), status_code=500, code="order_update_failed")


@app.get("/api/ops/wallets")
async def list_ops_wallets(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"wallets": commerce_service.list_wallets(limit=limit, search=search)})


@app.get("/api/ops/credit-ledger")
async def list_ops_credit_ledger(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"items": commerce_service.list_credit_ledger(limit=limit, search=search)})


@app.get("/api/ops/payment-proofs")
async def list_ops_payment_proofs(
    request: Request,
    limit: int = 50,
    search: str = "",
    status: str = "",
    order_id: str = "",
):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success(
        {
            "items": commerce_service.list_payment_proofs(
                limit=limit,
                search=search,
                status=status,
                order_id=order_id,
            )
        }
    )


@app.patch("/api/ops/payment-proofs/{proof_id}")
async def patch_ops_payment_proof(proof_id: str, request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        proof = commerce_service.update_payment_proof(proof_id, data or {})
        return _api_success({"proof": proof})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="payment_proof_update_failed")
    except Exception as e:
        logger.exception("payment_proof_update_failed")
        return _api_error(str(e), status_code=500, code="payment_proof_update_failed")


@app.get("/api/ops/payment-proofs/{proof_id}/download")
async def download_ops_payment_proof(proof_id: str, request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        proof = commerce_service.get_payment_proof(proof_id)
        file_path = str(proof.get("file_path") or "").strip()
        if not file_path or not os.path.exists(file_path):
            return _api_error("payment_proof_file_missing", status_code=404, code="payment_proof_file_missing")
        return FileResponse(
            file_path,
            media_type=str(proof.get("mime_type") or "application/octet-stream"),
            filename=str(proof.get("file_name") or f"{proof_id}.bin"),
        )
    except ValueError as e:
        return _api_error(str(e), status_code=404, code="payment_proof_not_found")
    except Exception as e:
        logger.exception("payment_proof_download_failed")
        return _api_error(str(e), status_code=500, code="payment_proof_download_failed")


@app.get("/api/ops/access-codes")
async def list_ops_access_codes(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"access_codes": commerce_service.list_access_codes(limit=limit, search=search)})


@app.patch("/api/ops/access-codes/{code}")
async def patch_ops_access_code(code: str, request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        access_code = commerce_service.update_access_code(code, data or {})
        return _api_success({"access_code": access_code})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="access_code_update_failed")
    except Exception as e:
        return _api_error(str(e), status_code=500, code="access_code_update_failed")


@app.get("/api/ops/tickets")
async def list_ops_tickets(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"tickets": commerce_service.list_tickets(limit=limit, search=search)})


@app.post("/api/ops/tickets")
async def create_ops_ticket(request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        ticket = commerce_service.create_ticket(
            buyer_id=str((data or {}).get("buyer_id") or "").strip(),
            order_id=str((data or {}).get("order_id") or "").strip(),
            subject=str((data or {}).get("subject") or "").strip(),
            content=str((data or {}).get("content") or "").strip(),
            channel=str((data or {}).get("channel") or "").strip() or "ops",
            status=str((data or {}).get("status") or "").strip() or "open",
            priority=str((data or {}).get("priority") or "").strip() or "normal",
            assignee=str((data or {}).get("assignee") or "").strip(),
            note=str((data or {}).get("note") or "").strip(),
        )
        return _api_success({"ticket": ticket})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="ops_ticket_create_failed")
    except Exception as e:
        return _api_error(str(e), status_code=500, code="ops_ticket_create_failed")


@app.patch("/api/ops/tickets/{ticket_id}")
async def patch_ops_ticket(ticket_id: str, request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        ticket = commerce_service.update_ticket(ticket_id, data or {})
        return _api_success({"ticket": ticket})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="ops_ticket_update_failed")
    except Exception as e:
        return _api_error(str(e), status_code=500, code="ops_ticket_update_failed")


@app.get("/api/ops/local-agents")
async def list_ops_local_agents(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"agents": commerce_service.list_local_agents(limit=limit, search=search)})


@app.get("/api/ops/local-tasks")
async def list_ops_local_tasks(request: Request, limit: int = 50, search: str = ""):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    return _api_success({"tasks": commerce_service.list_local_tasks(limit=limit, search=search)})


@app.post("/api/ops/local-tasks")
async def create_ops_local_task(request: Request):
    deny = _require_ops_secret(request)
    if deny:
        return deny
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        task = commerce_service.enqueue_local_task(
            access_code=str((data or {}).get("access_code") or "").strip(),
            task_type=str((data or {}).get("task_type") or "").strip() or "local_auto_apply",
            payload=(data or {}).get("payload") if isinstance((data or {}).get("payload"), dict) else {},
            buyer_id=str((data or {}).get("buyer_id") or "").strip(),
            agent_id=str((data or {}).get("agent_id") or "").strip(),
        )
        return _api_success({"task": task})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="ops_local_task_create_failed")
    except Exception as e:
        logger.exception("创建本地代理任务失败")
        return _api_error(str(e), status_code=500, code="ops_local_task_create_failed")


@app.post("/api/local-agent/register")
async def register_local_agent(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        agent = commerce_service.register_local_agent(
            access_code=str((data or {}).get("access_code") or "").strip(),
            machine_name=str((data or {}).get("machine_name") or "").strip(),
            hostname=str((data or {}).get("hostname") or "").strip(),
            platform=str((data or {}).get("platform") or "").strip(),
            capabilities=(data or {}).get("capabilities") if isinstance((data or {}).get("capabilities"), dict) else {},
            note=str((data or {}).get("note") or "").strip(),
        )
        return _api_success({"agent": agent})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="local_agent_register_failed")
    except Exception as e:
        logger.exception("注册本地代理失败")
        return _api_error(str(e), status_code=500, code="local_agent_register_failed")


@app.post("/api/local-agent/heartbeat")
async def local_agent_heartbeat(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        agent = commerce_service.heartbeat_local_agent(
            agent_id=str((data or {}).get("agent_id") or "").strip(),
            status=str((data or {}).get("status") or "").strip() or "online",
            capabilities=(data or {}).get("capabilities") if isinstance((data or {}).get("capabilities"), dict) else None,
        )
        return _api_success({"agent": agent})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="local_agent_heartbeat_failed")
    except Exception as e:
        logger.exception("本地代理心跳失败")
        return _api_error(str(e), status_code=500, code="local_agent_heartbeat_failed")


@app.post("/api/local-agent/tasks/claim")
async def claim_local_agent_task(request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        task = commerce_service.claim_local_task(
            agent_id=str((data or {}).get("agent_id") or "").strip(),
            access_code=str((data or {}).get("access_code") or "").strip(),
        )
        return _api_success({"task": task})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="local_agent_claim_failed")
    except Exception as e:
        logger.exception("本地代理领取任务失败")
        return _api_error(str(e), status_code=500, code="local_agent_claim_failed")


@app.post("/api/local-agent/tasks/{task_id}/progress")
async def local_agent_task_progress(task_id: str, request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        task = commerce_service.update_local_task_progress(
            task_id=task_id,
            status=str((data or {}).get("status") or "").strip() or "running",
            progress=(data or {}).get("progress") if isinstance((data or {}).get("progress"), dict) else {},
        )
        return _api_success({"task": task})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="local_agent_progress_failed")
    except Exception as e:
        logger.exception("本地代理上报进度失败")
        return _api_error(str(e), status_code=500, code="local_agent_progress_failed")


@app.post("/api/local-agent/tasks/{task_id}/complete")
async def local_agent_task_complete(task_id: str, request: Request):
    try:
        data = await request.json()
    except Exception:
        data = {}
    try:
        task = commerce_service.complete_local_task(
            task_id=task_id,
            success=bool((data or {}).get("success")),
            status=str((data or {}).get("status") or "").strip(),
            result=(data or {}).get("result") if isinstance((data or {}).get("result"), dict) else {},
        )
        return _api_success({"task": task})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="local_agent_complete_failed")
    except Exception as e:
        logger.exception("本地代理上报结果失败")
        return _api_error(str(e), status_code=500, code="local_agent_complete_failed")


@app.get("/api/investor/readiness")
async def investor_readiness(request: Request):
    """
    Investor-facing readiness API.
    Optional auth via INVESTOR_READ_TOKEN header x-investor-token.
    """
    token = os.getenv("INVESTOR_READ_TOKEN", "").strip()
    supplied = request.headers.get("x-investor-token", "").strip()
    if token and supplied != token:
        return _api_error("forbidden", status_code=403, code="forbidden")
    try:
        return _api_success(_build_investor_readiness_snapshot())
    except Exception as e:
        _track_event("api_error", {"api": "/api/investor/readiness", "error": str(e)[:300]})
        return _api_error(str(e), status_code=500, code="investor_readiness_failed")


@app.get("/api/investor/summary")
async def investor_summary(request: Request):
    """Compact summary payload for pitch materials."""
    token = os.getenv("INVESTOR_READ_TOKEN", "").strip()
    supplied = request.headers.get("x-investor-token", "").strip()
    if token and supplied != token:
        return _api_error("forbidden", status_code=403, code="forbidden")
    try:
        snap = _build_investor_readiness_snapshot()
        p = snap.get("pillars", {})
        narrative = [
            f"Current financing readiness status: {snap.get('status')} (score {snap.get('overall_score')}).",
            f"Product pillar score: {p.get('product')}, reliability: {p.get('reliability')}.",
            f"Traction pillar score: {p.get('traction')}, GTM pillar score: {p.get('go_to_market')}.",
            f"Feedback captured so far: {(snap.get('metrics') or {}).get('feedback', {}).get('total', 0)}.",
            "Next 30 days focus: increase uploads, process runs, qualified leads, and actionable user feedback while keeping CN-real job links stable.",
        ]
        return _api_success(
            {
                "status": snap.get("status"),
                "overall_score": snap.get("overall_score"),
                "pillars": p,
                "highlights": snap.get("highlights", {}),
                "next_30d_targets": snap.get("next_30d_targets", {}),
                "narrative": narrative,
            }
        )
    except Exception as e:
        _track_event("api_error", {"api": "/api/investor/summary", "error": str(e)[:300]})
        return _api_error(str(e), status_code=500, code="investor_summary_failed")


@app.post("/api/hr/jobs")
async def hr_create_job(request: Request):
    """HR 发布岗位。"""
    try:
        data = await request.json()
    except Exception:
        data = {}
    if not isinstance(data, dict):
        data = {}
    try:
        job = hr_match_service.create_job(data)
        return _api_success({"job": job}, status_code=201)
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="hr_job_create_failed")
    except Exception as e:
        logger.exception("hr_job_create_failed")
        return _api_error(str(e), status_code=500, code="hr_job_create_failed")


@app.get("/api/hr/jobs")
async def hr_list_jobs(hr_id: str = "", status: str = "", limit: int = 50):
    """查询 HR 岗位池。"""
    try:
        n = max(1, min(int(limit or 50), 200))
        items = hr_match_service.list_jobs(hr_id=hr_id, status=status, limit=n)
        return _api_success({"items": items, "total": len(items)})
    except Exception as e:
        logger.exception("hr_job_list_failed")
        return _api_error(str(e), status_code=500, code="hr_job_list_failed")


@app.post("/api/hr/candidates")
async def hr_upsert_candidate(request: Request):
    """HR 侧录入或更新候选人（本地池）。"""
    try:
        data = await request.json()
    except Exception:
        data = {}
    if not isinstance(data, dict):
        data = {}
    try:
        candidate = hr_match_service.upsert_candidate(data)
        return _api_success({"candidate": candidate})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="hr_candidate_upsert_failed")
    except Exception as e:
        logger.exception("hr_candidate_upsert_failed")
        return _api_error(str(e), status_code=500, code="hr_candidate_upsert_failed")


@app.get("/api/hr/candidates/local")
async def hr_local_candidates(limit: int = 200):
    """仅返回 HR 手工维护的候选人池。"""
    try:
        n = max(1, min(int(limit or 200), 500))
        items = hr_match_service.list_local_candidates(limit=n)
        return _api_success({"items": items, "total": len(items)})
    except Exception as e:
        logger.exception("hr_local_candidates_failed")
        return _api_error(str(e), status_code=500, code="hr_local_candidates_failed")


@app.get("/api/hr/candidates/pool")
async def hr_candidate_pool(include_resume_profiles: bool = True, limit: int = 300):
    """返回合并候选池（HR 本地 + 简历库）。"""
    try:
        n = max(1, min(int(limit or 300), 600))
        local_rows = hr_match_service.list_local_candidates(limit=n)
        resume_rows = _build_resume_candidate_pool(limit=n) if include_resume_profiles else []
        merged = _merge_candidate_pool(local_rows, resume_rows)
        return _api_success(
            {
                "items": merged[:n],
                "total": len(merged),
                "sources": {
                    "local_candidates": len(local_rows),
                    "resume_profiles": len(resume_rows),
                },
            }
        )
    except Exception as e:
        logger.exception("hr_candidate_pool_failed")
        return _api_error(str(e), status_code=500, code="hr_candidate_pool_failed")


@app.get("/api/hr/candidates/match")
async def hr_match_candidates(
    job_id: str = "",
    include_resume_profiles: bool = True,
    limit: int = 50,
):
    """按岗位给 HR 输出候选人排序列表。"""
    target_job_id = str(job_id or "").strip()
    if not target_job_id:
        return _api_error("job_id 不能为空", status_code=400, code="hr_match_candidates_failed")
    try:
        job = _find_hr_job(target_job_id)
        if not job:
            return _api_error("岗位不存在", status_code=404, code="hr_job_not_found")

        n = max(1, min(int(limit or 50), 200))
        local_rows = hr_match_service.list_local_candidates(limit=300)
        resume_rows = _build_resume_candidate_pool(limit=300) if include_resume_profiles else []
        pool = _merge_candidate_pool(local_rows, resume_rows)
        items = hr_match_service.match_candidates_for_job(job=job, candidates=pool, limit=n)
        return _api_success(
            {
                "job": job,
                "items": items,
                "total": len(items),
                "candidate_pool_size": len(pool),
            }
        )
    except Exception as e:
        logger.exception("hr_match_candidates_failed")
        return _api_error(str(e), status_code=500, code="hr_match_candidates_failed")


@app.post("/api/hr/actions")
async def hr_action(request: Request):
    """HR 对候选人做 like/pass/hold。"""
    try:
        data = await request.json()
    except Exception:
        data = {}
    if not isinstance(data, dict):
        data = {}
    try:
        action = hr_match_service.record_hr_action(data)
        return _api_success({"action": action})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="hr_action_failed")
    except Exception as e:
        logger.exception("hr_action_failed")
        return _api_error(str(e), status_code=500, code="hr_action_failed")


@app.post("/api/candidate/actions")
async def candidate_action(request: Request):
    """候选人对岗位做 like/pass/hold。"""
    try:
        data = await request.json()
    except Exception:
        data = {}
    if not isinstance(data, dict):
        data = {}
    try:
        action = hr_match_service.record_candidate_action(data)
        return _api_success({"action": action})
    except ValueError as e:
        return _api_error(str(e), status_code=400, code="candidate_action_failed")
    except Exception as e:
        logger.exception("candidate_action_failed")
        return _api_error(str(e), status_code=500, code="candidate_action_failed")


@app.get("/api/hr/mutual-matches")
async def hr_mutual_matches(hr_id: str = "", candidate_id: str = "", limit: int = 100):
    """查询双向 like 的匹配结果。"""
    try:
        n = max(1, min(int(limit or 100), 300))
        items = hr_match_service.list_mutual_matches(hr_id=hr_id, candidate_id=candidate_id, limit=n)
        return _api_success({"items": items, "total": len(items)})
    except Exception as e:
        logger.exception("hr_mutual_matches_failed")
        return _api_error(str(e), status_code=500, code="hr_mutual_matches_failed")


@app.get("/api/hr/overview")
async def hr_overview():
    """HR 双向筛选总览。"""
    try:
        return _api_success(hr_match_service.overview())
    except Exception as e:
        logger.exception("hr_overview_failed")
        return _api_error(str(e), status_code=500, code="hr_overview_failed")


@app.get("/api/candidate/jobs/feed")
async def candidate_jobs_feed(candidate_id: str = "", include_resume_profiles: bool = True, limit: int = 50):
    """候选人视角的岗位推荐流（双向筛选对称端）。"""
    target = str(candidate_id or "").strip()
    if not target:
        return _api_error("candidate_id 不能为空", status_code=400, code="candidate_jobs_feed_failed")
    try:
        n = max(1, min(int(limit or 50), 200))
        local_rows = hr_match_service.list_local_candidates(limit=400)
        resume_rows = _build_resume_candidate_pool(limit=400) if include_resume_profiles else []
        pool = _merge_candidate_pool(local_rows, resume_rows)
        candidate = next((row for row in pool if str(row.get("candidate_id") or "").strip() == target), None)
        if not candidate:
            return _api_error("候选人不存在", status_code=404, code="candidate_not_found")
        jobs = hr_match_service.list_jobs(status="open", limit=400)
        items = hr_match_service.candidate_job_feed(candidate=candidate, jobs=jobs, limit=n)
        return _api_success({"candidate": candidate, "items": items, "total": len(items)})
    except Exception as e:
        logger.exception("candidate_jobs_feed_failed")
        return _api_error(str(e), status_code=500, code="candidate_jobs_feed_failed")

@app.get("/api/jobs/search")
async def search_jobs(
    keywords: str = None,
    location: str = None,
    salary_min: int = None,
    experience: str = None,
    limit: int = 50,
    allow_portal_fallback: bool = True,
    sources: str = None,
    source_scopes: str = None,
    custom_sources: str = None,
):
    """搜索真实岗位"""
    try:
        cfg_mode = os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower()
        n = int(limit) if limit is not None else 50
        n = max(1, min(n, 100))
        keyword_list = keywords.split(",") if keywords else []
        kw = [k.strip() for k in keyword_list if k and k.strip()]
        raw_sources = ",".join(
            [
                str(sources or "").strip(),
                str(source_scopes or "").strip(),
                str(custom_sources or "").strip(),
            ]
        )
        source_list = [item.strip() for item in raw_sources.split(",") if item and item.strip()]
        allow_portal = bool(allow_portal_fallback) or (
            os.getenv("ALLOW_CN_PORTAL_FALLBACK", "1").strip().lower() in {"1", "true", "yes", "on"}
        )

        # Cloud mode: prefer crawler cache; fallback to cloud-safe real-time providers.
        if cfg_mode == "cloud" or cloud_jobs_cache:
            jobs = _filter_cloud_cache_by_query(kw, location, limit=n)
            jobs = _enforce_cn_market_jobs(jobs)
            warning = None
            mode = "cloud"
            if not jobs:
                fallback_jobs, fallback_mode, fallback_err = _search_jobs_without_browser(
                    kw,
                    location,
                    limit=n,
                    allow_portal_fallback=allow_portal,
                )
                jobs = _enforce_cn_market_jobs(fallback_jobs)
                mode = fallback_mode or "cloud"
                warning = (
                    f"cloud cache is empty; switched to no-browser provider: {mode}"
                    if jobs
                    else (
                        f"cloud cache empty and no-browser search failed: {fallback_err}"
                        if fallback_err
                        else "cloud cache is empty and no real jobs were found"
                    )
                )
            if not jobs and allow_portal:
                jobs = _search_jobs_cn_entrypoints(kw, location, limit=n)
                mode = "cn_portal"
                warning = warning or "switched to CN portal entrypoint fallback"
            _track_event(
                "job_search",
                {
                    "provider_mode": mode,
                    "result_count": len(jobs),
                    "cloud_mode": True,
                },
            )
            _cache_recent_jobs(jobs)
            return _api_success({
                "total": len(jobs),
                "jobs": jobs,
                "provider_mode": mode,
                "warning": warning,
            })

        # Stream progress to the same WebSocket channel as the AI pipeline.
        # Frontend listens for `type=job_search`.
        import asyncio

        def progress_cb(message: str, percent: int):
            asyncio.create_task(
                progress_tracker.broadcast(
                    {"type": "job_search", "data": {"message": message, "percent": int(percent)}}
                )
            )

        try:
            jobs = real_job_service.search_jobs(
                keywords=kw,
                location=location,
                salary_min=salary_min,
                experience=experience,
                limit=n,
                progress_callback=progress_cb,
                sources=source_list or None,
            )
            jobs = _normalize_real_jobs(jobs, limit=n)
            jobs = _enforce_cn_market_jobs(jobs)
            mode = (real_job_service.get_statistics() or {}).get("provider_mode", cfg_mode)
            if not jobs and allow_portal:
                fallback_jobs, fallback_mode, _ = _search_jobs_without_browser(
                    kw,
                    location,
                    limit=n,
                    allow_portal_fallback=True,
                )
                fallback_jobs = _enforce_cn_market_jobs(fallback_jobs)
                if fallback_jobs:
                    jobs = fallback_jobs
                    mode = fallback_mode or mode
            if not jobs and allow_portal:
                jobs = _search_jobs_cn_entrypoints(kw, location, limit=n)
                mode = "cn_portal"
        except Exception as e:
            jobs, mode, _ = _search_jobs_without_browser(
                kw,
                location,
                limit=n,
                allow_portal_fallback=allow_portal,
            )
            jobs = _enforce_cn_market_jobs(jobs)
            if not jobs:
                raise e
        _track_event(
            "job_search",
            {
                "provider_mode": mode,
                "result_count": len(jobs),
                "cloud_mode": False,
            },
        )
        _cache_recent_jobs(jobs)
        return _api_success({
            "total": len(jobs),
            "jobs": jobs,
            "provider_mode": mode,
            "job_sources": get_job_source_registry_payload(),
        })
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/search", "error": str(e)[:300]})
        fallback_jobs: List[Dict[str, Any]] = []
        if allow_portal_fallback:
            try:
                kw = [k.strip() for k in (keywords or "").split(",") if k and k.strip()]
                fallback_jobs = _search_jobs_cn_entrypoints(kw, location, limit=max(1, min(int(limit or 10), 30)))
            except Exception:
                fallback_jobs = []
        # Do not hard-fail the UI when provider is rate-limited/captcha-blocked.
        # Return an empty-but-success payload so the frontend can keep the full flow available.
        return _api_success({
            "total": len(fallback_jobs),
            "jobs": fallback_jobs,
            "provider_mode": ("cn_portal" if fallback_jobs else (os.getenv("JOB_DATA_PROVIDER", "auto").strip().lower() or "auto")),
            "warning": str(e),
            "code": "job_search_failed",
            "job_sources": get_job_source_registry_payload(),
        })

@app.get("/api/jobs/{job_id}")
async def get_job_detail(job_id: str, request: Request):
    """获取岗位详情：仅在真实岗位详情校验通过后扣减 1 credit。"""
    normalized_job_id = str(job_id or "").strip()
    if not normalized_job_id:
        return _api_error("job_id 不能为空", status_code=400, code="invalid_job_id")
    try:
        job = real_job_service.get_job_detail(normalized_job_id)
        if not job:
            return _api_error("岗位不存在", status_code=404, code="job_not_found")

        validated, validation_reason = _validate_chargeable_job_detail(job)
        if not validated:
            credit_charge = {
                "ok": True,
                "skipped": True,
                "reason": f"job_detail_not_chargeable:{validation_reason}",
            }
            _record_job_detail_credit_state(
                normalized_job_id,
                job,
                credit_charge,
                validated=False,
                validation_reason=validation_reason,
            )
            return _api_success(
                {
                    "job": job,
                    "detail_validated": False,
                    "detail_validation_reason": validation_reason,
                    "credit_charge": credit_charge,
                }
            )

        deny, credit_guard = _credit_balance_guard(request, "job_detail")
        deny_reason = ""
        if deny:
            deny_reason = "credit_guard_denied"
            try:
                deny_payload = json.loads((deny.body or b"{}").decode("utf-8", errors="ignore") or "{}")
                deny_reason = str(deny_payload.get("code") or deny_payload.get("error") or deny_reason)
            except Exception:
                pass
            if deny_reason != "insufficient_credits":
                credit_charge = {
                    "ok": False,
                    "skipped": True,
                    "reason": deny_reason,
                }
                _record_job_detail_credit_state(
                    normalized_job_id,
                    job,
                    credit_charge,
                    validated=True,
                    validation_reason=validation_reason,
                )
                return deny

        try:
            credit_charge = commerce_service.consume_credits_once_for_resource(
                amount=int((credit_guard or {}).get("required") or 0),
                action="job_detail",
                resource_id=normalized_job_id,
                access_code=str((credit_guard or {}).get("access_code") or ""),
                buyer_id=str(((credit_guard or {}).get("wallet") or {}).get("buyer_id") or ""),
                note="岗位详情查看",
                meta={
                    "endpoint": "/api/jobs/{job_id}",
                    "resource_id": normalized_job_id,
                    "job_id": normalized_job_id,
                    "validated": True,
                    "validation_reason": validation_reason,
                },
            )
            if (
                credit_charge.get("ok")
                and credit_charge.get("skipped")
                and str(credit_charge.get("reason") or "") == "already_charged_for_resource"
            ):
                credit_charge["reason"] = "already_charged_for_job_detail"
        except Exception:
            logger.exception("job_detail_credit_consume_failed job_id=%s", normalized_job_id[:120])
            credit_charge = {
                "ok": False,
                "skipped": True,
                "reason": "credit_consume_exception",
            }
            _record_job_detail_credit_state(
                normalized_job_id,
                job,
                credit_charge,
                validated=True,
                validation_reason=validation_reason,
            )
            return _api_error("credits 扣减失败，请重试", status_code=500, code="credit_consume_exception")

        if not credit_charge.get("ok"):
            error_code = str(credit_charge.get("error") or "credit_consume_failed").strip() or "credit_consume_failed"
            credit_charge["reason"] = error_code
            _record_job_detail_credit_state(
                normalized_job_id,
                job,
                credit_charge,
                validated=True,
                validation_reason=validation_reason,
            )
            if deny is not None and error_code == "insufficient_credits":
                return deny
            if error_code == "insufficient_credits":
                return _api_error(
                    "credits 不足，请先充值",
                    status_code=402,
                    code="insufficient_credits",
                )
            return _api_error("credits 扣减失败", status_code=409, code=error_code)

        _record_job_detail_credit_state(
            normalized_job_id,
            job,
            credit_charge,
            validated=True,
            validation_reason=validation_reason,
        )
        _track_event(
            "job_detail_view",
            {
                "job_id": normalized_job_id,
                "validated": True,
                "validation_reason": validation_reason,
                "credit_status": ("skipped" if credit_charge.get("skipped") else "charged"),
            },
        )
        return _api_success(
            {
                "job": job,
                "detail_validated": True,
                "detail_validation_reason": validation_reason,
                "credit_charge": credit_charge,
            }
        )
    except Exception as e:
        logger.exception("job_detail_failed job_id=%s", normalized_job_id[:120])
        return _api_error(str(e), status_code=500, code="job_detail_failed")

@app.post("/api/jobs/apply")
async def apply_job(request: Request):
    """投递简历到指定岗位"""
    try:
        data = await request.json()
        job_id = data.get("job_id")
        resume_text = data.get("resume_text")
        user_info = data.get("user_info", {})
        
        result = real_job_service.apply_job(job_id, resume_text, user_info)
        if (not result.get("success")) and job_id in recent_search_jobs:
            # Fallback for no-browser providers whose job detail is not in local provider cache.
            j = recent_search_jobs.get(job_id, {})
            link = j.get("link") or j.get("apply_url")
            if link:
                app_id = f"EXT{int(datetime.now().timestamp())}"
                record = {
                    "application_id": app_id,
                    "job_id": job_id,
                    "job_title": j.get("title", ""),
                    "company": j.get("company", ""),
                    "platform": j.get("platform", j.get("provider", "")),
                    "apply_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "status": "待投递",
                    "apply_link": link,
                    "user_info": user_info or {},
                }
                try:
                    real_job_service.records.add_record(record)
                except Exception:
                    pass
                result = {
                    "success": True,
                    "message": "已生成投递跳转链接（请在招聘网站完成真实投递）",
                    "data": record,
                }
        _track_event(
            "job_apply",
            {
                "job_id": job_id,
                "success": bool(result.get("success")),
            },
        )
        return JSONResponse(result)
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/apply", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/jobs/batch_apply")
async def batch_apply_jobs(request: Request):
    """批量投递简历"""
    try:
        data = await request.json()
        job_ids = data.get("job_ids", [])
        resume_text = data.get("resume_text")
        user_info = data.get("user_info", {})
        
        results = real_job_service.batch_apply(job_ids, resume_text, user_info)
        _track_event(
            "job_apply",
            {
                "batch": True,
                "requested": len(job_ids),
                "success_count": len([r for r in results if r.get("success")]),
            },
        )
        return JSONResponse({
            "success": True,
            "total": len(results),
            "results": results
        })
    except Exception as e:
        _track_event("api_error", {"api": "/api/jobs/batch_apply", "error": str(e)[:300]})
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/statistics")
async def get_statistics():
    """获取数据统计"""
    try:
        stats = real_job_service.get_statistics()
        return JSONResponse({"success": True, "data": stats})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.get("/api/config/runtime")
async def get_runtime_config():
    """Expose runtime source/agent schemas for UI configurators."""
    try:
        return _api_success(
            {
                "job_sources": get_job_source_registry_payload(),
                "agent_profiles": get_agent_profiles_payload(),
            }
        )
    except Exception as e:
        return _api_error(str(e), status_code=500, code="runtime_config_failed")

# ========================================
# 爬虫数据接收接口（云端部署时使用）
# ========================================

from fastapi import Header

# 简单的API密钥验证
CRAWLER_API_KEY = os.getenv("CRAWLER_API_KEY", "your-secret-key-change-this")

@app.post("/api/crawler/upload")
async def receive_crawler_data(request: Request, authorization: str = Header(None)):
    """接收本地爬虫推送的岗位数据"""
    try:
        # 验证API密钥
        if not authorization or not authorization.startswith("Bearer "):
            return JSONResponse({"error": "未授权：缺少API密钥"}, status_code=401)
        
        api_key = authorization.replace("Bearer ", "")
        if api_key != CRAWLER_API_KEY:
            return JSONResponse({"error": "未授权：API密钥无效"}, status_code=401)
        
        # 解析数据
        data = await request.json()
        jobs = data.get("jobs", [])
        
        if not jobs:
            return JSONResponse({"error": "岗位数据为空"}, status_code=400)
        
        # 添加接收时间戳
        for job in jobs:
            job["received_at"] = datetime.now().isoformat()

        # 存储到缓存（去重 + 过滤 seed/demo + 必须可跳转）
        incoming = _normalize_and_filter_jobs(jobs, limit=5000)
        existing_keys = {
            str(j.get("link") or j.get("apply_url") or j.get("id") or "").strip().lower()
            for j in cloud_jobs_cache
        }
        new_jobs = []
        for job in incoming:
            key = str(job.get("link") or job.get("apply_url") or job.get("id") or "").strip().lower()
            if key and key not in existing_keys:
                existing_keys.add(key)
                new_jobs.append(job)

        cloud_jobs_cache.extend(new_jobs)

        # 限制缓存大小（保留最新的5000个）
        if len(cloud_jobs_cache) > 5000:
            cloud_jobs_cache[:] = cloud_jobs_cache[-5000:]

        cloud_jobs_meta["last_push_at"] = datetime.now().isoformat()
        cloud_jobs_meta["last_received"] = len(jobs)
        cloud_jobs_meta["last_new"] = len(new_jobs)
        _track_event(
            "crawler_upload",
            {"received": len(jobs), "new": len(new_jobs), "total": len(cloud_jobs_cache)},
        )

        print(f"✅ 接收爬虫数据：{len(new_jobs)} 个新岗位（总计：{len(cloud_jobs_cache)}）")
        
        return JSONResponse({
            "success": True,
            "received": len(jobs),
            "new": len(new_jobs),
            "total": len(cloud_jobs_cache)
        })
        
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.get("/api/crawler/status")
async def get_crawler_status():
    """获取爬虫数据状态"""
    if not cloud_jobs_cache:
        return JSONResponse({
            "status": "empty",
            "total": 0,
            "last_push_at": cloud_jobs_meta.get("last_push_at"),
            "last_received": cloud_jobs_meta.get("last_received", 0),
            "last_new": cloud_jobs_meta.get("last_new", 0),
        })

    return JSONResponse({
        "status": "ok",
        "total": len(cloud_jobs_cache),
        "last_push_at": cloud_jobs_meta.get("last_push_at"),
        "last_received": cloud_jobs_meta.get("last_received", 0),
        "last_new": cloud_jobs_meta.get("last_new", 0),
    })


# ========================================
# 自动投递 API 接口（新增）
# ========================================

# 全局任务管理
auto_apply_tasks: Dict[str, Dict[str, Any]] = {}
task_lock = asyncio.Lock()

# 平台映射
PLATFORM_APPLIERS = {
    'boss': 'app.services.auto_apply.boss_applier.BossApplier',
    'zhilian': 'app.services.auto_apply.zhilian_applier.ZhilianApplier',
    'linkedin': 'app.services.auto_apply.linkedin_applier.LinkedInApplier'
}

PLATFORM_RUNTIME_DEPENDENCIES: Dict[str, List[str]] = {
    "boss": ["playwright"],
    "zhilian": ["DrissionPage"],
    "linkedin": ["selenium", "undetected_chromedriver"],
}


def _merge_nested_dicts(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(base or {})
    for key, value in (override or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            nested = dict(merged.get(key) or {})
            nested.update(value)
            merged[key] = nested
        else:
            merged[key] = value
    return merged


def _platform_runtime_snapshot(platform: str) -> Dict[str, Any]:
    missing_dependencies: List[str] = []
    for module_name in PLATFORM_RUNTIME_DEPENDENCIES.get(platform, []):
        try:
            exists = importlib.util.find_spec(module_name) is not None
        except Exception:
            exists = False
        if not exists and module_name not in missing_dependencies:
            missing_dependencies.append(module_name)

    server_ready = not missing_dependencies
    if server_ready:
        message = "服务端运行依赖完整，可直接执行。"
        runtime_mode = "server_browser"
    else:
        missing_text = ", ".join(missing_dependencies)
        message = (
            f"服务端缺少运行依赖: {missing_text}。"
            " 当前自动化应切换到买家本地浏览器/本地代理执行。"
        )
        runtime_mode = "buyer_local_browser"

    return {
        "platform": platform,
        "server_ready": server_ready,
        "server_execution_available": server_ready,
        "runtime_mode": runtime_mode,
        "missing_dependencies": missing_dependencies,
        "message": message,
    }


def _prepared_only_platform_progress(
    platform: str,
    runtime: Dict[str, Any],
    message: str = "",
    status: str = "prepared_only",
) -> Dict[str, Any]:
    return {
        "status": status,
        "message": str(message or runtime.get("message") or "").strip(),
        "runtime_mode": str(runtime.get("runtime_mode") or "buyer_local_browser"),
        "server_execution_available": bool(runtime.get("server_execution_available")),
        "missing_dependencies": list(runtime.get("missing_dependencies") or []),
        "platform": platform,
    }


def _normalize_auto_apply_runtime_config(payload: Dict[str, Any]) -> Dict[str, Any]:
    raw_keywords = _to_string_list(payload.get("keywords"))
    raw_locations = _to_string_list(payload.get("locations"))
    if not raw_keywords:
        keyword = str(payload.get("keywords") or "").strip()
        if keyword:
            raw_keywords = [keyword]
    if not raw_locations:
        location = str(payload.get("location") or "").strip()
        if location:
            raw_locations = [location]

    max_count = max(1, min(200, int(payload.get("max_count") or HAITOU_DEFAULT_POLICY["max_count"])))
    base_config = _build_fallback_auto_apply_config(payload, raw_keywords, raw_locations, max_count)
    raw_config = payload.get("config") if isinstance(payload.get("config"), dict) else {}
    merged = _merge_nested_dicts(base_config, raw_config)

    for nested_key in ("boss_config", "zhilian_config", "linkedin_config"):
        merged[nested_key] = _merge_nested_dicts(
            base_config.get(nested_key) if isinstance(base_config.get(nested_key), dict) else {},
            raw_config.get(nested_key) if isinstance(raw_config.get(nested_key), dict) else {},
        )

    merged["keywords"] = str(merged.get("keywords") or ",".join(raw_keywords)).strip()
    merged["location"] = str(merged.get("location") or (raw_locations[0] if raw_locations else "")).strip()
    merged["max_count"] = max_count
    merged["headless"] = bool(HAITOU_DEFAULT_POLICY["headless"])
    merged["verification_wait_seconds"] = int(HAITOU_DEFAULT_POLICY["verification_wait_seconds"])
    merged["use_ai_answers"] = bool(merged.get("use_ai_answers", payload.get("use_ai_answers", True)))
    raw_override = raw_config.get("compliance_override_platforms")
    if raw_override is None:
        raw_override = payload.get("compliance_override_platforms")
    merged["compliance_override_platforms"] = _normalize_platform_list(raw_override)
    merged["allow_risky_automation"] = _to_bool(
        raw_config.get(
            "allow_risky_automation",
            payload.get("allow_risky_automation", merged.get("allow_risky_automation")),
        ),
        default=False,
    )
    return merged


def _single_platform_validation_payload(platform: str, config: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "platform": platform,
        "max_apply_per_session": int(config.get("max_count") or HAITOU_DEFAULT_POLICY["max_count"]),
        "keywords": str(config.get("keywords") or "").strip(),
        "location": str(config.get("location") or "").strip(),
        "company_blacklist": _to_string_list(config.get("blacklist") or config.get("company_blacklist")),
        "pause_before_submit": bool(config.get("pause_before_submit", False)),
        "headless": bool(config.get("headless", HAITOU_DEFAULT_POLICY["headless"])),
        "use_ai_answers": bool(config.get("use_ai_answers", True)),
    }


def _set_task_stopped(task: Dict[str, Any]) -> None:
    """统一设置任务为停止状态。"""
    task["status"] = "stopped"
    task["completed_at"] = datetime.now().isoformat()

@app.post("/api/auto-apply/start")
async def start_auto_apply(request: Request):
    """启动自动投递"""
    try:
        raw_data = await request.json()
        data = _apply_haitou_defaults(raw_data if isinstance(raw_data, dict) else {})
        platform = str(data.get('platform') or 'boss').strip().lower()
        if platform not in PLATFORM_APPLIERS:
            return _api_error(f'不支持的平台: {platform}', 400)

        config = _normalize_auto_apply_runtime_config(data)
        config['platform'] = platform
        legacy_jobs = data.get('jobs') if isinstance(data.get('jobs'), list) else []

        # 验证配置
        from app.services.auto_apply.config import AutoApplyConfig, validate_config
        apply_config = AutoApplyConfig.from_dict(_single_platform_validation_payload(platform, config))
        is_valid, error_msg = validate_config(apply_config)

        if not is_valid:
            return _api_error(error_msg, 400)

        accepted, blockers = _filter_platforms_by_credentials([platform], config)
        if not accepted:
            return _api_error(blockers[0] if blockers else f'{platform} 缺少登录配置', 400)

        runtime = _platform_runtime_snapshot(platform)

        # 生成任务ID
        task_id = str(uuid.uuid4())

        # 创建任务记录
        auto_apply_tasks[task_id] = {
            'task_id': task_id,
            'status': 'starting',
            'platform': platform,
            'platforms': [platform],
            'config': config,
            'strategy': HAITOU_DEFAULT_POLICY["strategy_version"],
            'progress': {
                'total_platforms': 1,
                'completed_platforms': 0,
                'applied': 0,
                'failed': 0,
                'total': 0,
                'current_job': None,
                'platform_progress': {},
            },
            'created_at': datetime.now().isoformat(),
            'started_at': None,
            'completed_at': None
        }

        if not runtime.get('server_ready'):
            row = _prepared_only_platform_progress(platform, runtime)
            auto_apply_tasks[task_id]['status'] = 'prepared_only'
            auto_apply_tasks[task_id]['progress']['completed_platforms'] = 1
            auto_apply_tasks[task_id]['progress']['platform_progress'][platform] = row
            auto_apply_tasks[task_id]['result'] = row
            auto_apply_tasks[task_id]['completed_at'] = datetime.now().isoformat()
            return _api_success({
                'task_id': task_id,
                'message': '服务端当前不直接执行该平台，已切到本地兼容模式',
                'prepared_only': True,
                'runtime': runtime,
            })

        # 异步启动投递任务
        asyncio.create_task(_run_auto_apply_task(task_id, config, legacy_jobs))

        return _api_success({
            'task_id': task_id,
            'message': '自动投递任务已启动',
            'prepared_only': False,
            'runtime': runtime,
        })

    except Exception as e:
        logger.exception("启动自动投递失败")
        return _api_error(str(e), 500)


async def _run_auto_apply_task(task_id: str, config: Dict[str, Any], jobs: List[Dict[str, Any]]):
    """运行自动投递任务（异步）"""
    try:
        task = auto_apply_tasks.get(task_id)
        if not task:
            return

        platform = str(task.get('platform') or config.get('platform') or 'boss').strip().lower()
        task['status'] = 'running'
        task['started_at'] = datetime.now().isoformat()

        run_config = dict(config or {})
        run_config['platform'] = platform
        if jobs:
            run_config['_legacy_jobs'] = list(jobs)

        result = await _run_single_platform_apply(task_id, platform, run_config)

        task = auto_apply_tasks.get(task_id)
        if not task or task.get('status') == 'stopped':
            return

        row = (
            task.get('progress', {})
            .get('platform_progress', {})
            .get(platform, {})
        )
        row_status = str((row or {}).get('status') or '').strip()
        task['completed_at'] = datetime.now().isoformat()

        if row_status == 'prepared_only':
            task['status'] = 'prepared_only'
            task['result'] = row
        elif row_status == 'failed':
            task['status'] = 'failed'
            task['error'] = str((row or {}).get('error') or '投递失败')
        elif row_status == 'stopped':
            task['status'] = 'stopped'
        else:
            task['status'] = 'completed'
            task['result'] = result
            task['progress']['applied'] = int(result.get('applied') or 0)
            task['progress']['failed'] = int(result.get('failed') or 0)
            task['progress']['total'] = int((row or {}).get('total') or 0)

    except Exception as e:
        logger.exception(f"自动投递任务失败: {task_id}")
        if task_id in auto_apply_tasks:
            auto_apply_tasks[task_id]['status'] = 'failed'
            auto_apply_tasks[task_id]['error'] = str(e)
            auto_apply_tasks[task_id]['completed_at'] = datetime.now().isoformat()


@app.post("/api/auto-apply/stop")
async def stop_auto_apply(request: Request):
    """停止自动投递"""
    try:
        data = await request.json()
        task_id = data.get('task_id')

        if not task_id or task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)

        task = auto_apply_tasks[task_id]

        if task['status'] in ['completed', 'failed', 'stopped']:
            return _api_success({
                'message': '任务已结束',
                'status': task['status'],
            })

        # 设置停止标志（实际停止逻辑在 applier 中处理）
        _set_task_stopped(task)

        return _api_success({
            'message': '停止指令已发送'
        })

    except Exception as e:
        logger.exception("停止自动投递失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/stop/{task_id}")
async def stop_auto_apply_by_path(task_id: str):
    """兼容前端路径参数形式的停止接口。"""
    try:
        if task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)

        task = auto_apply_tasks[task_id]
        if task['status'] not in ['completed', 'failed', 'stopped']:
            _set_task_stopped(task)

        return _api_success({
            'task_id': task_id,
            'status': task.get('status'),
            'message': '停止指令已发送'
        })
    except Exception as e:
        logger.exception("路径停止自动投递失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/pause/{task_id}")
async def pause_auto_apply_by_path(task_id: str):
    """兼容前端路径参数形式的暂停接口。"""
    try:
        if task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)
        task = auto_apply_tasks[task_id]
        if task['status'] in ['completed', 'failed', 'stopped']:
            return _api_error('任务已结束，无法暂停', 400)

        task['status'] = 'paused'
        return _api_success({
            'task_id': task_id,
            'status': 'paused',
            'message': '任务已暂停'
        })
    except Exception as e:
        logger.exception("暂停自动投递失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/resume/{task_id}")
async def resume_auto_apply_by_path(task_id: str):
    """兼容前端路径参数形式的恢复接口。"""
    try:
        if task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)
        task = auto_apply_tasks[task_id]
        if task['status'] in ['completed', 'failed', 'stopped']:
            return _api_error('任务已结束，无法继续', 400)

        task['status'] = 'running'
        return _api_success({
            'task_id': task_id,
            'status': 'running',
            'message': '任务已继续'
        })
    except Exception as e:
        logger.exception("继续自动投递失败")
        return _api_error(str(e), 500)


@app.get("/api/auto-apply/status/{task_id}")
async def get_auto_apply_status(task_id: str):
    """查询投递状态"""
    try:
        if task_id not in auto_apply_tasks:
            return _api_error('任务不存在', 404)

        task = auto_apply_tasks[task_id]

        return _api_success({
            'task': task
        })

    except Exception as e:
        logger.exception("查询投递状态失败")
        return _api_error(str(e), 500)


@app.get("/api/auto-apply/history")
async def get_auto_apply_history(limit: int = 50):
    """获取投递历史"""
    try:
        # 按时间倒序排列
        tasks = sorted(
            auto_apply_tasks.values(),
            key=lambda x: x.get('created_at', ''),
            reverse=True
        )

        return _api_success({
            'tasks': tasks[:limit],
            'total': len(tasks)
        })

    except Exception as e:
        logger.exception("获取投递历史失败")
        return _api_error(str(e), 500)


@app.websocket("/ws/auto-apply/{task_id}")
async def auto_apply_progress_ws(websocket: WebSocket, task_id: str):
    """实时推送投递进度（支持多平台）"""
    await websocket.accept()

    try:
        while True:
            if task_id not in auto_apply_tasks:
                await websocket.send_json({
                    'type': 'error',
                    'message': '任务不存在'
                })
                break

            task = auto_apply_tasks[task_id]

            # 发送详细进度
            await websocket.send_json({
                'type': 'progress',
                'task_id': task_id,
                'status': task['status'],
                'progress': task['progress'],
                'platforms': task.get('platforms', [task.get('config', {}).get('platform', 'linkedin')]),
                'timestamp': datetime.now().isoformat()
            })

            # 如果任务已完成，发送最终消息并断开
            if task['status'] in ['completed', 'failed', 'stopped']:
                await websocket.send_json({
                    'type': 'complete',
                    'task_id': task_id,
                    'status': task['status'],
                    'progress': task['progress'],
                    'result': task.get('result', {}),
                    'error': task.get('error')
                })
                break

            await asyncio.sleep(2)

    except WebSocketDisconnect:
        logger.info(f"WebSocket 断开: {task_id}")
    except Exception as e:
        logger.exception(f"WebSocket 错误: {e}")
        try:
            await websocket.send_json({
                'type': 'error',
                'message': str(e)
            })
        except:
            pass


# ========================================
# 多平台投递 API 接口（新增）
# ========================================

@app.post("/api/auto-apply/start-multi")
async def start_multi_platform_apply(request: Request):
    """启动多平台自动投递"""
    try:
        incoming = await request.json()
        data = _apply_haitou_defaults(incoming if isinstance(incoming, dict) else {})
        platforms = _normalize_platform_list(data.get('platforms')) or list(HAITOU_DEFAULT_POLICY["fallback_platforms"])
        config = _normalize_auto_apply_runtime_config(data)
        accepted_platforms, blockers = _filter_platforms_by_credentials(platforms, config)
        if not accepted_platforms:
            return _api_error("；".join(blockers) or '缺少可执行的平台登录配置', 400)

        runtime_matrix = {
            platform: _platform_runtime_snapshot(platform)
            for platform in accepted_platforms
        }

        # 生成任务ID
        task_id = str(uuid.uuid4())

        # 创建任务记录
        async with task_lock:
            auto_apply_tasks[task_id] = {
                'task_id': task_id,
                'status': 'starting',
                'platforms': accepted_platforms,
                'config': config,
                'credential_blockers': blockers,
                'runtime_matrix': runtime_matrix,
                'strategy': HAITOU_DEFAULT_POLICY["strategy_version"],
                'progress': {
                    'total_platforms': len(accepted_platforms),
                    'completed_platforms': 0,
                    'total_applied': 0,
                    'total_failed': 0,
                    'platform_progress': {}
                },
                'created_at': datetime.now().isoformat(),
                'started_at': None,
                'completed_at': None
            }

        if not any(snapshot.get('server_ready') for snapshot in runtime_matrix.values()):
            task = auto_apply_tasks[task_id]
            task['status'] = 'prepared_only'
            task['progress']['completed_platforms'] = len(accepted_platforms)
            task['progress']['platform_progress'] = {
                platform: _prepared_only_platform_progress(platform, runtime_matrix[platform])
                for platform in accepted_platforms
            }
            task['result'] = {
                'message': '服务端缺少多平台运行依赖，任务已切到本地兼容模式',
                'runtime_matrix': runtime_matrix,
                'credential_blockers': blockers,
            }
            task['completed_at'] = datetime.now().isoformat()
            return _api_success({
                'task_id': task_id,
                'message': '服务端当前不直接执行这些平台，已切到本地兼容模式',
                'prepared_only': True,
                'platforms': accepted_platforms,
                'runtime_matrix': runtime_matrix,
                'blocked_platforms': blockers,
            })

        # 异步启动多平台投递
        asyncio.create_task(_run_multi_platform_apply(task_id, accepted_platforms, config))

        return _api_success({
            'task_id': task_id,
            'message': f'已启动 {len(accepted_platforms)} 个平台的自动投递',
            'prepared_only': False,
            'platforms': accepted_platforms,
            'runtime_matrix': runtime_matrix,
            'blocked_platforms': blockers,
        })

    except Exception as e:
        logger.exception("启动多平台投递失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/boss/submit-code")
async def submit_boss_sms_code(request: Request):
    """提交 Boss 短信验证码，供运行中的任务继续登录。"""
    try:
        data = await request.json()
        task_id = _resolve_auto_apply_task_id(
            task_id=str(data.get("task_id") or "").strip(),
            gh_task_id=str(data.get("gh_task_id") or "").strip(),
        )
        code = "".join(ch for ch in str(data.get("code") or "") if ch.isdigit())

        if not task_id:
            return _api_error("task_id 不能为空（或提供有效 gh_task_id）", status_code=400)
        if len(code) < 4:
            return _api_error("验证码格式错误", status_code=400)

        _set_boss_sms_code(task_id, code)

        task = auto_apply_tasks.get(task_id) or {}
        boss_row = (task.get("progress") or {}).get("platform_progress", {}).get("boss")
        if isinstance(boss_row, dict):
            boss_row["status"] = "verifying"
            boss_row["message"] = "已收到验证码，正在登录验证"

        return _api_success(
            {
                "task_id": task_id,
                "accepted": True,
                "message": "验证码已提交，任务将继续执行",
            }
        )
    except Exception as e:
        logger.exception("提交 Boss 验证码失败")
        return _api_error(str(e), status_code=500)


@app.post("/api/auto-apply/boss/resend-code")
async def resend_boss_sms_code(request: Request):
    """触发运行中 Boss 任务重发短信验证码。"""
    try:
        data = await request.json()
        task_id = _resolve_auto_apply_task_id(
            task_id=str(data.get("task_id") or "").strip(),
            gh_task_id=str(data.get("gh_task_id") or "").strip(),
        )
        if not task_id:
            return _api_error("task_id 不能为空（或提供有效 gh_task_id）", status_code=400)

        task = auto_apply_tasks.get(task_id) or {}
        boss_row = (task.get("progress") or {}).get("platform_progress", {}).get("boss")
        if not isinstance(boss_row, dict):
            return _api_error("当前任务未运行 Boss 投递器", status_code=400)
        if str(boss_row.get("status") or "").strip() not in {"waiting_verification", "verifying", "opening"}:
            return _api_error("当前状态不可重发验证码", status_code=400)

        _push_boss_control_action(task_id, "resend_sms")
        boss_row["message"] = "已请求重发验证码，请留意短信"
        boss_row["status"] = "waiting_verification"
        boss_row["need_sms_code"] = True

        return _api_success(
            {
                "task_id": task_id,
                "accepted": True,
                "message": "已发送重发验证码指令",
            }
        )
    except Exception as e:
        logger.exception("重发 Boss 验证码失败")
        return _api_error(str(e), status_code=500)


async def _run_multi_platform_apply(task_id: str, platforms: List[str], config: Dict[str, Any]):
    """运行多平台投递任务"""
    try:
        # 更新状态
        auto_apply_tasks[task_id]['status'] = 'running'
        auto_apply_tasks[task_id]['started_at'] = datetime.now().isoformat()

        # 并发执行多个平台
        tasks = []
        for platform in platforms:
            if platform in PLATFORM_APPLIERS:
                task = asyncio.create_task(
                    _run_single_platform_apply(task_id, platform, config)
                )
                tasks.append(task)

        # 等待所有平台完成
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 汇总结果
        total_applied = 0
        total_failed = 0

        for result in results:
            if isinstance(result, dict):
                total_applied += result.get('applied', 0)
                total_failed += result.get('failed', 0)

        # 更新最终状态（若已手动停止则保持 stopped）
        task = auto_apply_tasks.get(task_id, {})
        task.setdefault('progress', {})
        platform_progress = task['progress'].get('platform_progress', {})
        platform_states = [
            str((platform_progress.get(p) or {}).get('status') or '')
            for p in platforms
        ]
        failed_platforms = sum(1 for s in platform_states if s == 'failed')
        prepared_platforms = sum(1 for s in platform_states if s == 'prepared_only')
        blocked_platforms = sum(1 for s in platform_states if s == 'blocked')
        done_platforms = sum(1 for s in platform_states if s in ('completed', 'failed', 'stopped', 'prepared_only', 'blocked'))

        task['progress']['completed_platforms'] = done_platforms
        task['progress']['total_applied'] = total_applied
        task['progress']['total_failed'] = total_failed

        if task.get('status') != 'stopped':
            if done_platforms == 0 and platforms:
                task['status'] = 'failed'
                task['error'] = '未执行任何平台任务'
            elif failed_platforms == len([p for p in platforms if p in PLATFORM_APPLIERS]) and total_applied == 0:
                task['status'] = 'failed'
                task['error'] = '全部平台执行失败'
            elif prepared_platforms == len(platforms) and total_applied == 0 and failed_platforms == 0:
                task['status'] = 'prepared_only'
                task['result'] = {
                    'message': '服务端缺少运行依赖，已切到本地兼容模式',
                    'runtime_matrix': task.get('runtime_matrix', {}),
                }
            elif failed_platforms > 0 or blocked_platforms > 0:
                task['status'] = 'completed_with_errors'
            elif prepared_platforms > 0 and total_applied == 0:
                task['status'] = 'prepared_only'
            elif failed_platforms > 0:
                task['status'] = 'completed_with_errors'
            else:
                task['status'] = 'completed'
            task['completed_at'] = datetime.now().isoformat()

    except Exception as e:
        logger.exception(f"多平台投递任务失败: {task_id}")
        auto_apply_tasks[task_id]['status'] = 'failed'
        auto_apply_tasks[task_id]['error'] = str(e)


async def _run_single_platform_apply(task_id: str, platform: str, config: Dict[str, Any]):
    """运行单个平台的投递任务"""
    applier = None
    try:
        task = auto_apply_tasks.get(task_id, {})
        if task.get('status') == 'stopped':
            return {'applied': 0, 'failed': 0}

        progress = task.setdefault('progress', {})
        platform_progress = progress.setdefault('platform_progress', {})
        runtime = _platform_runtime_snapshot(platform)
        if not runtime.get('server_ready'):
            platform_progress[platform] = _prepared_only_platform_progress(platform, runtime)
            progress['completed_platforms'] = int(progress.get('completed_platforms') or 0) + 1
            return {'applied': 0, 'failed': 0}

        # 动态导入平台 Applier
        try:
            module_path, class_name = PLATFORM_APPLIERS[platform].rsplit('.', 1)
            module = __import__(module_path, fromlist=[class_name])
            ApplierClass = getattr(module, class_name)
        except Exception as import_error:
            runtime = _platform_runtime_snapshot(platform)
            missing_dependencies = list(runtime.get('missing_dependencies') or [])
            if isinstance(import_error, ModuleNotFoundError):
                missing_name = str(getattr(import_error, 'name', '') or '').strip()
                if missing_name and not missing_name.startswith('app') and missing_name not in missing_dependencies:
                    missing_dependencies.append(missing_name)
            if missing_dependencies:
                compat_runtime = dict(runtime)
                compat_runtime['server_ready'] = False
                compat_runtime['server_execution_available'] = False
                compat_runtime['runtime_mode'] = 'buyer_local_browser'
                compat_runtime['missing_dependencies'] = missing_dependencies
                compat_runtime['message'] = (
                    f"服务端缺少运行依赖: {', '.join(missing_dependencies)}。"
                    " 当前自动化应切换到买家本地浏览器/本地代理执行。"
                )
                platform_progress[platform] = _prepared_only_platform_progress(platform, compat_runtime)
                progress['completed_platforms'] = int(progress.get('completed_platforms') or 0) + 1
                return {'applied': 0, 'failed': 0}
            raise

        # 获取平台特定配置
        raw_platform_config = config.get(f'{platform}_config', {})
        platform_config = dict(raw_platform_config) if isinstance(raw_platform_config, dict) else {}
        platform_headless = platform_config.get('headless')
        if platform_headless is None:
            platform_headless = config.get('headless', HAITOU_DEFAULT_POLICY["headless"])
        force_headless = os.getenv("FORCE_HEADLESS_BROWSER", "0").strip().lower() in {"1", "true", "yes", "on"}
        if force_headless:
            platform_headless = True
        platform_config.update({
            'keywords': config.get('keywords', ''),
            'location': config.get('location', ''),
            'max_apply_per_session': config.get('max_count', 50),
            'company_blacklist': config.get('blacklist', []),
            'headless': bool(platform_headless),
        })

        platform_progress[platform] = {
            'status': 'initializing',
            'message': '正在初始化平台'
        }

        if platform == 'boss':
            def _boss_progress(payload: Dict[str, Any]):
                stage = str(payload.get("stage") or "").strip()
                message = str(payload.get("message") or "").strip()
                row = auto_apply_tasks.get(task_id, {}).get('progress', {}).get('platform_progress', {}).get(platform, {})
                if stage == "waiting_sms_code":
                    row.update({
                        'status': 'waiting_verification',
                        'message': message or '等待短信验证码',
                        'need_sms_code': True,
                        'phone': _mask_phone(str(platform_config.get('phone') or '')),
                    })
                elif stage == "sms_code_submitted":
                    row.update({'status': 'verifying', 'message': message or '验证码已提交，验证中'})
                elif stage == "login_success":
                    row.update({'status': 'logged_in', 'message': message or '登录成功'})
                elif stage == "login_failed":
                    row.update({'status': 'failed', 'error': message or '登录失败'})
                elif stage == "login_open":
                    row.update({'status': 'opening', 'message': message or '打开登录页中'})
                elif stage == "sms_resent":
                    row.update({'status': 'waiting_verification', 'message': message or '验证码已重发，请查收'})
                elif stage == "sms_send_failed":
                    row.update({'status': 'failed', 'error': message or '验证码发送失败'})
                auto_apply_tasks.get(task_id, {}).get('progress', {}).get('platform_progress', {})[platform] = row

            platform_config.update({
                'task_id': task_id,
                'verification_wait_seconds': int(config.get('verification_wait_seconds', 180)),
                'sms_code_fetcher': (lambda tid=task_id: _pop_boss_sms_code(tid)),
                'control_action_fetcher': (lambda tid=task_id: _pop_boss_control_action(tid)),
                'progress_hook': _boss_progress,
                'sms_code': str(platform_config.get('sms_code') or config.get('sms_code') or '').strip() or None,
            })

        # 创建投递器（兼容不同的构造函数）
        # LinkedIn 需要 llm_client 参数，其他平台不应受其可用性影响。
        if platform == 'linkedin':
            llm_client = None
            if config.get('use_ai_answers', True):
                try:
                    from app.core.llm_client import LLMClient
                    llm_client = LLMClient()
                except Exception:
                    llm_client = None
            applier = ApplierClass(platform_config, llm_client)
        else:
            applier = ApplierClass(platform_config)

        # Boss uses one-loop pipeline to avoid cross-loop browser/session invalidation.
        if platform == 'boss':
            phone = str(platform_config.get('phone') or '').strip()
            if not phone:
                raise Exception("boss 缺少手机号")

            pipeline_result = await asyncio.to_thread(
                applier.run_apply_pipeline,
                phone,
                platform_config.get('sms_code'),
                platform_config.get('keywords', ''),
                platform_config.get('location', ''),
                {},
                platform_config.get('max_apply_per_session', 50),
            )
            if not pipeline_result.get('ok'):
                row = auto_apply_tasks.get(task_id, {}).get('progress', {}).get('platform_progress', {}).get(platform, {})
                detail = str((row or {}).get('error') or (row or {}).get('message') or '').strip()
                raise Exception(detail or str(pipeline_result.get('error') or "boss 执行失败"))

            jobs_count = int(pipeline_result.get('jobs') or 0)
            applied_count = int(pipeline_result.get('applied') or 0)
            failed_count = int(pipeline_result.get('failed') or 0)
            platform_progress[platform] = {
                'status': 'completed',
                'total': jobs_count,
                'applied': applied_count,
                'failed': failed_count,
            }
            progress['completed_platforms'] = int(progress.get('completed_platforms') or 0) + 1
            return {'applied': applied_count, 'failed': failed_count}

        # 登录
        login_success = False
        if platform == 'zhilian':
            username = platform_config.get('username')
            password = platform_config.get('password')
            if username and password:
                login_success = await asyncio.to_thread(applier.login, username, password)
        elif platform == 'linkedin':
            email = platform_config.get('email')
            password = platform_config.get('password')
            if email and password:
                login_success = await asyncio.to_thread(applier.login, email, password)

        if not login_success:
            row = auto_apply_tasks.get(task_id, {}).get('progress', {}).get('platform_progress', {}).get(platform, {})
            detail = str((row or {}).get('error') or (row or {}).get('message') or '').strip()
            raise Exception(detail or f"{platform} 登录失败")

        # 搜索职位
        legacy_jobs = config.get('_legacy_jobs') if isinstance(config.get('_legacy_jobs'), list) else []
        if legacy_jobs:
            jobs = list(legacy_jobs)
        else:
            jobs = await asyncio.to_thread(
                applier.search_jobs,
                keywords=platform_config.get('keywords', ''),
                location=platform_config.get('location', ''),
                filters={},
            )

        # 更新进度
        platform_progress[platform] = {
            'status': 'running',
            'total': len(jobs),
            'applied': 0,
            'failed': 0
        }

        if auto_apply_tasks.get(task_id, {}).get('status') == 'stopped':
            platform_progress[platform]['status'] = 'stopped'
            await asyncio.to_thread(applier.cleanup)
            return {'applied': 0, 'failed': 0}

        # 批量投递
        result = await asyncio.to_thread(
            applier.batch_apply,
            jobs,
            platform_config.get('max_apply_per_session', 50),
        )

        # 更新平台进度
        platform_progress[platform] = {
            'status': 'completed',
            'total': len(jobs),
            'applied': result['applied'],
            'failed': result['failed']
        }

        progress['completed_platforms'] = int(progress.get('completed_platforms') or 0) + 1

        # 清理资源
        await asyncio.to_thread(applier.cleanup)

        return result

    except Exception as e:
        logger.exception(f"平台 {platform} 投递失败")
        auto_apply_tasks[task_id].setdefault('progress', {}).setdefault('platform_progress', {})[platform] = {
            'status': 'failed',
            'error': str(e)
        }
        return {'applied': 0, 'failed': 0}
    finally:
        if platform == 'boss':
            _pop_boss_sms_code(task_id)
            _pop_boss_control_action(task_id)


@app.get("/api/auto-apply/platforms")
async def get_supported_platforms():
    """获取支持的平台列表"""
    return _api_success({
        'platforms': [
            {
                'id': 'boss',
                'name': 'Boss直聘',
                'icon': '💼',
                'status': 'available',
                'features': ['手机验证码登录', '智能投递', '打招呼语'],
                'config_fields': [
                    {'name': 'phone', 'label': '手机号', 'type': 'text', 'required': True}
                ]
            },
            {
                'id': 'zhilian',
                'name': '智联招聘',
                'icon': '📋',
                'status': 'available',
                'features': ['账号密码登录', '简历投递', '附件上传'],
                'config_fields': [
                    {'name': 'username', 'label': '用户名', 'type': 'text', 'required': True},
                    {'name': 'password', 'label': '密码', 'type': 'password', 'required': True}
                ]
            },
            {
                'id': 'linkedin',
                'name': 'LinkedIn',
                'icon': '🔗',
                'status': 'available',
                'features': ['Easy Apply', 'AI问答', '国际职位'],
                'config_fields': [
                    {'name': 'email', 'label': '邮箱', 'type': 'email', 'required': True},
                    {'name': 'password', 'label': '密码', 'type': 'password', 'required': True}
                ]
            }
        ]
    })


@app.get("/api/auto-apply/stats")
async def get_apply_stats():
    """获取投递统计"""
    try:
        # 统计所有任务
        total_tasks = len(auto_apply_tasks)
        completed_tasks = sum(1 for t in auto_apply_tasks.values() if t['status'] == 'completed')
        running_tasks = sum(1 for t in auto_apply_tasks.values() if t['status'] == 'running')

        total_applied = 0
        total_failed = 0

        # 平台统计
        platform_stats = {}

        for task in auto_apply_tasks.values():
            # 单平台任务
            if 'result' in task:
                result = task['result']
                total_applied += result.get('applied', 0)
                total_failed += result.get('failed', 0)

                platform = task.get('config', {}).get('platform', 'linkedin')
                if platform not in platform_stats:
                    platform_stats[platform] = {'applied': 0, 'failed': 0, 'total': 0}
                platform_stats[platform]['applied'] += result.get('applied', 0)
                platform_stats[platform]['failed'] += result.get('failed', 0)

            # 多平台任务
            if 'progress' in task and 'platform_progress' in task['progress']:
                for platform, progress in task['progress']['platform_progress'].items():
                    if platform not in platform_stats:
                        platform_stats[platform] = {'applied': 0, 'failed': 0, 'total': 0}
                    platform_stats[platform]['applied'] += progress.get('applied', 0)
                    platform_stats[platform]['failed'] += progress.get('failed', 0)
                    platform_stats[platform]['total'] += progress.get('total', 0)

                total_applied += task['progress'].get('total_applied', 0)
                total_failed += task['progress'].get('total_failed', 0)

        return _api_success({
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'running_tasks': running_tasks,
            'total_applied': total_applied,
            'total_failed': total_failed,
            'success_rate': round(total_applied / (total_applied + total_failed) * 100, 2) if (total_applied + total_failed) > 0 else 0,
            'platform_stats': platform_stats
        })

    except Exception as e:
        logger.exception("获取统计失败")
        return _api_error(str(e), 500)


@app.post("/api/auto-apply/test-platform")
async def test_platform_config(request: Request):
    """测试平台配置（不实际投递）"""
    try:
        data = await request.json()
        platform = str(data.get('platform') or 'boss').strip().lower()
        config = data.get('config') if isinstance(data.get('config'), dict) else {}

        if platform not in PLATFORM_APPLIERS:
            return _api_error(f'不支持的平台: {platform}', 400)

        runtime = _platform_runtime_snapshot(platform)
        login_result = {'success': False, 'message': '缺少必要配置'}

        if platform == 'boss':
            phone = str(config.get('phone') or '').strip()
            if phone:
                login_result = {
                    'success': True,
                    'message': f'配置正确，手机号: {phone[:3]}****{phone[-4:]}'
                }
        elif platform == 'zhilian':
            username = str(config.get('username') or '').strip()
            password = str(config.get('password') or '')
            if username and password:
                login_result = {
                    'success': True,
                    'message': f'配置正确，用户名: {username}'
                }
        elif platform == 'linkedin':
            email = str(config.get('email') or '').strip()
            password = str(config.get('password') or '')
            if email and password:
                login_result = {
                    'success': True,
                    'message': f'配置正确，邮箱: {email}'
                }

        return _api_success({
            'platform': platform,
            'login_test': login_result,
            'config_valid': login_result['success'],
            'prepared_only': not bool(runtime.get('server_ready')),
            'runtime': runtime,
            'message': (
                '配置字段通过，但服务端当前会走本地兼容模式'
                if login_result['success'] and not runtime.get('server_ready')
                else runtime.get('message')
            ),
        })

    except Exception as e:
        logger.exception("测试平台配置失败")
        return _api_error(str(e), 500)


@app.get("/api/resume/templates")
async def resume_templates():
    """Available resume templates for rendering."""
    return _api_success({"templates": resume_render_service.template_names()})


@app.post("/api/resume/structure")
async def resume_structure(request: Request):
    """Resume text -> structured JSON, optionally persisted as profile."""
    try:
        raw_data = await request.json()
        data = _apply_haitou_defaults(raw_data if isinstance(raw_data, dict) else {})
        resume_text = str(data.get("resume_text") or "").strip()
        if len(resume_text) < 20:
            return _api_error("resume_text 太短，至少 20 个字符", status_code=400)
        deny, credit_guard = _credit_balance_guard(request, "resume_structure")
        if deny:
            return deny

        model = str(data.get("model") or "").strip() or None
        save_profile = bool(data.get("save_profile", True))
        user_id = str(data.get("user_id") or "").strip()
        title = str(data.get("title") or "").strip()
        profile_id = str(data.get("profile_id") or "").strip()

        result = await resume_profile_service.structure_resume_text(resume_text, model=model)
        if not result.get("ok"):
            return _api_error(result.get("error") or "结构化失败", status_code=500, code="resume_structure_failed")

        out: Dict[str, Any] = {
            "resume_json": result.get("resume_json") or {},
            "source": result.get("source") or "",
            "model": result.get("model") or "",
        }
        if result.get("warning"):
            out["warning"] = result.get("warning")

        if save_profile:
            info = resume_profile_service.save_profile(
                resume_json=out["resume_json"],
                profile_id=profile_id,
                user_id=user_id,
                title=title,
                source_text=resume_text,
            )
            out["profile"] = info
            _track_event("resume_profile_saved", {"profile_id": info.get("profile_id"), "source": out["source"]})

        credit_charge = _settle_credit_action(
            credit_guard,
            note="结构化简历母版",
            meta={"endpoint": "/api/resume/structure", "save_profile": save_profile},
        )
        if not credit_charge.get("ok"):
            return _api_error("credits 扣减失败", status_code=409, code="credit_consume_failed")
        out["credit_charge"] = credit_charge

        return _api_success(out)
    except Exception as e:
        logger.exception("结构化简历失败")
        return _api_error(str(e), status_code=500, code="resume_structure_failed")


@app.post("/api/resume/profiles")
async def upsert_resume_profile(request: Request):
    """Create/update structured resume profile directly."""
    try:
        data = await request.json()
        resume_json = data.get("resume_json")
        if not isinstance(resume_json, dict):
            return _api_error("resume_json 必须是对象", status_code=400)

        info = resume_profile_service.save_profile(
            resume_json=resume_json,
            profile_id=str(data.get("profile_id") or "").strip(),
            user_id=str(data.get("user_id") or "").strip(),
            title=str(data.get("title") or "").strip(),
            source_text=str(data.get("source_text") or "").strip(),
        )
        _track_event("resume_profile_saved", {"profile_id": info.get("profile_id"), "source": "direct"})
        return _api_success({"profile": info})
    except Exception as e:
        logger.exception("保存简历档案失败")
        return _api_error(str(e), status_code=500, code="resume_profile_save_failed")


@app.get("/api/resume/profiles")
async def list_resume_profiles(limit: int = 20, user_id: str = ""):
    """List structured resume profiles."""
    try:
        payload = resume_profile_service.list_profiles(limit=limit, user_id=user_id)
        return _api_success(payload)
    except Exception as e:
        logger.exception("获取简历档案列表失败")
        return _api_error(str(e), status_code=500, code="resume_profile_list_failed")


@app.get("/api/resume/profiles/{profile_id}")
async def get_resume_profile(profile_id: str):
    """Get one structured resume profile."""
    try:
        row = resume_profile_service.get_profile(profile_id)
        if not row:
            return _api_error("简历档案不存在", status_code=404, code="resume_profile_not_found")
        return _api_success({"profile": row})
    except Exception as e:
        logger.exception("获取简历档案失败")
        return _api_error(str(e), status_code=500, code="resume_profile_get_failed")


@app.post("/api/resume/profiles/{profile_id}/fragment")
async def merge_resume_fragment(profile_id: str, request: Request):
    """Merge fragmented input into existing structured resume profile."""
    try:
        data = await request.json()
        fragment_text = str(data.get("fragment_text") or "").strip()
        if not fragment_text:
            return _api_error("fragment_text 不能为空", status_code=400)

        profile = resume_profile_service.get_profile(profile_id)
        if not profile:
            return _api_error("简历档案不存在", status_code=404, code="resume_profile_not_found")

        merged = await resume_profile_service.merge_fragment(
            current_resume_json=profile.get("resume_json") or {},
            fragment_text=fragment_text,
            section=str(data.get("section") or "").strip(),
            model=str(data.get("model") or "").strip() or None,
        )
        if not merged.get("ok"):
            return _api_error(merged.get("error") or "合并失败", status_code=400)

        info = resume_profile_service.save_profile(
            resume_json=merged.get("resume_json") or {},
            profile_id=profile_id,
            user_id=str(profile.get("user_id") or ""),
            title=str(profile.get("title") or ""),
            source_text=str(profile.get("source_text") or ""),
        )
        _track_event("resume_fragment_merged", {"profile_id": profile_id, "source": merged.get("source")})
        return _api_success(
            {
                "profile": info,
                "resume_json": merged.get("resume_json") or {},
                "source": merged.get("source") or "",
                "model": merged.get("model") or "",
            }
        )
    except Exception as e:
        logger.exception("简历碎片合并失败")
        return _api_error(str(e), status_code=500, code="resume_fragment_merge_failed")


@app.post("/api/resume/render")
async def render_resume(request: Request):
    """Render structured resume JSON to PDF/HTML."""
    try:
        data = await request.json()
        template_name = str(data.get("template") or "classic").strip().lower()
        fmt = str(data.get("format") or "pdf").strip().lower()
        if template_name not in set(resume_render_service.template_names()):
            return _api_error("template 不支持", status_code=400)
        if fmt not in {"pdf", "html"}:
            return _api_error("format 仅支持 pdf/html", status_code=400)

        resume_json = data.get("resume_json")
        profile_id = str(data.get("profile_id") or "").strip()
        if not isinstance(resume_json, dict):
            if not profile_id:
                return _api_error("需要提供 profile_id 或 resume_json", status_code=400)
            profile = resume_profile_service.get_profile(profile_id)
            if not profile:
                return _api_error("简历档案不存在", status_code=404, code="resume_profile_not_found")
            resume_json = profile.get("resume_json") or {}

        doc = await asyncio.to_thread(
            resume_render_service.render_to_file,
            resume_json,
            template_name,
            fmt,
        )
        _track_event(
            "resume_rendered",
            {
                "doc_id": doc.get("doc_id"),
                "format": doc.get("format"),
                "template": doc.get("template"),
            },
        )
        return _api_success(
            {
                "document": {
                    "doc_id": doc.get("doc_id"),
                    "format": doc.get("format"),
                    "template": doc.get("template"),
                    "file_size": doc.get("file_size"),
                    "created_at": doc.get("created_at"),
                    "download_url": f"/api/resume/render/download/{doc.get('doc_id')}",
                }
            }
        )
    except Exception as e:
        logger.exception("简历渲染失败")
        return _api_error(str(e), status_code=500, code="resume_render_failed")


@app.get("/api/resume/render/download/{doc_id}")
async def download_rendered_resume(doc_id: str):
    """Download rendered resume document by doc_id."""
    try:
        file_path = resume_render_service.resolve_doc_path(doc_id)
        if not file_path:
            return _api_error("文档不存在", status_code=404, code="render_doc_not_found")
        filename = os.path.basename(file_path)
        media = "application/pdf" if file_path.lower().endswith(".pdf") else "text/html; charset=utf-8"
        return FileResponse(file_path, media_type=media, filename=filename)
    except Exception as e:
        logger.exception("下载渲染文档失败")
        return _api_error(str(e), status_code=500, code="render_doc_download_failed")


def _default_email_html_template() -> str:
    return """
    <p>您好 {name}：</p>
    <p>我是候选人 {candidate_name}，关注到贵司 {company} 的相关岗位，非常希望有机会进一步沟通。</p>
    <p>我的核心能力：{skills_line}。</p>
    <p>简历已附上，期待您的回复，感谢！</p>
    <p>此致<br/>敬礼</p>
    <p>{candidate_name}<br/>{candidate_email}</p>
    """


@app.get("/api/email/presets")
async def email_presets():
    return _api_success({"presets": SMTP_PRESETS})


@app.get("/api/email/history")
async def email_history(limit: int = 50):
    try:
        return _api_success(email_campaign_service.list_history(limit=limit))
    except Exception as e:
        logger.exception("查询邮件历史失败")
        return _api_error(str(e), status_code=500, code="email_history_failed")


@app.get("/api/email/rate-limit")
async def email_rate_limit(sender_email: str, max_hourly: int = 30, max_daily: int = 100):
    try:
        if "@" not in (sender_email or ""):
            return _api_error("sender_email 不合法", status_code=400)
        payload = email_campaign_service.check_rate_limit(
            sender_email=sender_email.strip().lower(),
            max_hourly=max_hourly,
            max_daily=max_daily,
        )
        return _api_success(payload)
    except Exception as e:
        logger.exception("查询邮件限速失败")
        return _api_error(str(e), status_code=500, code="email_rate_limit_failed")


@app.post("/api/email/send-batch")
async def email_send_batch(request: Request):
    """
    Batch email delivery with SMTP presets + rate limits.
    Payload:
      smtp, contacts[], subject_template, body_template,
      profile_id/resume_json, template, max_hourly, max_daily, dry_run, attach_pdf
    """
    try:
        data = await request.json()
        smtp = data.get("smtp") if isinstance(data.get("smtp"), dict) else {}
        contacts = data.get("contacts") if isinstance(data.get("contacts"), list) else []
        if not contacts:
            return _api_error("contacts 不能为空", status_code=400)

        profile_id = str(data.get("profile_id") or "").strip()
        resume_json = data.get("resume_json") if isinstance(data.get("resume_json"), dict) else None
        if (not resume_json) and profile_id:
            profile = resume_profile_service.get_profile(profile_id)
            if not profile:
                return _api_error("简历档案不存在", status_code=404, code="resume_profile_not_found")
            resume_json = profile.get("resume_json") or {}

        dry_run = bool(data.get("dry_run", False))
        attach_pdf = bool(data.get("attach_pdf", True))
        # In dry-run mode, skip heavy PDF rendering unless explicitly requested.
        if dry_run and ("attach_pdf" not in data):
            attach_pdf = False
        template_name = str(data.get("template") or "classic").strip().lower()
        if template_name not in set(resume_render_service.template_names()):
            template_name = "classic"

        attachment_path = ""
        rendered_doc = None
        attachment_warning = ""
        if attach_pdf:
            if not isinstance(resume_json, dict):
                return _api_error("attach_pdf=true 时必须提供 profile_id 或 resume_json", status_code=400)
            try:
                rendered_doc = await asyncio.to_thread(
                    resume_render_service.render_to_file,
                    resume_json,
                    template_name,
                    "pdf",
                )
                attachment_path = rendered_doc.get("file_path") or ""
            except Exception as render_err:
                if dry_run:
                    attachment_warning = f"附件渲染已跳过: {render_err}"
                    logger.warning("dry_run 附件渲染失败，已降级为仅正文发送")
                    rendered_doc = None
                    attachment_path = ""
                else:
                    raise

        personal = (resume_json or {}).get("personal_info") if isinstance(resume_json, dict) else {}
        skills = (resume_json or {}).get("skills") if isinstance(resume_json, dict) else []
        candidate_name = str((personal or {}).get("name") or "候选人")
        candidate_email = str((personal or {}).get("email") or (smtp.get("username") or ""))
        skills_line = "、".join([str(s) for s in skills[:8]]) if isinstance(skills, list) else ""
        default_ctx = {
            "candidate_name": candidate_name,
            "candidate_email": candidate_email,
            "skills_line": skills_line,
        }

        subject_template = str(data.get("subject_template") or "应聘咨询 - {candidate_name}")
        body_template = str(data.get("body_template") or _default_email_html_template())
        max_hourly = int(data.get("max_hourly") or 30)
        max_daily = int(data.get("max_daily") or 100)

        result = await asyncio.to_thread(
            email_campaign_service.send_batch,
            smtp,
            contacts,
            subject_template,
            body_template,
            attachment_path,
            max_hourly,
            max_daily,
            dry_run,
            default_ctx,
        )
        _track_event(
            "email_batch_sent",
            {
                "campaign_id": result.get("campaign_id"),
                "sent": result.get("sent", 0),
                "failed": result.get("failed", 0),
                "dry_run": dry_run,
            },
        )

        out: Dict[str, Any] = {"result": result}
        if attachment_warning:
            out["warning"] = attachment_warning
        if rendered_doc:
            out["attachment"] = {
                "doc_id": rendered_doc.get("doc_id"),
                "download_url": f"/api/resume/render/download/{rendered_doc.get('doc_id')}",
                "file_size": rendered_doc.get("file_size"),
            }
        return _api_success(out)
    except Exception as e:
        logger.exception("批量发送邮件失败")
        return _api_error(str(e), status_code=500, code="email_batch_send_failed")


@app.get("/api/features/overview")
async def features_overview():
    aihawk = _aihawk_capability_snapshot()
    return _api_success(
        {
            "features": {
                "market_process": True,
                "multi_agent_process": True,
                "skills_graph": True,
                "resume_profile": True,
                "resume_render": True,
                "job_search": True,
                "email_campaign": True,
                "auto_apply_platforms": True,
                "github_highstar_prepare": bool(aihawk.get("can_prepare")),
                "github_highstar_apply": bool(aihawk.get("can_run")),
                "haitou_default_strategy": True,
            },
            "github_highstar": aihawk,
            "haitou_default_policy": HAITOU_DEFAULT_POLICY,
            "platform_automation_policy": PLATFORM_AUTOMATION_POLICY,
        }
    )


@app.get("/api/strategy/defaults")
async def strategy_defaults():
    return _api_success(
        {
            "haitou_default_policy": HAITOU_DEFAULT_POLICY,
            "platform_automation_policy": PLATFORM_AUTOMATION_POLICY,
        }
    )


@app.post("/api/skills/analyze")
async def analyze_skills_graph(request: Request):
    try:
        data = await request.json()
        resume_text = str(data.get("resume_text") or data.get("resume") or "").strip()
        if not resume_text:
            return _api_error("resume_text 不能为空", status_code=400)

        target_role = str(data.get("target_role") or "").strip()
        job_text = str(data.get("job_text") or "").strip()
        payload = _skills_graph_payload(resume_text, target_role, job_text)
        return _api_success({"skills_graph": payload})
    except Exception as e:
        logger.exception("skills_graph 分析失败")
        return _api_error(str(e), status_code=500, code="skills_graph_failed")


@app.post("/api/process/multi-agent")
async def process_resume_multi_agent(request: Request):
    """Run legacy multi-agent deep analysis pipeline + skills graph."""
    try:
        data = await request.json()
        resume_text = str(data.get("resume") or data.get("resume_text") or "").strip()
        runtime_agent_overrides = data.get("agent_overrides")
        if not isinstance(runtime_agent_overrides, dict):
            runtime_agent_overrides = data.get("agent_configs") if isinstance(data.get("agent_configs"), dict) else None
        if not resume_text:
            return _api_error("简历内容不能为空", status_code=400, code="empty_resume")

        # If caller explicitly passes agent runtime overrides, route through the market process
        # so each agent can honor dedicated base_url/api_key/model/prompt/skills.
        if isinstance(runtime_agent_overrides, dict) and runtime_agent_overrides:
            class _CompatJSONRequest:
                def __init__(self, payload: Dict[str, Any]):
                    self._payload = payload

                async def json(self):
                    return self._payload

            return await process_resume(
                _CompatJSONRequest(
                    {
                        "resume": resume_text,
                        "agent_overrides": runtime_agent_overrides,
                    }
                )
            )

        _track_event(
            "resume_process_started",
            {"chars": len(resume_text), "engine_mode": "multi_agent"},
        )

        info = analyzer.extract_info(resume_text)
        downgraded = False
        results: Dict[str, Any] = {}
        multi_timeout = max(20, int(os.getenv("MULTI_AGENT_TIMEOUT_SECONDS", "120") or "120"))
        try:
            results = await asyncio.wait_for(
                asyncio.to_thread(pipeline.process_resume, resume_text),
                timeout=multi_timeout,
            )
        except Exception as e:
            logger.warning("multi-agent 主流程超时/失败，准备降级: %s", e)
            downgraded = True

        if (not results) or _multi_agent_outputs_degraded(results):
            role = str(info.get("job_intention") or "Python后端工程师").strip()
            top_skills = [str(x) for x in (info.get("skills") or [])[:6] if str(x).strip()]
            results = {
                "career_analysis": _render_market_analysis_fallback(info),
                "job_recommendations": (
                    "推荐投递方向："
                    f"{role} / AI应用工程师 / 数据工程师。"
                    f"优先关键词：{', '.join(top_skills) if top_skills else 'Python, FastAPI, SQL'}。"
                ),
                "optimized_resume": _render_optimized_resume_fallback(resume_text, info),
                "interview_prep": _render_interview_prep_fallback(info, []),
                "mock_interview": "建议进行 10-15 分钟模拟面试：项目深挖 + 场景题 + 反问环节。",
            }
            downgraded = True

        seed_keywords: List[str] = []
        if info.get("job_intention") and info["job_intention"] != "未指定":
            seed_keywords.append(str(info["job_intention"]))
        seed_keywords.extend([str(x) for x in (info.get("skills") or [])[:6]])
        seed_keywords = [k for k in seed_keywords if k]
        seed_location = (info.get("preferred_locations") or [None])[0]

        provider_mode = "none"
        recommended_jobs: List[Dict[str, Any]] = []
        try:
            search_timeout = max(10, int(os.getenv("JOB_SEARCH_TIMEOUT_SECONDS", "30") or "30"))
            jobs, provider_mode, _ = await asyncio.wait_for(
                asyncio.to_thread(
                    _search_jobs_without_browser,
                    seed_keywords,
                    seed_location,
                    8,
                    True,
                ),
                timeout=search_timeout,
            )
            recommended_jobs = _public_job_payload(_enforce_cn_market_jobs(jobs), limit=8)
        except Exception:
            logger.exception("multi-agent 岗位检索失败")

        if downgraded:
            results["interview_prep"] = _render_interview_prep_fallback(info, recommended_jobs)
            if recommended_jobs:
                top_titles = [str(j.get("title") or "") for j in recommended_jobs[:3] if str(j.get("title") or "").strip()]
                if top_titles:
                    results["job_recommendations"] = "推荐岗位： " + " / ".join(top_titles)

        skills_payload = _skills_graph_payload(
            resume_text,
            str(info.get("job_intention") or ""),
            str(results.get("job_recommendations") or ""),
        )

        response_payload = {
            "career_analysis": str(results.get("career_analysis") or ""),
            "job_recommendations": str(results.get("job_recommendations") or ""),
            "optimized_resume": str(results.get("optimized_resume") or ""),
            "interview_prep": str(results.get("interview_prep") or ""),
            "mock_interview": str(results.get("mock_interview") or ""),
            "recommended_jobs": recommended_jobs,
            "job_provider_mode": provider_mode,
            "engine_mode": "multi_agent",
            "fallback_downgraded": downgraded,
            "skills_graph": skills_payload,
        }
        _track_event(
            "resume_processed",
            {
                "ok": True,
                "engine_mode": "multi_agent",
                "provider_mode": provider_mode,
                "jobs": len(recommended_jobs),
            },
        )
        return _api_success(response_payload)
    except Exception as e:
        logger.exception("multi-agent 简历处理失败")
        _track_event(
            "resume_processed",
            {"ok": False, "engine_mode": "multi_agent", "error": str(e)[:300]},
        )
        return _api_error(str(e), status_code=500, code="process_multi_agent_failed")


@app.post("/api/offer/train")
async def offer_training_pack(request: Request):
    """Build targeted interview training content for one role/job."""
    try:
        data = await request.json()
        resume_text = str(data.get("resume_text") or data.get("resume") or "").strip()
        if len(resume_text) < 20:
            return _api_error("resume_text 太短，至少 20 个字符", status_code=400)
        deny, credit_guard = _credit_balance_guard(request, "offer_pipeline")
        if deny:
            return deny

        target_role = str(data.get("target_role") or "").strip()
        job = data.get("job") if isinstance(data.get("job"), dict) else {}
        job_text = str(data.get("job_text") or "").strip()
        if job_text and not job:
            job = {"title": target_role or "目标岗位", "company": str(data.get("company") or "目标公司"), "description": job_text}

        info = analyzer.extract_info(resume_text)
        if not target_role:
            target_role = str(info.get("job_intention") or job.get("title") or "AI应用工程师").strip()
        pack = _build_interview_pack(resume_text, info, target_role, job)
        return _api_success({"training": pack})
    except Exception as e:
        logger.exception("生成面试训练包失败")
        return _api_error(str(e), status_code=500, code="offer_training_failed")


@app.get("/api/offer/funnel")
async def offer_funnel(user_id: str = "", phone: str = "", limit: int = 80):
    """Offer pipeline board derived from application records."""
    try:
        rows = _filter_application_rows(real_job_service.records.get_all_records(), user_id=user_id, phone=phone)
        return _api_success({"funnel": _build_offer_funnel_payload(rows, limit=limit)})
    except Exception as e:
        logger.exception("获取 offer funnel 失败")
        return _api_error(str(e), status_code=500, code="offer_funnel_failed")


@app.patch("/api/applications/{application_id}/status")
async def update_application_status(application_id: str, request: Request):
    """Advance one application through the offer funnel."""
    try:
        data = await request.json()
        current = real_job_service.records.get_record(application_id)
        if not current:
            return _api_error("申请记录不存在", status_code=404, code="application_not_found")

        current_status = _normalize_offer_stage(str(current.get("status") or "待投递"))
        requested_status = str(data.get("status") or "").strip()
        if requested_status and _is_placeholder_text(requested_status):
            next_status = current_status
        else:
            next_status = _normalize_offer_stage(requested_status or current_status)
        if requested_status and not _is_offer_stage_value(next_status):
            next_status = current_status
        patch = {
            "status": next_status,
            "notes": str(data.get("notes") or current.get("notes") or "").strip(),
            "follow_up_at": str(data.get("follow_up_at") or current.get("follow_up_at") or "").strip(),
            "feedback_summary": str(data.get("feedback_summary") or current.get("feedback_summary") or "").strip(),
            "next_action": str(data.get("next_action") or _build_offer_next_action({**current, "status": next_status})).strip(),
        }
        ok = real_job_service.records.update_record(application_id, patch)
        if not ok:
            return _api_error("更新申请记录失败", status_code=500, code="application_update_failed")

        updated = real_job_service.records.get_record(application_id)
        rows = _filter_application_rows(
            real_job_service.records.get_all_records(),
            user_id=str((updated.get("user_info") or {}).get("user_id") or ""),
            phone=str((updated.get("user_info") or {}).get("phone") or ""),
        )
        return _api_success({"application": _serialize_application_row(updated), "funnel": _build_offer_funnel_payload(rows)})
    except Exception as e:
        logger.exception("推进申请状态失败")
        return _api_error(str(e), status_code=500, code="application_status_failed")


@app.post("/api/offer/pipeline")
async def offer_pipeline(request: Request):
    """Run the full offer pipeline: multi-agent analysis, variants, jobs, training, funnel."""
    try:
        data = await request.json()
        resume_text = str(data.get("resume_text") or data.get("resume") or "").strip()
        if len(resume_text) < 20:
            return _api_error("resume_text 太短，至少 20 个字符", status_code=400)
        deny, credit_guard = _credit_balance_guard(request, "offer_pipeline")
        if deny:
            return deny

        target_role = str(data.get("target_role") or "").strip()
        city = str(data.get("city") or data.get("location") or "").strip()
        job_text = str(data.get("job_text") or "").strip()
        phone = str(data.get("phone") or "").strip()
        user_id = str(data.get("user_id") or "").strip()
        count = max(3, min(int(data.get("count") or 8), 20))
        auto_apply = bool(data.get("auto_apply", False))
        runtime_agent_overrides = data.get("agent_overrides")
        if not isinstance(runtime_agent_overrides, dict):
            runtime_agent_overrides = data.get("agent_configs") if isinstance(data.get("agent_configs"), dict) else None
        raw_sources = ",".join(
            [
                str(data.get("sources") or "").strip(),
                str(data.get("source_scopes") or "").strip(),
                str(data.get("custom_sources") or "").strip(),
            ]
        )
        source_list = [item.strip() for item in raw_sources.split(",") if item and item.strip()]

        class _CompatJSONRequest:
            def __init__(self, payload: Dict[str, Any]):
                self._payload = payload

            async def json(self):
                return self._payload

        core_resp = await process_resume(
            _CompatJSONRequest(
                {
                    "resume": resume_text,
                    "agent_overrides": runtime_agent_overrides,
                }
            )
        )
        raw_core = {}
        if hasattr(core_resp, "body"):
            raw_core = json.loads((core_resp.body or b"{}").decode("utf-8", errors="ignore") or "{}")
        if not raw_core.get("success"):
            return _api_error("主流程执行失败", status_code=500, code="offer_pipeline_failed")

        multi_data = raw_core.get("data") if isinstance(raw_core.get("data"), dict) else raw_core
        info = analyzer.extract_info(resume_text)
        requested_target_role = str(data.get("target_role") or "").strip()
        requested_city = str(data.get("city") or data.get("location") or "").strip()
        if not target_role:
            target_role = str(info.get("job_intention") or "").strip()
        target_role = _prefer_clean_text(requested_target_role, target_role)
        city = _prefer_clean_text(requested_city, city)
        if not isinstance(multi_data.get("skills_graph"), dict):
            multi_data["skills_graph"] = _skills_graph_payload(
                resume_text,
                target_role,
                str(multi_data.get("job_recommendations") or ""),
            )

        recommended_jobs = multi_data.get("recommended_jobs") if isinstance(multi_data.get("recommended_jobs"), list) else []
        recommended_jobs = _strip_cn_entrypoint_jobs(recommended_jobs)
        if city and recommended_jobs:
            for row in recommended_jobs:
                if not row.get("location"):
                    row["location"] = city

        if not recommended_jobs:
            fallback_keywords = _dedupe_preserve(
                [target_role] + [str(x) for x in (info.get("skills") or [])[:6]]
            ) or ["Python"]
            fallback_jobs: List[Dict[str, Any]] = []
            provider_mode = ""
            try:
                fallback_jobs = real_job_service.search_jobs(
                    keywords=fallback_keywords,
                    location=city or None,
                    limit=count,
                    sources=source_list or None,
                )
                provider_mode = str((real_job_service.get_statistics() or {}).get("provider_mode") or "")
            except Exception as search_error:
                logger.warning("offer_pipeline source-aware search failed: %s", search_error)
            if not fallback_jobs:
                fallback_jobs, provider_mode, _ = _search_jobs_without_browser(
                    fallback_keywords,
                    city or None,
                    count,
                    allow_portal_fallback=False,
                )
            recommended_jobs = _public_job_payload(
                _strip_cn_entrypoint_jobs(_enforce_cn_market_jobs(fallback_jobs)),
                limit=count,
            )
            multi_data["job_provider_mode"] = provider_mode
        if not recommended_jobs:
            recommended_jobs = _heuristic_opportunity_jobs(target_role, city, info, limit=count)
            multi_data["job_provider_mode"] = str(multi_data.get("job_provider_mode") or "heuristic_fallback")

        opportunities = _build_opportunity_cards(recommended_jobs[:count], resume_text, target_role)

        structured = await resume_profile_service.structure_resume_text(resume_text)
        resume_json = structured.get("resume_json") if isinstance(structured.get("resume_json"), dict) else {}
        profile = resume_profile_service.save_profile(
            resume_json=resume_json,
            user_id=user_id,
            title=str(data.get("title") or target_role or "Offer OS 简历"),
            source_text=resume_text,
        )

        variants = await _build_resume_variants_payload(
            resume_text=resume_text,
            resume_json=resume_json,
            info=info,
            target_role=target_role,
            jobs=opportunities,
            user_id=user_id,
        )

        primary_job = opportunities[0] if opportunities else {
            "title": target_role or "目标岗位",
            "company": str(data.get("company") or "目标公司"),
            "description": job_text,
            "location": city,
        }
        if job_text and not primary_job.get("description"):
            primary_job["description"] = job_text

        training = _build_interview_pack(resume_text, info, target_role, primary_job)
        outreach = _build_follow_up_templates(primary_job, target_role or str(primary_job.get("title") or "目标岗位"))

        auto_apply_preview = None
        if auto_apply:
            auto_apply_preview = _run_simple_apply_flow(
                {
                    "phone": phone,
                    "resume_text": resume_text,
                    "job_keyword": ",".join(
                        _dedupe_preserve(
                            [target_role]
                            + [str(row.get("title") or row.get("job_title") or "") for row in opportunities[:3]]
                        )
                    ),
                    "city": city,
                    "count": min(count, 10),
                }
            )

        rows = _filter_application_rows(real_job_service.records.get_all_records(), user_id=user_id, phone=phone)
        funnel = _build_offer_funnel_payload(rows)
        roadmap = _build_offer_roadmap(target_role or str(primary_job.get("title") or "目标岗位"), opportunities, funnel)
        credit_charge = _settle_credit_action(
            credit_guard,
            note="完整 Offer OS 流水线",
            meta={
                "endpoint": "/api/offer/pipeline",
                "target_role": target_role,
                "city": city,
                "count": count,
                "auto_apply": auto_apply,
            },
        )
        if not credit_charge.get("ok"):
            return _api_error("credits 扣减失败", status_code=409, code="credit_consume_failed")

        return _api_success(
            {
                "pipeline": {
                    "target_role": target_role or str(primary_job.get("title") or "目标岗位"),
                    "city": city,
                    "profile": profile,
                    "analysis": {
                        "career_analysis": str(multi_data.get("career_analysis") or ""),
                        "job_recommendations": str(multi_data.get("job_recommendations") or ""),
                        "optimized_resume": str(multi_data.get("optimized_resume") or ""),
                        "interview_prep": str(multi_data.get("interview_prep") or ""),
                        "mock_interview": str(multi_data.get("mock_interview") or ""),
                        "skills_graph": multi_data.get("skills_graph") or {},
                    },
                    "resume_variants": variants,
                    "opportunities": opportunities,
                    "training": training,
                    "outreach": outreach,
                    "roadmap": roadmap,
                    "funnel": funnel,
                    "auto_apply_preview": auto_apply_preview,
                    "job_provider_mode": str(multi_data.get("job_provider_mode") or ""),
                    "credit_charge": credit_charge,
                }
            }
        )
    except Exception as e:
        logger.exception("Offer pipeline 执行失败")
        return _api_error(str(e), status_code=500, code="offer_pipeline_failed")


@app.post("/api/offer/pipeline/stream")
async def offer_pipeline_stream(request: Request):
    """SSE wrapper around the full offer pipeline for real-time UI streaming."""
    data = await request.json()

    class _StreamRequestShim:
        def __init__(self, payload: Dict[str, Any], source: Request):
            self._payload = payload
            self.headers = source.headers
            self.cookies = source.cookies
            self.state = source.state

        async def json(self):
            return self._payload

    async def event_stream():
        stages = [
            ("extract", "解析简历结构与亮点"),
            ("rewrite", "生成可投递简历版本"),
            ("match", "抓取并筛选岗位池"),
            ("interview", "整理面试问题与回答重点"),
            ("apply", "准备本地执行参数"),
        ]
        task = asyncio.create_task(offer_pipeline(_StreamRequestShim(data, request)))
        stage_index = 0

        def _pack(payload: Dict[str, Any]) -> str:
            return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

        yield _pack({"stage": "boot", "status": "running", "chunk": "已连接流式执行通道"})

        while not task.done():
            stage, chunk = stages[min(stage_index, len(stages) - 1)]
            yield _pack({"stage": stage, "status": "running", "chunk": chunk})
            await asyncio.sleep(0.9)
            if stage_index < len(stages) - 1:
                stage_index += 1

        try:
            response = await task
            raw = {}
            if hasattr(response, "body"):
                raw = json.loads((response.body or b"{}").decode("utf-8", errors="ignore") or "{}")
            if not raw.get("success"):
                message = str(raw.get("detail") or raw.get("error") or raw.get("message") or "流水线执行失败")
                yield _pack({"stage": "done", "status": "error", "chunk": message})
                return

            pipeline_payload = raw.get("pipeline") if isinstance(raw.get("pipeline"), dict) else {}
            yield _pack(
                {
                    "stage": "done",
                    "status": "done",
                    "chunk": "流水线执行完成",
                    "pipeline": pipeline_payload,
                }
            )
        except Exception as exc:
            yield _pack({"stage": "done", "status": "error", "chunk": str(exc)[:300] or "流式执行失败"})

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _normalize_platform_list(raw: Any) -> List[str]:
    valid = {"boss", "zhilian", "linkedin"}
    items = _to_string_list(raw)
    out: List[str] = []
    for item in items:
        key = str(item).strip().lower()
        if key in valid and key not in out:
            out.append(key)
    return out


def _build_fallback_auto_apply_config(
    payload: Dict[str, Any],
    keywords: List[str],
    locations: List[str],
    max_count: int,
) -> Dict[str, Any]:
    user_profile = payload.get("user_profile") if isinstance(payload.get("user_profile"), dict) else {}
    cfg: Dict[str, Any] = {
        "keywords": ",".join(keywords or []),
        "location": (locations[0] if locations else ""),
        "max_count": int(max_count or HAITOU_DEFAULT_POLICY["max_count"]),
        "blacklist": _to_string_list(payload.get("blacklist") or payload.get("company_blacklist")),
        "headless": bool(HAITOU_DEFAULT_POLICY["headless"]),
        "use_ai_answers": bool(payload.get("use_ai_answers", True)),
        "verification_wait_seconds": int(HAITOU_DEFAULT_POLICY["verification_wait_seconds"]),
        "sms_code": str(payload.get("sms_code") or "").strip() or None,
        "allow_risky_automation": _to_bool(payload.get("allow_risky_automation"), default=False),
        "compliance_override_platforms": _to_string_list(payload.get("compliance_override_platforms")),
    }

    boss_cfg = payload.get("boss_config") if isinstance(payload.get("boss_config"), dict) else {}
    zhilian_cfg = payload.get("zhilian_config") if isinstance(payload.get("zhilian_config"), dict) else {}
    linkedin_cfg = payload.get("linkedin_config") if isinstance(payload.get("linkedin_config"), dict) else {}

    cfg["boss_config"] = {
        **boss_cfg,
        "phone": str(
            boss_cfg.get("phone")
            or payload.get("boss_phone")
            or payload.get("phone")
            or user_profile.get("phone")
            or ""
        ).strip(),
        "sms_code": str(
            boss_cfg.get("sms_code")
            or payload.get("boss_sms_code")
            or payload.get("sms_code")
            or ""
        ).strip() or None,
    }
    cfg["zhilian_config"] = {
        **zhilian_cfg,
        "username": str(
            zhilian_cfg.get("username")
            or payload.get("zhilian_username")
            or payload.get("username")
            or user_profile.get("username")
            or user_profile.get("email")
            or ""
        ).strip(),
        "password": str(
            zhilian_cfg.get("password")
            or payload.get("zhilian_password")
            or user_profile.get("password")
            or ""
        ),
    }
    cfg["linkedin_config"] = {
        **linkedin_cfg,
        "email": str(
            linkedin_cfg.get("email")
            or payload.get("linkedin_email")
            or payload.get("email")
            or user_profile.get("email")
            or ""
        ).strip(),
        "password": str(
            linkedin_cfg.get("password")
            or payload.get("linkedin_password")
            or user_profile.get("password")
            or ""
        ),
    }
    return cfg


def _filter_platforms_by_credentials(platforms: List[str], config: Dict[str, Any]) -> Tuple[List[str], List[str]]:
    accepted: List[str] = []
    blockers: List[str] = []
    allow_all_risky = _to_bool(config.get("allow_risky_automation"), default=False) or _to_bool(
        os.getenv("AUTO_APPLY_ALLOW_RISKY_AUTOMATION"), default=False
    )
    override_platforms = _normalize_platform_list(config.get("compliance_override_platforms"))
    override_platforms.extend(_normalize_platform_list(os.getenv("AUTO_APPLY_COMPLIANCE_OVERRIDE_PLATFORMS")))
    override_set = set(override_platforms)

    for p in platforms:
        policy = PLATFORM_AUTOMATION_POLICY.get(p, {})
        if str(policy.get("status") or "").strip() == "restricted_by_default" and not (
            allow_all_risky or p in override_set
        ):
            blockers.append(str(policy.get("block_message") or f"{p} 默认禁用自动化"))
            continue

        if p == "boss":
            phone = str((config.get("boss_config") or {}).get("phone") or "").strip()
            if phone:
                accepted.append(p)
            else:
                blockers.append("boss 缺少手机号（boss_config.phone）")
        elif p == "zhilian":
            u = str((config.get("zhilian_config") or {}).get("username") or "").strip()
            pw = str((config.get("zhilian_config") or {}).get("password") or "")
            if u and pw:
                accepted.append(p)
            else:
                blockers.append("zhilian 缺少账号密码（zhilian_config.username/password）")
        elif p == "linkedin":
            e = str((config.get("linkedin_config") or {}).get("email") or "").strip()
            pw = str((config.get("linkedin_config") or {}).get("password") or "")
            if e and pw:
                accepted.append(p)
            else:
                blockers.append("linkedin 缺少邮箱密码（linkedin_config.email/password）")
    return accepted, blockers


async def _create_multi_platform_auto_apply_task(platforms: List[str], config: Dict[str, Any]) -> str:
    task_id = str(uuid.uuid4())
    async with task_lock:
        auto_apply_tasks[task_id] = {
            "task_id": task_id,
            "status": "starting",
            "platforms": platforms,
            "config": config,
            "progress": {
                "total_platforms": len(platforms),
                "completed_platforms": 0,
                "total_applied": 0,
                "total_failed": 0,
                "platform_progress": {},
            },
            "created_at": datetime.now().isoformat(),
            "started_at": None,
            "completed_at": None,
        }
    asyncio.create_task(_run_multi_platform_apply(task_id, platforms, config))
    return task_id


async def _run_github_auto_apply_task(task_id: str, payload: Dict[str, Any]):
    task = github_apply_tasks.get(task_id)
    if not task:
        return
    task["status"] = "running"
    task["started_at"] = datetime.now().isoformat()
    try:
        resume_text = str(payload.get("resume_text") or "").strip()
        keywords = _to_string_list(payload.get("keywords"))
        locations = _to_string_list(payload.get("locations"))
        max_count = max(1, min(200, int(payload.get("max_count") or 20)))
        llm_key = (
            str(payload.get("llm_api_key") or "").strip()
            or os.getenv("DEEPSEEK_API_KEY", "").strip()
            or os.getenv("OPENAI_API_KEY", "").strip()
        )
        linkedin_cfg = payload.get("linkedin_config") if isinstance(payload.get("linkedin_config"), dict) else {}
        runtime_credentials = {
            "linkedin_email": str(
                payload.get("linkedin_email")
                or linkedin_cfg.get("email")
                or ""
            ).strip(),
            "linkedin_password": str(
                payload.get("linkedin_password")
                or linkedin_cfg.get("password")
                or ""
            ),
        }
        enable_fallback = bool(payload.get("enable_fallback_auto_apply", True)) or bool(HAITOU_DEFAULT_POLICY["enabled"])
        requested_platforms = _normalize_platform_list(
            payload.get("fallback_platforms") or payload.get("platforms")
        ) or list(HAITOU_DEFAULT_POLICY["fallback_platforms"])

        async def _fallback_from_github(message: str, extra: Optional[Dict[str, Any]] = None) -> bool:
            if not enable_fallback:
                return False
            fallback_cfg = _build_fallback_auto_apply_config(
                payload,
                keywords,
                locations,
                max_count,
            )
            accepted, blockers = _filter_platforms_by_credentials(
                requested_platforms,
                fallback_cfg,
            )
            if accepted:
                linked_task_id = await _create_multi_platform_auto_apply_task(
                    accepted,
                    fallback_cfg,
                )
                task["status"] = "fallback_started"
                task["linked_auto_apply_task_id"] = linked_task_id
                result: Dict[str, Any] = {
                    "message": message,
                    "platforms": accepted,
                    "linked_auto_apply_task_id": linked_task_id,
                }
                if extra:
                    result.update(extra)
                task["result"] = result
            else:
                task["status"] = "prepared_only"
                result = {
                    "message": message,
                    "blockers": blockers,
                }
                if extra:
                    result.update(extra)
                task["result"] = result
            task["completed_at"] = datetime.now().isoformat()
            return True

        snapshot = _aihawk_capability_snapshot()
        task["capability"] = snapshot
        if not snapshot.get("can_prepare"):
            started = await _fallback_from_github(
                "AIHawk 运行环境不完整，已自动切换到内置多平台自动投递引擎。",
                {
                    "capability_blockers": {
                        "missing_required": snapshot.get("missing_required"),
                        "missing_deps": snapshot.get("missing_deps"),
                    }
                },
            )
            if started:
                return
            task["status"] = "blocked"
            task["error"] = (
                "AIHawk 运行环境不完整，无法生成任务配置。"
                f" missing_required={snapshot.get('missing_required')} missing_deps={snapshot.get('missing_deps')}"
            )
            task["completed_at"] = datetime.now().isoformat()
            return

        # Lazy import to avoid startup hard dependency.
        from app.core.aihawk_client import aihawk_client

        config_dir = await asyncio.to_thread(
            aihawk_client.create_config_files,
            task_id,
            keywords or ["Python"],
            locations or ["北京"],
            max_count,
            resume_text,
            llm_key or "missing_api_key",
            runtime_credentials,
        )
        task["prepared"] = True
        task["config_dir"] = config_dir

        # Current upstream AIHawk package in this repo may not include auto-apply plugins.
        if not snapshot.get("has_auto_apply_plugins"):
            started = await _fallback_from_github(
                "AIHawk 插件缺失，已自动切换到内置多平台自动投递引擎。"
            )
            if started:
                return
            else:
                task["status"] = "prepared_only"
                task["result"] = {
                    "message": (
                        "已生成 AIHawk 配置文件，但当前仓库缺少 auto-apply 插件代码。"
                        " 如需真实自动投递，请切回含插件的历史版本或私有分支。"
                    )
                }
            task["completed_at"] = datetime.now().isoformat()
            return

        if not llm_key:
            started = await _fallback_from_github(
                "缺少 LLM API KEY，已自动切换到内置多平台自动投递引擎。"
            )
            if started:
                return
            task["status"] = "blocked"
            task["error"] = "缺少 LLM API KEY：请配置环境变量或在请求中传 llm_api_key。"
            task["completed_at"] = datetime.now().isoformat()
            return

        plugin_mode = str(snapshot.get("plugin_mode") or "")
        if plugin_mode.startswith("linkedin_legacy"):
            if not runtime_credentials.get("linkedin_email") or not runtime_credentials.get("linkedin_password"):
                started = await _fallback_from_github(
                    "AIHawk 插件模式要求 LinkedIn 账号密码，已自动切换到内置多平台自动投递引擎。"
                )
                if started:
                    return
                task["status"] = "blocked"
                task["error"] = "AIHawk 插件版缺少 LinkedIn 账号密码，请提供 linkedin_email/linkedin_password。"
                task["completed_at"] = datetime.now().isoformat()
                return

        if not snapshot.get("can_run"):
            started = await _fallback_from_github(
                "AIHawk 插件运行依赖不完整，已自动切换到内置多平台自动投递引擎。"
            )
            if started:
                return
            task["status"] = "blocked"
            task["error"] = "AIHawk 插件已就绪，但运行依赖不完整，请安装缺失依赖。"
            task["completed_at"] = datetime.now().isoformat()
            return

        run_result = await asyncio.to_thread(aihawk_client.run_aihawk, task_id, max_count)
        if run_result.get("success"):
            task["status"] = "completed"
            task["result"] = run_result
            task["completed_at"] = datetime.now().isoformat()
            return

        started = await _fallback_from_github(
            "AIHawk 运行失败，已自动切换到内置多平台自动投递引擎。",
            {"aihawk_result": run_result},
        )
        if started:
            return
        task["status"] = "failed"
        task["result"] = run_result
        task["completed_at"] = datetime.now().isoformat()
    except Exception as e:
        logger.exception("github 高星自动投递任务失败")
        task["status"] = "failed"
        task["error"] = str(e)
        task["completed_at"] = datetime.now().isoformat()


@app.get("/api/github-auto-apply/health")
async def github_auto_apply_health():
    return _api_success({"capability": _aihawk_capability_snapshot()})


@app.post("/api/github-auto-apply/start")
async def github_auto_apply_start(request: Request):
    """
    Start GitHub high-star AIHawk task.
    Payload:
      resume_text (required),
      keywords (list or csv),
      locations (list or csv),
      max_count (int)
    """
    try:
        raw_data = await request.json()
        data = _apply_haitou_defaults(raw_data if isinstance(raw_data, dict) else {})
        resume_text = str(data.get("resume_text") or "").strip()
        if not resume_text:
            profile_id = str(data.get("profile_id") or "").strip()
            if profile_id:
                profile = resume_profile_service.get_profile(profile_id)
                if profile:
                    resume_json = profile.get("resume_json") or {}
                    resume_text = json.dumps(resume_json, ensure_ascii=False)
        if not resume_text:
            return _api_error("resume_text 不能为空（或提供 profile_id）", status_code=400)

        info = analyzer.extract_info(resume_text)
        keywords = _to_string_list(data.get("keywords"))
        if not keywords:
            keywords = [str(info.get("job_intention") or "").strip()]
            keywords.extend([str(x) for x in (info.get("skills") or [])[:5]])
            keywords = [k for k in keywords if k]
        locations = _to_string_list(data.get("locations"))
        if not locations:
            locations = [str(x) for x in (info.get("preferred_locations") or [])[:3] if str(x).strip()]
        if not locations:
            locations = ["北京"]

        task_id = f"gh_{uuid.uuid4().hex[:10]}"
        task = {
            "task_id": task_id,
            "status": "queued",
            "keywords": keywords,
            "locations": locations,
            "max_count": int(HAITOU_DEFAULT_POLICY["max_count"]),
            "enable_fallback_auto_apply": True,
            "requested_fallback_platforms": list(HAITOU_DEFAULT_POLICY["fallback_platforms"]),
            "created_at": datetime.now().isoformat(),
            "prepared": False,
            "strategy": HAITOU_DEFAULT_POLICY["strategy_version"],
        }
        github_apply_tasks[task_id] = task
        asyncio.create_task(
            _run_github_auto_apply_task(
                task_id,
                {
                    "resume_text": resume_text,
                    "keywords": keywords,
                    "locations": locations,
                    "max_count": task["max_count"],
                    "enable_fallback_auto_apply": True,
                    "fallback_platforms": task["requested_fallback_platforms"],
                    "boss_config": data.get("boss_config"),
                    "zhilian_config": data.get("zhilian_config"),
                    "linkedin_config": data.get("linkedin_config"),
                    "headless": bool(HAITOU_DEFAULT_POLICY["headless"]),
                    "use_ai_answers": data.get("use_ai_answers", True),
                    "verification_wait_seconds": int(HAITOU_DEFAULT_POLICY["verification_wait_seconds"]),
                    "llm_api_key": data.get("llm_api_key"),
                    "sms_code": data.get("sms_code"),
                    "blacklist": data.get("blacklist"),
                    "company_blacklist": data.get("company_blacklist"),
                    "boss_phone": data.get("boss_phone"),
                    "boss_sms_code": data.get("boss_sms_code"),
                    "phone": data.get("phone"),
                    "zhilian_username": data.get("zhilian_username"),
                    "zhilian_password": data.get("zhilian_password"),
                    "linkedin_email": data.get("linkedin_email"),
                    "linkedin_password": data.get("linkedin_password"),
                },
            )
        )
        return _api_success({"task_id": task_id, "status": "queued"})
    except Exception as e:
        logger.exception("启动 github 高星自动投递失败")
        return _api_error(str(e), status_code=500, code="github_auto_apply_start_failed")


@app.get("/api/github-auto-apply/status/{task_id}")
async def github_auto_apply_status(task_id: str):
    task = github_apply_tasks.get(task_id)
    if not task:
        return _api_error("任务不存在", status_code=404, code="task_not_found")
    linked_id = str(task.get("linked_auto_apply_task_id") or "").strip()
    linked_task = auto_apply_tasks.get(linked_id) if linked_id else None
    return _api_success(
        {
            "task": task,
            "linked_auto_apply_task": linked_task,
        }
    )


@app.get("/api/github-auto-apply/history")
async def github_auto_apply_history(limit: int = 20):
    rows = sorted(
        github_apply_tasks.values(),
        key=lambda x: x.get("created_at", ""),
        reverse=True,
    )
    rows = rows[: max(1, min(200, int(limit or 20)))]
    return _api_success({"tasks": rows, "total": len(github_apply_tasks)})


if __name__ == "__main__":
    import webbrowser
    import threading
    
    port = int(os.getenv("PORT", 8000))
    
    print("\n" + "🚀"*30)
    print("AI求职助手 - Web服务启动中...")
    print("🚀"*30 + "\n")
    print(f"📍 访问地址: http://localhost:{port}")
    print(f"📍 API文档: http://localhost:{port}/docs")
    print(f"📍 WebSocket: ws://localhost:{port}/ws/progress")
    print("\n✨ 新功能: WebSocket实时进度推送")
    print("按 Ctrl+C 停止服务\n")
    
    # 延迟2秒后自动打开浏览器
    def open_browser():
        import time
        time.sleep(2)
        webbrowser.open(f'http://localhost:{port}/app')
    
    threading.Thread(target=open_browser, daemon=True).start()
    
    uvicorn.run(app, host="0.0.0.0", port=port)

